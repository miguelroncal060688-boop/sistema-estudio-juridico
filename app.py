
import streamlit as st
import pandas as pd
import os
import hashlib
import shutil
from datetime import date, datetime

# ==========================================================
# FUNCIONES BASE
# ==========================================================
def cargar_csv(nombre):
    if os.path.exists(nombre):
        return pd.read_csv(nombre)
    return pd.DataFrame()

def guardar_csv(df, nombre):
    df.to_csv(nombre, index=False)


# ==========================================================
# ALIAS DE FUNCIONES (COMPATIBILIDAD)
# ==========================================================
def load_df(nombre):
    return cargar_csv(f"{nombre}.csv")

def save_df(nombre, df):
    guardar_csv(df, f"{nombre}.csv")

def normalize_key(x):
    return str(x).strip().upper()

# ==========================================================
# BLOQUE A — VISIBILIDAD POR ROL (CASOS)
# ==========================================================
def _rol_actual():
    return str(st.session_state.get("rol","")).strip().lower()

def _usuario_actual():
    return str(
        st.session_state.get("usuario")
        or st.session_state.get("Usuario")
        or st.session_state.get("user")
        or ""
    ).strip()

def _nombre_abogado_del_usuario():
    """
    Devuelve el NOMBRE del abogado asociado al usuario (si rol=Abogado),
    usando usuarios.AbogadoID -> abogados.ID -> abogados.Nombre
    """
    try:
        usuario = _usuario_actual()
        if not usuario:
            return ""

        df_u = load_df("usuarios")
        if df_u.empty or "Usuario" not in df_u.columns:
            return ""

        fila = df_u[df_u["Usuario"].astype(str) == usuario]
        if fila.empty:
            return ""

        abogado_id = str(fila.iloc[0].get("AbogadoID","")).strip()
        if not abogado_id:
            return ""

        df_a = load_df("abogados")
        if df_a.empty or "ID" not in df_a.columns:
            return ""

        match = df_a[df_a["ID"].astype(str) == abogado_id]
        if match.empty:
            return ""

        if "Nombre" in match.columns:
            return str(match.iloc[0]["Nombre"]).strip()

        return ""
    except Exception:
        return ""

def filtrar_casos_por_rol(df_casos):
    """
    Reglas:
    - Admin / Personal Administrativo: ve todo
    - Abogado: ve casos donde:
        * es Abogado principal
        * está en AbogadosExtra (defensa conjunta)
        * está en Delegados
    - Secretaria/o / Asistente: solo casos Delegados
    """
    if df_casos is None or df_casos.empty:
        return df_casos

    rol = _rol_actual()
    usuario = _usuario_actual()

    # Roles con vista total
    if rol in ["admin", "personal administrativo"]:
        return df_casos

    df = df_casos.copy()

    for col in ["Abogado", "AbogadosExtra", "Delegados"]:
        if col not in df.columns:
            df[col] = ""
        df[col] = df[col].astype(str)

    # Abogado
    if rol == "abogado":
        nombre_abogado = _nombre_abogado_del_usuario()

        m_deleg = df["Delegados"].str.contains(usuario, case=False, na=False) if usuario else False

        if nombre_abogado:
            m1 = df["Abogado"].str.contains(nombre_abogado, case=False, na=False)
            m2 = df["AbogadosExtra"].str.contains(nombre_abogado, case=False, na=False)
            return df[m1 | m2 | m_deleg].copy()

        return df[m_deleg].copy() if isinstance(m_deleg, pd.Series) else df.iloc[0:0].copy()

    # Secretaria / Asistente
    if rol in ["secretaria", "secretaria/o", "asistente"]:
        if not usuario:
            return df.iloc[0:0].copy()
        return df[df["Delegados"].str.contains(usuario, case=False, na=False)].copy()

    # Otros roles: conservador
    if not usuario:
        return df.iloc[0:0].copy()
    return df[df["Delegados"].str.contains(usuario, case=False, na=False)].copy()

# ==========================================================
# MARCA 004 – VERSIÓN ESTABLE OPERATIVA
# Estado: FUNCIONA TODO – NO MODIFICAR NI REDUCIR
# Añadidos: Edit/Borrar Honorarios + Ficha/Historial Actuaciones (con link OneDrive)
#           + Módulo Consultas (con proforma e historial)
# ==========================================================
APP_VERSION = "MARCA 006"  # Integración completa 1-19

# ==========================================================
# CONFIGURACIÓN GENERAL
# ==========================================================
APP_NAME = "Estudio Jurídico Roncal Liñan y Asociados"
CONTROL_PASSWORD = st.secrets.get("CONTROL_PASSWORD", "control123")  # clave panel de control
ADMIN_BOOTSTRAP_PASSWORD = st.secrets.get("ADMIN_BOOTSTRAP_PASSWORD", "estudio123")
PASSWORD_PEPPER = st.secrets.get("PASSWORD_PEPPER", "")


DATA_DIR = "."
BACKUP_DIR = os.path.join(DATA_DIR, "backups")
UPLOADS_DIR = os.path.join(DATA_DIR, "uploads")
GENERADOS_DIR = os.path.join(DATA_DIR, "generados")

os.makedirs(BACKUP_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(GENERADOS_DIR, exist_ok=True)

st.set_page_config(page_title=f"⚖️ {APP_NAME} – {APP_VERSION}", layout="wide")

# ==========================================================
# ARCHIVOS
# ==========================================================
FILES = {
    "usuarios": "usuarios.csv",
    "clientes": "clientes.csv",
    "abogados": "abogados.csv",
    "casos": "casos.csv",

    "honorarios": "honorarios.csv",
    "honorarios_etapas": "honorarios_etapas.csv",

    "pagos_honorarios": "pagos_honorarios.csv",
    "cuota_litis": "cuota_litis.csv",
    "pagos_litis": "pagos_litis.csv",

    "cuotas": "cuotas.csv",
    "actuaciones": "actuaciones.csv",
    "documentos": "documentos.csv",
    "plantillas": "plantillas_contratos.csv",

    # ✅ NUEVO: consultas
    "consultas": "consultas.csv",
    "instancias": "instancias.csv",
    "honorarios_tipo": "honorarios_tipo.csv",
    "contratos": "contratos.csv",
    "auditoria_mod": "auditoria_mod.csv",

    # ✅ NUEVO (NECESARIO PARA ROLES Y PERMISOS)
    "permisos": "permisos.csv",
}
# ==========================================================
# ESQUEMAS
# ==========================================================
SCHEMAS = {
    # ======================
    # USUARIOS / SEGURIDAD
    # ======================
    "usuarios": ["Usuario","PasswordHash","Rol","AbogadoID","Activo","Creado"],

    # ✅ NUEVO – ROLES Y PERMISOS
    "permisos": ["Rol","Ver","Agregar","Modificar","Borrar"],

    # ======================
    # ENTIDADES PRINCIPALES
    # ======================
    "clientes": [
        "ID","TipoCliente","Nombre","DNI","Celular","Correo","Direccion",
        "Observaciones","ContactoEmergencia","CelularEmergencia",
        "RazonSocial","RUC","RepresentanteLegal",
        "PartidaElectronica","SedeRegistral"
    ],

    "abogados": [
        "ID","Nombre","DNI","Celular","Correo","Colegiatura",
        "ColegioProfesional","Domicilio Procesal","ReferenciaDomicilio",
        "Casilla Electronica","DistritoJudicial","Casilla Judicial","Notas"
    ],

    "casos": [
        "ID","Cliente","Abogado","Expediente","Año","Materia","Instancia",
        "Pretension","Juzgado","DistritoJudicial",
        "Contraparte","ContraparteDoc",
        "Observaciones","EstadoCaso","FechaInicio"
    ],

    # ======================
    # HONORARIOS / PAGOS
    # ======================
    "honorarios": ["ID","Caso","Monto Pactado","Notas","FechaRegistro"],
    "honorarios_etapas": ["ID","Caso","Etapa","Monto Pactado","Notas","FechaRegistro"],
    "pagos_honorarios": ["ID","Caso","Etapa","FechaPago","Monto","Observacion"],

    "cuota_litis": ["ID","Caso","Monto Base","Porcentaje","Notas","FechaRegistro"],
    "pagos_litis": ["ID","Caso","FechaPago","Monto","Observacion"],
    "cuotas": ["ID","Caso","Tipo","NroCuota","FechaVenc","Monto","Notas"],

    # ======================
    # ACTUACIONES / DOCUMENTOS
    # ======================
    "actuaciones": [
        "ID","Caso","Cliente","Fecha","TipoActuacion","Resumen",
        "ProximaAccion","FechaProximaAccion","LinkOneDrive",
        "CostasAranceles","Gastos","Notas","GastosPagado"
    ],

    "documentos": ["ID","Caso","Tipo","NombreArchivo","Ruta","Fecha","Notas"],
    "plantillas": ["ID","Nombre","Contenido","Notas","Creado"],

    # ======================
    # OTROS MÓDULOS
    # ======================
    "consultas": [
        "ID","Fecha","Cliente","Caso","Abogado","Consulta","Estrategia",
        "CostoConsulta","HonorariosPropuestos","Proforma",
        "LinkOneDrive","Notas"
    ],

    "instancias": [
        "ID","Caso","TipoInstancia","EstadoActual",
        "Resultado","Accion","Honorarios","FechaRegistro"
    ],

    "honorarios_tipo": ["ID","Caso","Tipo","Monto","Notas","FechaRegistro"],

    "contratos": [
        "ID","Numero","Año","Sigla","NombreContrato",
        "Caso","Cliente","Abogado","Estado","Archivo","Fecha"
    ],

    "auditoria_mod": [
        "ID","Fecha","Usuario","Rol","Accion",
        "Entidad","EntidadID","Detalle"
    ],
}

# ============================
# CONSTANTES
# ============================
ETAPAS_HONORARIOS = ["Primera instancia", "Segunda instancia", "Casación", "Otros"]
TIPOS_CUOTA = ["Honorarios", "CuotaLitis"]

# ==========================================================
# UTILIDADES
# ==========================================================
def sha256(text: str) -> str:
    base = (PASSWORD_PEPPER or "") + str(text)
    return hashlib.sha256(base.encode("utf-8")).hexdigest()

def normalize_key(x) -> str:
    if pd.isna(x):
        return ""
    return str(x).strip()

def drop_unnamed(df: pd.DataFrame) -> pd.DataFrame:
    return df.loc[:, ~df.columns.str.contains(r"^Unnamed", case=False, na=False)]

def safe_float_series(s):
    return pd.to_numeric(s, errors="coerce").fillna(0.0)

def money(x):
    try:
        return float(x)
    except Exception:
        return 0.0

def to_date_safe(x):
    if pd.isna(x) or str(x).strip()=="":
        return None
    try:
        return pd.to_datetime(x).date()
    except Exception:
        return None

def backup_file(path: str):
    if not os.path.exists(path):
        return
    base = os.path.basename(path)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    dst = os.path.join(BACKUP_DIR, f"{base}.{stamp}.bak")
    try:
        shutil.copy2(path, dst)
    except Exception:
        pass

# ==========================================================
# ensure_csv (NO destruye datos; guarda corrupto antes)
# ==========================================================
def ensure_csv(key: str):
    path = FILES[key]
    cols = SCHEMAS.get(key, ["ID"])

    if not os.path.exists(path):
        pd.DataFrame(columns=cols).to_csv(path, index=False)
        return

    try:
        if os.path.getsize(path) == 0:
            pd.DataFrame(columns=cols).to_csv(path, index=False)
            return
    except OSError:
        pass

    try:
        df = pd.read_csv(path)
    except pd.errors.EmptyDataError:
        try:
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            shutil.copy2(path, os.path.join(BACKUP_DIR, f"{os.path.basename(path)}.{stamp}.corrupt.bak"))
        except Exception:
            pass
        pd.DataFrame(columns=cols).to_csv(path, index=False)
        return

    df = drop_unnamed(df)

    for c in cols:
        if c not in df.columns:
            df[c] = ""

    df = df.reindex(columns=cols)
    df.to_csv(path, index=False)

# ==========================================================
# load_df (migraciones suaves)
# ==========================================================
def load_df(key: str) -> pd.DataFrame:
    ensure_csv(key)
    try:
        df = pd.read_csv(FILES[key])
    except pd.errors.EmptyDataError:
        df = pd.DataFrame(columns=SCHEMAS[key])

    df = drop_unnamed(df)

    # Migración actuaciones si venían con nombres antiguos
    if key == "actuaciones":
        rename_map = {
            "ActuaciónID": "ID",
            "CasoID": "Caso",
            "TipoActuación": "TipoActuacion",
            "PróximaAcción": "ProximaAccion",
            "FechaPróximaAcción": "FechaProximaAccion",
            "Link": "LinkOneDrive",
            "Enlace": "LinkOneDrive",
        }
        df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns}, inplace=True)

    # Migración pagos_honorarios si era esquema antiguo sin Etapa
    if key == "pagos_honorarios":
        if "Etapa" not in df.columns:
            df["Etapa"] = ""

    # Migración casos: si no existía Instancia
    if key == "casos":
        if "Instancia" not in df.columns:
            df["Instancia"] = ""

    # Migración consultas si existiera algo previo
    if key == "consultas":
        for col in SCHEMAS["consultas"]:
            if col not in df.columns:
                df[col] = ""

    df = df.reindex(columns=SCHEMAS[key])
    return df

def save_df(key: str, df: pd.DataFrame):
    backup_file(FILES[key])
    df = drop_unnamed(df)
    df = df.reindex(columns=SCHEMAS[key])
    df.to_csv(FILES[key], index=False)

def next_id(df: pd.DataFrame, col="ID") -> int:
    if df.empty:
        return 1
    m = pd.to_numeric(df[col], errors="coerce").max()
    return int(m) + 1 if pd.notna(m) else len(df) + 1

def ensure_ids(df: pd.DataFrame) -> pd.DataFrame:
    if "ID" not in df.columns:
        return df
    ids = pd.to_numeric(df["ID"], errors="coerce")
    max_id = int(ids.max()) if pd.notna(ids.max()) else 0
    for idx in df[ids.isna()].index.tolist():
        max_id += 1
        df.at[idx, "ID"] = max_id
    return df

def add_row(df: pd.DataFrame, row_dict: dict, schema_key: str) -> pd.DataFrame:
    df2 = pd.concat([df, pd.DataFrame([row_dict])], ignore_index=True)
    df2 = df2.reindex(columns=SCHEMAS[schema_key])
    return df2

def brand_header():
    st.markdown(
        f"""
        <div style="padding:10px 0 0 0;">
            <h1 style="margin-bottom:0;">⚖️ {APP_NAME}</h1>
            <p style="margin-top:2px;color:#666;">Sistema de gestión de clientes, casos, pagos, cuotas, documentos y contratos</p>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.caption(f"Versión: {APP_VERSION}")

def require_admin():
    if st.session_state.get("rol") != "admin":
        st.error("❌ Solo ADMIN puede acceder aquí.")
        st.stop()

# ==========================================================
# INICIALIZACIÓN DE ARCHIVOS
# ==========================================================
for k in FILES:
    ensure_csv(k)

# ==========================================================
# ASEGURAR ADMIN (sin borrar nada)
# ==========================================================
usuarios = load_df("usuarios")
if usuarios[usuarios["Usuario"].astype(str) == "admin"].empty:
    usuarios = add_row(usuarios, {
        "Usuario": "admin",
        "PasswordHash": sha256(ADMIN_BOOTSTRAP_PASSWORD),
        "Rol": "admin",
        "AbogadoID": "",
        "Activo": "1",
        "Creado": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }, "usuarios")
    save_df("usuarios", usuarios)

# ==========================================================
# LOGIN
# ==========================================================
if "usuario" not in st.session_state:
    st.session_state.usuario = None
if "rol" not in st.session_state:
    st.session_state.rol = None
if "abogado_id" not in st.session_state:
    st.session_state.abogado_id = ""

def login_ui():
    brand_header()
    st.subheader("Ingreso al Sistema")

    u = st.text_input("Usuario", value="")
    p = st.text_input("Contraseña", type="password", value="")

    if st.button("Ingresar"):
        users = load_df("usuarios")
        users = users[users["Activo"].astype(str) == "1"].copy()
        row = users[users["Usuario"].astype(str) == str(u)].copy()
        if row.empty or row.iloc[0]["PasswordHash"] != sha256(p):
            st.error("Credenciales incorrectas")
            st.stop()

        st.session_state.usuario = u
        st.session_state.rol = row.iloc[0]["Rol"]
        st.session_state.abogado_id = str(row.iloc[0]["AbogadoID"]) if "AbogadoID" in row.columns else ""
        st.rerun()

if st.session_state.usuario is None:
    login_ui()
    st.stop()

# ==========================================================
# CARGA DATA (con IDs asegurados)
# ==========================================================
clientes = load_df("clientes")
abogados = load_df("abogados")
casos = load_df("casos")

honorarios = ensure_ids(load_df("honorarios"))
honorarios_etapas = ensure_ids(load_df("honorarios_etapas"))
pagos_honorarios = ensure_ids(load_df("pagos_honorarios"))

cuota_litis = ensure_ids(load_df("cuota_litis"))
pagos_litis = ensure_ids(load_df("pagos_litis"))

cuotas = ensure_ids(load_df("cuotas"))
actuaciones = ensure_ids(load_df("actuaciones"))
documentos = ensure_ids(load_df("documentos"))
plantillas = ensure_ids(load_df("plantillas"))
consultas = ensure_ids(load_df("consultas"))
usuarios = load_df("usuarios")

# normalizar claves
if "Expediente" in casos.columns:
    casos["Expediente"] = casos["Expediente"].apply(normalize_key)

for df in [honorarios, honorarios_etapas, pagos_honorarios, cuota_litis, pagos_litis, cuotas, actuaciones, documentos]:
    if "Caso" in df.columns:
        df["Caso"] = df["Caso"].apply(normalize_key)

# normalizar Tipo de cuotas si venía con espacios
if not cuotas.empty and "Tipo" in cuotas.columns:
    cuotas["Tipo"] = cuotas["Tipo"].astype(str).str.replace(" ", "", regex=False)
    cuotas["Tipo"] = cuotas["Tipo"].replace({"CuotaLitis":"CuotaLitis","Honorarios":"Honorarios"})
    save_df("cuotas", cuotas)

# guardar IDs reparados (NO borra nada)
save_df("honorarios", honorarios)
save_df("honorarios_etapas", honorarios_etapas)
save_df("pagos_honorarios", pagos_honorarios)
save_df("cuota_litis", cuota_litis)
save_df("pagos_litis", pagos_litis)
save_df("cuotas", cuotas)
save_df("actuaciones", actuaciones)
save_df("documentos", documentos)
save_df("plantillas", plantillas)
save_df("consultas", consultas)

# ==========================================================
# FINANZAS
# ==========================================================
def canon_last_by_case(df: pd.DataFrame, case_col="Caso"):
    if df.empty:
        return df
    tmp = df.copy()
    tmp[case_col] = tmp[case_col].apply(normalize_key)
    tmp["_id"] = pd.to_numeric(tmp["ID"], errors="coerce").fillna(0).astype(int)
    tmp.sort_values([case_col, "_id"], inplace=True)
    tmp = tmp.groupby(case_col, as_index=False).tail(1)
    tmp.drop(columns=["_id"], inplace=True, errors="ignore")
    return tmp


def gastos_actuaciones_por_caso():
    """
    Suma gastos de actuaciones por caso (INFORMATIVO).
    NO se consideran deuda general.
    """
    if actuaciones.empty:
        return pd.DataFrame(columns=["Caso", "GastosActuaciones"])

    df = actuaciones.copy()
    df["Caso"] = df["Caso"].apply(normalize_key)

    df["CostasAranceles"] = pd.to_numeric(df.get("CostasAranceles", 0), errors="coerce").fillna(0.0)
    df["Gastos"] = pd.to_numeric(df.get("Gastos", 0), errors="coerce").fillna(0.0)
    df["GastosActuaciones"] = df["CostasAranceles"] + df["Gastos"]

    out = df.groupby("Caso", as_index=False)["GastosActuaciones"].sum()
    return out


def resumen_financiero_df():
    """
    Resumen por caso (Expediente):

    ✅ El Saldo Total incluye SOLO:
       - Honorarios pendientes
       - Cuota litis pendiente

    ✅ Los gastos de actuaciones:
       - Se muestran como INFORMATIVOS
       - NO inflan deuda ni saldo total
    """
    if casos.empty:
        return pd.DataFrame(columns=[
            "Expediente","Cliente","Materia",
            "Honorario Pactado","Honorario Pagado","Honorario Pendiente",
            "Cuota Litis Calculada","Pagado Litis","Saldo Litis",
            "Gastos Actuaciones","Saldo Total"
        ])

    # Normalizar claves en dataframes financieros
    for df in [honorarios, honorarios_etapas, pagos_honorarios, cuota_litis, pagos_litis]:
        if df is not None and not df.empty and "Caso" in df.columns:
            df["Caso"] = df["Caso"].apply(normalize_key)

    # Preparar cuota litis con cálculo por fila
    cl = cuota_litis.copy()
    if not cl.empty:
        cl["Monto Base"] = safe_float_series(cl.get("Monto Base", 0))
        cl["Porcentaje"] = safe_float_series(cl.get("Porcentaje", 0))
        cl["CuotaCalc"] = cl["Monto Base"] * cl["Porcentaje"] / 100.0
    else:
        cl = pd.DataFrame(columns=["Caso", "CuotaCalc"])

    # Gastos informativos por actuaciones
    gastos_df = gastos_actuaciones_por_caso()
    gastos_map = {}
    if not gastos_df.empty:
        gastos_map = dict(
            zip(
                gastos_df["Caso"].astype(str),
                safe_float_series(gastos_df["GastosActuaciones"]).tolist()
            )
        )

    rows = []
    for _, c in casos.iterrows():
        exp = normalize_key(c.get("Expediente",""))

        # 1) Honorarios pactados
        etapas_exp = honorarios_etapas[honorarios_etapas["Caso"] == exp].copy()
        if not etapas_exp.empty:
            pactado = safe_float_series(etapas_exp.get("Monto Pactado", 0)).sum()
        else:
            sub_h = honorarios[honorarios["Caso"] == exp].copy()
            pactado = safe_float_series(sub_h.get("Monto Pactado", 0)).sum()

        # 2) Pagos honorarios
        sub_ph = pagos_honorarios[pagos_honorarios["Caso"] == exp].copy()
        pagado_h = safe_float_series(sub_ph.get("Monto", 0)).sum()

        # 3) Cuota litis calculada
        sub_cl = cl[cl["Caso"] == exp].copy()
        calc = safe_float_series(sub_cl.get("CuotaCalc", 0)).sum()

        # 4) Pagos litis
        sub_pl = pagos_litis[pagos_litis["Caso"] == exp].copy()
        pagado_l = safe_float_series(sub_pl.get("Monto", 0)).sum()

        # 5) Pendientes reales
        pend_h = max(0.0, float(pactado) - float(pagado_h))
        pend_l = max(0.0, float(calc) - float(pagado_l))

        # 6) Gastos actuaciones (INFORMATIVO)
        gastos_act = float(gastos_map.get(exp, 0.0))

        # ✅ 7) Saldo total (SIN gastos)
        saldo_total = float(pend_h + pend_l)

        rows.append([
            exp, c.get("Cliente",""), c.get("Materia",""),
            float(pactado), float(pagado_h), float(pend_h),
            float(calc), float(pagado_l), float(pend_l),
            float(gastos_act), float(saldo_total)
        ])

    return pd.DataFrame(rows, columns=[
        "Expediente","Cliente","Materia",
        "Honorario Pactado","Honorario Pagado","Honorario Pendiente",
        "Cuota Litis Calculada","Pagado Litis","Saldo Litis",
        "Gastos Actuaciones","Saldo Total"
    ])


def cuotas_status_all():
    if cuotas.empty:
        return pd.DataFrame()

    df = cuotas.copy()
    df["Caso"] = df["Caso"].apply(normalize_key)
    df["Monto"] = safe_float_series(df["Monto"])
    df["FechaVenc_dt"] = df["FechaVenc"].apply(to_date_safe)

    ph = pagos_honorarios.copy()
    pl = pagos_litis.copy()
    ph["Caso"] = ph["Caso"].apply(normalize_key)
    pl["Caso"] = pl["Caso"].apply(normalize_key)
    ph["Monto"] = safe_float_series(ph["Monto"])
    pl["Monto"] = safe_float_series(pl["Monto"])

    def calc_for_type(tipo, pagos_df):
        sub = df[df["Tipo"] == tipo].copy()
        if sub.empty:
            return sub

        sub["_sort_date"] = sub["FechaVenc_dt"].apply(lambda d: d if d else date(2100,1,1))
        sub["NroCuota"] = pd.to_numeric(sub["NroCuota"], errors="coerce").fillna(0).astype(int)
        sub.sort_values(["Caso","_sort_date","NroCuota"], inplace=True)

        pagado_por_caso = pagos_df.groupby("Caso")["Monto"].sum().to_dict()
        remaining = {k: float(v) for k, v in pagado_por_caso.items()}

        asignados, saldos, estados, dias = [], [], [], []
        today = date.today()
        for _, r in sub.iterrows():
            caso = r["Caso"]
            monto = float(r["Monto"])
            venc = r["FechaVenc_dt"]

            rem = remaining.get(caso, 0.0)
            asign = min(rem, monto) if monto > 0 else 0.0
            remaining[caso] = rem - asign
            saldo = monto - asign

            if monto == 0:
                est = "Sin monto"
            elif saldo <= 0.00001:
                est = "Pagada"
            elif asign > 0:
                est = "Parcial"
            else:
                est = "Pendiente"

            dv = None if venc is None else (venc - today).days
            asignados.append(asign); saldos.append(saldo); estados.append(est); dias.append(dv)

        sub["PagadoAsignado"] = asignados
        sub["SaldoCuota"] = saldos
        sub["Estado"] = estados
        sub["DiasParaVencimiento"] = dias
        sub.drop(columns=["_sort_date"], inplace=True, errors="ignore")
        return sub

    out_h = calc_for_type("Honorarios", ph)
    out_l = calc_for_type("CuotaLitis", pl)
    out = pd.concat([out_h, out_l], ignore_index=True) if (not out_h.empty or not out_l.empty) else pd.DataFrame()
    return out
# ==========================================================
# PANEL DE CONTROL (RESET OCULTO)
# ==========================================================
def reset_suave():
    casos_set = set(casos["Expediente"].tolist())

    def clean(keyname):
        df = load_df(keyname)
        if "Caso" in df.columns:
            df["Caso"] = df["Caso"].apply(normalize_key)
            df = df[df["Caso"].isin(casos_set)].copy()
        df = ensure_ids(df)
        save_df(keyname, df)

    for keyname in ["honorarios","honorarios_etapas","pagos_honorarios","cuota_litis","pagos_litis","cuotas","actuaciones","documentos","consultas"]:
        clean(keyname)

def reset_total(borrar_archivos=False):
    for k in FILES:
        backup_file(FILES[k])

    for k in FILES:
        if k == "usuarios":
            continue
        pd.DataFrame(columns=SCHEMAS[k]).to_csv(FILES[k], index=False)

    users = pd.DataFrame(columns=SCHEMAS["usuarios"])
    users = pd.concat([users, pd.DataFrame([{
        "Usuario":"admin",
        "PasswordHash": sha256(ADMIN_BOOTSTRAP_PASSWORD),
        "Rol":"admin",
        "AbogadoID":"",
        "Activo":"1",
        "Creado": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }])], ignore_index=True)
    users.to_csv(FILES["usuarios"], index=False)

    if borrar_archivos:
        for folder in [UPLOADS_DIR, GENERADOS_DIR]:
            try:
                if os.path.exists(folder):
                    shutil.rmtree(folder)
                os.makedirs(folder, exist_ok=True)
            except Exception:
                pass

# ==========================================================
# SIDEBAR
# ==========================================================
st.sidebar.markdown(f"### 🏷️ {APP_VERSION}")
st.sidebar.write(f"👤 Usuario: {st.session_state.usuario} ({st.session_state.rol})")

# Cerrar sesión
if st.sidebar.button("Cerrar sesión"):
    st.session_state.usuario = None
    st.session_state.rol = None
    st.session_state.abogado_id = ""
    st.rerun()

# Panel de control protegido (se mantiene igual)
with st.sidebar.expander("🔒 Panel de control", expanded=False):
    pwd = st.text_input("Clave del panel", type="password")
    if pwd == CONTROL_PASSWORD:
        st.success("Acceso concedido")
        if st.button("✅ Reset suave (limpiar fantasmas)"):
            reset_suave()
            st.success("✅ Reset suave aplicado")
            st.rerun()

        wipe = st.checkbox("Borrar también uploads/ y generados/ (solo reset total)", value=False)
        if st.button("🧨 Reset total (borra todo)"):
            reset_total(borrar_archivos=wipe)
            st.success("✅ Reset total aplicado. Usuario: admin. Contraseña: la configurada en Secrets.")
            st.rerun()
    else:
        st.info("Panel protegido. (Pide la clave)")

# ==========================================================
# MENÚ MARCA 004 (FILTRADO POR ROL – SIN REDUCIR)
# ==========================================================
rol = str(st.session_state.get("rol","")).strip().lower()

menu_items = [
    "Dashboard",
    "Ficha del Caso",
    "Clientes",
    "Abogados",
    "Casos",
    "Honorarios",
    "Pagos Honorarios",
    "Cuota Litis",
    "Pagos Cuota Litis",
    "Cronograma de Cuotas",
    "Actuaciones",
    "Consultas",
    "Documentos",
    "Plantillas de Contrato",
    "Generar Contrato",
    "Repositorio Contratos",
    "Instancias",
    "Usuarios",
    "Reportes",
    "Auditoría",
]

# -------- FILTRO POR ROL --------
if rol == "abogado":
    menu_items = [
        m for m in menu_items
        if m not in [
            "Honorarios",
            "Pagos Honorarios",
            "Cuota Litis",
            "Pagos Cuota Litis",
            "Cronograma de Cuotas",
            "Usuarios",
            "Auditoría",
        ]
    ]

elif rol in ["secretaria", "secretaria/o", "asistente"]:
    menu_items = [
        m for m in menu_items
        if m not in [
            "Honorarios",
            "Pagos Honorarios",
            "Cuota Litis",
            "Pagos Cuota Litis",
            "Cronograma de Cuotas",
            "Usuarios",
            "Auditoría",
            "Abogados",
        ]
    ]

elif rol == "solo lectura":
    menu_items = [
        m for m in menu_items
        if m in ["Casos", "Documentos"]
    ]

# Admin / Personal Administrativo ve todo (sin cambios)

menu = st.sidebar.radio("📌 Menú", menu_items)

# ==========================================================
# UI HEADER
# ==========================================================
brand_header()

# ==========================================================
# DASHBOARD COMPLETO (ROBUSTO: suma honorarios y litis, pendiente correcto)
# ==========================================================
if menu == "Dashboard":

    # =========================
    # Helpers locales (seguros)
    # =========================
    def _to_num(s):
        return pd.to_numeric(s, errors="coerce").fillna(0.0)

    def _norm(s):
        return normalize_key(s)

    # =========================
    # Construir resumen por caso (sin depender del "último registro")
    # =========================
    rows = []
    if casos.empty:
        df_res = pd.DataFrame(columns=[
            "Expediente","Cliente","Materia",
            "Honorario Pactado","Honorario Pagado","Honorario Pendiente",
            "Cuota Litis Calculada","Pagado Litis","Saldo Litis"
        ])
    else:
        # Copias defensivas
        h = honorarios.copy() if 'honorarios' in globals() else load_df("honorarios")
        he = honorarios_etapas.copy() if 'honorarios_etapas' in globals() else load_df("honorarios_etapas")
        ph = pagos_honorarios.copy() if 'pagos_honorarios' in globals() else load_df("pagos_honorarios")

        cl = cuota_litis.copy() if 'cuota_litis' in globals() else load_df("cuota_litis")
        pl = pagos_litis.copy() if 'pagos_litis' in globals() else load_df("pagos_litis")

        # Normalizar claves
        for df in [h, he, ph, cl, pl]:
            if df is not None and not df.empty and "Caso" in df.columns:
                df["Caso"] = df["Caso"].apply(_norm)

        # Normalizar montos
        if h is not None and not h.empty and "Monto Pactado" in h.columns:
            h["Monto Pactado"] = _to_num(h["Monto Pactado"])
        if he is not None and not he.empty and "Monto Pactado" in he.columns:
            he["Monto Pactado"] = _to_num(he["Monto Pactado"])
        if ph is not None and not ph.empty and "Monto" in ph.columns:
            ph["Monto"] = _to_num(ph["Monto"])
        if pl is not None and not pl.empty and "Monto" in pl.columns:
            pl["Monto"] = _to_num(pl["Monto"])

        # Cuota litis: calcular CuotaCalc por fila y sumar (todas)
        if cl is not None and not cl.empty:
            mb = "Monto Base" if "Monto Base" in cl.columns else None
            pc = "Porcentaje" if "Porcentaje" in cl.columns else None
            if mb is not None:
                cl[mb] = _to_num(cl[mb])
            else:
                cl["Monto Base"] = 0.0
                mb = "Monto Base"
            if pc is not None:
                cl[pc] = _to_num(cl[pc])
            else:
                cl["Porcentaje"] = 0.0
                pc = "Porcentaje"
            cl["CuotaCalc"] = cl[mb] * cl[pc] / 100.0
        else:
            cl = pd.DataFrame(columns=["Caso","CuotaCalc"])

        # Armar resumen por expediente
        for _, c in casos.iterrows():
            exp = _norm(c.get("Expediente",""))

            # 1) Honorario pactado:
            #    Si existen honorarios por etapa, sumarlos; si no, sumar todos los honorarios totales.
            pactado = 0.0
            if he is not None and not he.empty:
                sub_et = he[he["Caso"] == exp]
                if not sub_et.empty:
                    pactado = float(sub_et["Monto Pactado"].sum())

            if pactado == 0.0 and h is not None and not h.empty:
                sub_h = h[h["Caso"] == exp]
                pactado = float(sub_h["Monto Pactado"].sum()) if not sub_h.empty else 0.0

            # 2) Pagos honorarios (sumar todos)
            pagado_h = 0.0
            if ph is not None and not ph.empty:
                sub_ph = ph[ph["Caso"] == exp]
                pagado_h = float(sub_ph["Monto"].sum()) if not sub_ph.empty else 0.0

            # 3) Cuota litis calculada (sumar todas)
            calc_litis = 0.0
            if cl is not None and not cl.empty and "CuotaCalc" in cl.columns:
                sub_cl = cl[cl["Caso"] == exp]
                calc_litis = float(sub_cl["CuotaCalc"].sum()) if not sub_cl.empty else 0.0

            # 4) Pagos litis (sumar todos)
            pagado_l = 0.0
            if pl is not None and not pl.empty:
                sub_pl = pl[pl["Caso"] == exp]
                pagado_l = float(sub_pl["Monto"].sum()) if not sub_pl.empty else 0.0

            # 5) Pendientes (nunca negativos)
            pend_h = max(0.0, pactado - pagado_h)
            pend_l = max(0.0, calc_litis - pagado_l)

            rows.append([
                exp,
                c.get("Cliente",""),
                c.get("Materia",""),
                float(pactado),
                float(pagado_h),
                float(pend_h),
                float(calc_litis),
                float(pagado_l),
                float(pend_l)
            ])

        df_res = pd.DataFrame(rows, columns=[
            "Expediente","Cliente","Materia",
            "Honorario Pactado","Honorario Pagado","Honorario Pendiente",
            "Cuota Litis Calculada","Pagado Litis","Saldo Litis"
        ])

    # =========================
    # Estado de cuotas (vencidas / por vencer) - se mantiene igual que lo tenías
    # =========================
    df_estado = cuotas_status_all()

    # =========================
    # Totales para métricas
    # =========================
    total_pactado = df_res["Honorario Pactado"].sum() if not df_res.empty else 0
    total_pagado_h = df_res["Honorario Pagado"].sum() if not df_res.empty else 0
    total_pend_h = df_res["Honorario Pendiente"].sum() if not df_res.empty else 0

    total_litis = df_res["Cuota Litis Calculada"].sum() if not df_res.empty else 0
    total_pagado_l = df_res["Pagado Litis"].sum() if not df_res.empty else 0
    total_pend_l = df_res["Saldo Litis"].sum() if not df_res.empty else 0

    # =========================
    # UI
    # =========================
    st.subheader("📊 Dashboard General")

    r1c1, r1c2, r1c3 = st.columns(3)
    r1c1.metric("👥 Clientes", f"{len(clientes)}")
    r1c2.metric("👨‍⚖️ Abogados", f"{len(abogados)}")
    r1c3.metric("📁 Casos", f"{len(casos)}")

    st.markdown("### 💰 Indicadores Económicos")
    c1, c2, c3 = st.columns(3)
    c1.metric("Honorarios pactados (S/)", f"{total_pactado:,.2f}")
    c2.metric("Honorarios pagados (S/)", f"{total_pagado_h:,.2f}")
    c3.metric("Honorarios pendientes (S/)", f"{total_pend_h:,.2f}")

    c4, c5, c6 = st.columns(3)
    c4.metric("Cuota litis calculada (S/)", f"{total_litis:,.2f}")
    c5.metric("Cuota litis pagada (S/)", f"{total_pagado_l:,.2f}")
    c6.metric("Cuota litis pendiente (S/)", f"{total_pend_l:,.2f}")

    st.divider()
    st.markdown("### 📌 Detalle por caso")
    st.dataframe(df_res, use_container_width=True)

    st.divider()
    st.markdown("### 📅 Cuotas vencidas / por vencer")
    if df_estado is None or df_estado.empty or "SaldoCuota" not in df_estado.columns:
        st.info("Aún no hay cronograma calculable.")
    else:
        df_estado = df_estado.copy()
        df_estado["SaldoCuota"] = pd.to_numeric(df_estado["SaldoCuota"], errors="coerce").fillna(0.0)

        df_pend = df_estado[df_estado["SaldoCuota"] > 0].copy()
        vencidas = df_pend[df_pend["DiasParaVencimiento"].notna() & (df_pend["DiasParaVencimiento"] < 0)]
        por_vencer = df_pend[df_pend["DiasParaVencimiento"].notna() & (df_pend["DiasParaVencimiento"].between(0, 7))]

        st.markdown("**Vencidas**")
        st.dataframe(vencidas, use_container_width=True)
        st.markdown("**Por vencer (7 días)**")
        st.dataframe(por_vencer, use_container_width=True)

    st.download_button(
        "⬇️ Descargar reporte casos (CSV)",
        df_res.to_csv(index=False).encode("utf-8"),
        "reporte_casos.csv"
    )
# ==========================================================
# FICHA DEL CASO (CON DESCARGA WORD)
# ==========================================================
if menu == "Ficha del Caso":
    st.subheader("📁 Ficha del Caso")

    if casos.empty:
        st.info("Primero registra casos.")
    else:
        exp = st.selectbox("Expediente", casos["Expediente"].tolist())
        exp_n = normalize_key(exp)

        tabs = st.tabs([
            "Datos", "Pagos", "Cronograma",
            "Actuaciones", "Documentos", "Estado de Cuenta"
        ])

        # =========================
        # DATOS
        # =========================
        with tabs[0]:
            df_caso = casos[casos["Expediente"] == exp]
            st.dataframe(df_caso, use_container_width=True)

        # =========================
        # PAGOS
        # =========================
        with tabs[1]:
            st.markdown("### Pagos Honorarios")
            df_ph = pagos_honorarios[pagos_honorarios["Caso"] == exp_n]
            st.dataframe(df_ph, use_container_width=True)

            st.markdown("### Pagos Cuota Litis")
            df_pl = pagos_litis[pagos_litis["Caso"] == exp_n]
            st.dataframe(df_pl, use_container_width=True)

        # =========================
        # CRONOGRAMA
        # =========================
        with tabs[2]:
            st.markdown("### Cuotas registradas")
            st.dataframe(cuotas[cuotas["Caso"] == exp_n], use_container_width=True)

            st.markdown("### Estado cuotas")
            estado_cuotas = cuotas_status_all()
            if estado_cuotas is None or estado_cuotas.empty:
                st.info("No hay estado de cuotas disponible.")
            else:
                st.dataframe(
                    estado_cuotas[estado_cuotas["Caso"] == exp_n],
                    use_container_width=True
                )

        # =========================
        # ACTUACIONES
        # =========================
        with tabs[3]:
            df_act = actuaciones[actuaciones["Caso"] == exp_n].sort_values("Fecha", ascending=False)
            st.dataframe(df_act, use_container_width=True)

        # =========================
        # DOCUMENTOS
        # =========================
        with tabs[4]:
            df_doc = documentos[documentos["Caso"] == exp_n].sort_values("Fecha", ascending=False)
            st.dataframe(df_doc, use_container_width=True)

        # =========================
        # ESTADO DE CUENTA + WORD
        # =========================
        with tabs[5]:
            df_res = resumen_financiero_df()
            fila = df_res[df_res["Expediente"] == exp_n]

            if fila.empty:
                st.info("Sin estado de cuenta.")
                r = None
            else:
                r = fila.iloc[0]

                a, b, c = st.columns(3)
                a.metric("Honorario pactado", f"S/ {r['Honorario Pactado']:,.2f}")
                b.metric("Pagado honorarios", f"S/ {r['Honorario Pagado']:,.2f}")
                c.metric("Saldo honorarios", f"S/ {r['Honorario Pendiente']:,.2f}")

                d, e, g = st.columns(3)
                d.metric("Cuota litis calc.", f"S/ {r['Cuota Litis Calculada']:,.2f}")
                e.metric("Pagado litis", f"S/ {r['Pagado Litis']:,.2f}")
                g.metric("Saldo litis", f"S/ {r['Saldo Litis']:,.2f}")

            # =========================
            # GASTOS CLIENTE (INFORMATIVO)
            # =========================
            acts = actuaciones[actuaciones["Caso"] == exp_n].copy()
            acts["CostasAranceles"] = pd.to_numeric(acts.get("CostasAranceles",0), errors="coerce").fillna(0.0)
            acts["Gastos"] = pd.to_numeric(acts.get("Gastos",0), errors="coerce").fillna(0.0)
            acts["Total"] = acts["CostasAranceles"] + acts["Gastos"]
            acts["GastosPagado"] = acts.get("GastosPagado","0").astype(str)

            pend = acts.loc[acts["GastosPagado"]!="1","Total"].sum()
            pag = acts.loc[acts["GastosPagado"]=="1","Total"].sum()

            g1, g2 = st.columns(2)
            g1.metric("⏳ Gastos pendientes", f"S/ {pend:,.2f}")
            g2.metric("✅ Gastos pagados", f"S/ {pag:,.2f}")

            # =========================
            # EXPORTAR A WORD
            # =========================
            from io import BytesIO
            from docx import Document

            colw1, colw2 = st.columns(2)

            # ---- WORD: ESTADO DE CUENTA ----
            with colw1:
                if st.button("🧾 Descargar Estado de Cuenta (Word)"):
                    doc = Document()
                    doc.add_heading("ESTADO DE CUENTA DEL CASO", level=1)

                    doc.add_paragraph(f"Expediente: {exp}")
                    doc.add_paragraph(f"Cliente: {df_caso.iloc[0].get('Cliente','')}")
                    doc.add_paragraph("")

                    if r is not None:
                        doc.add_heading("Resumen Financiero", level=2)
                        doc.add_paragraph(f"Honorario pendiente: S/ {r['Honorario Pendiente']:,.2f}")
                        doc.add_paragraph(f"Saldo litis: S/ {r['Saldo Litis']:,.2f}")
                        doc.add_paragraph(f"Saldo total (sin gastos): S/ {r['Saldo Total']:,.2f}")

                    doc.add_heading("Gastos del Cliente", level=2)
                    doc.add_paragraph(f"Gastos pendientes: S/ {pend:,.2f}")
                    doc.add_paragraph(f"Gastos pagados: S/ {pag:,.2f}")

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "⬇️ Descargar Word",
                        buffer,
                        file_name=f"estado_cuenta_{exp.replace('/','_')}.docx"
                    )

            # ---- WORD: FICHA COMPLETA ----
            with colw2:
                if st.button("📁 Descargar Ficha Completa (Word)"):
                    doc = Document()
                    doc.add_heading("FICHA DEL CASO", level=1)

                    doc.add_heading("Datos del Caso", level=2)
                    for col in df_caso.columns:
                        doc.add_paragraph(f"{col}: {df_caso.iloc[0].get(col,'')}")

                    doc.add_heading("Resumen Financiero", level=2)
                    if r is not None:
                        for k in [
                            "Honorario Pactado","Honorario Pagado","Honorario Pendiente",
                            "Cuota Litis Calculada","Pagado Litis","Saldo Litis","Saldo Total"
                        ]:
                            doc.add_paragraph(f"{k}: S/ {r[k]:,.2f}")

                    doc.add_heading("Gastos del Cliente", level=2)
                    doc.add_paragraph(f"Gastos pendientes: S/ {pend:,.2f}")
                    doc.add_paragraph(f"Gastos pagados: S/ {pag:,.2f}")

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.download_button(
                        "⬇️ Descargar Ficha (Word)",
                        buffer,
                        file_name=f"ficha_caso_{exp.replace('/','_')}.docx"
                    )

# ==========================================================
# ==========================================================
# CLIENTES (CRUD) – Natural / Jurídica + Emergencia + Datos empresa
# ==========================================================
if menu == "Clientes":
 st.subheader("👥 Clientes")
 df_cli = load_df("clientes")
 is_readonly = st.session_state.get('rol') == 'asistente'
 accion = st.radio("Acción", ["Nuevo","Editar","Eliminar"], horizontal=True)
 if accion == "Nuevo":
  with st.form("nuevo_cliente"):
   tipo = st.selectbox("Tipo de cliente", ["Natural","Jurídica"], index=0)
   nombre = st.text_input("Nombre (Persona Natural)")
   dni = st.text_input("DNI")
   celular = st.text_input("Celular")
   correo = st.text_input("Correo")
   direccion = st.text_input("Dirección")
   contacto_em = st.text_input("Contacto de emergencia")
   cel_em = st.text_input("Celular de emergencia")
   st.markdown("**Datos de Persona Jurídica (si corresponde)**")
   razon = st.text_input("Razón Social")
   ruc = st.text_input("RUC")
   rep = st.text_input("Representante Legal")
   partida = st.text_input("Partida Electrónica")
   sede = st.text_input("Sede Registral")
   obs = st.text_area("Observaciones")
   submit = st.form_submit_button("Guardar", disabled=is_readonly)
  if submit:
   new_id = next_id(df_cli)
   row = {
    "ID": new_id,
    "TipoCliente": tipo,
    "Nombre": nombre,
    "DNI": dni,
    "Celular": celular,
    "Correo": correo,
    "Direccion": direccion,
    "Observaciones": obs,
    "ContactoEmergencia": contacto_em,
    "CelularEmergencia": cel_em,
    "RazonSocial": razon if tipo == "Jurídica" else "",
    "RUC": ruc if tipo == "Jurídica" else "",
    "RepresentanteLegal": rep if tipo == "Jurídica" else "",
    "PartidaElectronica": partida if tipo == "Jurídica" else "",
    "SedeRegistral": sede if tipo == "Jurídica" else "",
   }
   df_cli = add_row(df_cli, row, "clientes")
   save_df("clientes", df_cli)
   try: _audit_log('ADD','clientes', new_id, nombre)
   except Exception: pass
   st.success("✅ Cliente registrado")
   st.rerun()
 elif accion == "Editar" and not df_cli.empty:
  sel = st.selectbox("Cliente ID", df_cli["ID"].tolist(), key='cli_edit_id')
  fila = df_cli[df_cli["ID"] == sel].iloc[0]
  with st.form("edit_cliente"):
   tipo = st.selectbox("Tipo de cliente", ["Natural","Jurídica"], index=0 if str(fila.get('TipoCliente','Natural'))!='Jurídica' else 1)
   nombre = st.text_input("Nombre (Persona Natural)", value=str(fila.get("Nombre","")))
   dni = st.text_input("DNI", value=str(fila.get("DNI","")))
   celular = st.text_input("Celular", value=str(fila.get("Celular","")))
   correo = st.text_input("Correo", value=str(fila.get("Correo","")))
   direccion = st.text_input("Dirección", value=str(fila.get("Direccion","")))
   contacto_em = st.text_input("Contacto de emergencia", value=str(fila.get("ContactoEmergencia","")))
   cel_em = st.text_input("Celular de emergencia", value=str(fila.get("CelularEmergencia","")))
   st.markdown("**Datos de Persona Jurídica (si corresponde)**")
   razon = st.text_input("Razón Social", value=str(fila.get("RazonSocial","")))
   ruc = st.text_input("RUC", value=str(fila.get("RUC","")))
   rep = st.text_input("Representante Legal", value=str(fila.get("RepresentanteLegal","")))
   partida = st.text_input("Partida Electrónica", value=str(fila.get("PartidaElectronica","")))
   sede = st.text_input("Sede Registral", value=str(fila.get("SedeRegistral","")))
   obs = st.text_area("Observaciones", value=str(fila.get("Observaciones","")))
   submit = st.form_submit_button("Guardar cambios", disabled=is_readonly)
  if submit:
   idx = df_cli.index[df_cli["ID"] == sel][0]
   df_cli.loc[idx, [
    "TipoCliente","Nombre","DNI","Celular","Correo","Direccion","Observaciones",
    "ContactoEmergencia","CelularEmergencia",
    "RazonSocial","RUC","RepresentanteLegal","PartidaElectronica","SedeRegistral"
   ]] = [
    tipo, nombre, dni, celular, correo, direccion, obs,
    contacto_em, cel_em,
    (razon if tipo=="Jurídica" else ""),
    (ruc if tipo=="Jurídica" else ""),
    (rep if tipo=="Jurídica" else ""),
    (partida if tipo=="Jurídica" else ""),
    (sede if tipo=="Jurídica" else "")
   ]
   save_df("clientes", df_cli)
   try: _audit_log('UPDATE','clientes', sel, nombre)
   except Exception: pass
   st.success("✅ Actualizado")
   st.rerun()
 elif accion == "Eliminar" and not df_cli.empty:
  sel = st.selectbox("Cliente ID a eliminar", df_cli["ID"].tolist(), key='cli_del_id')
  if st.button("Eliminar cliente", disabled=is_readonly, key='cli_del_btn'):
   df_cli = df_cli[df_cli["ID"] != sel].copy()
   save_df("clientes", df_cli)
   try: _audit_log('DELETE','clientes', sel, '')
   except Exception: pass
   st.success("✅ Eliminado")
   st.rerun()
 st.dataframe(df_cli, use_container_width=True)
 st.download_button("⬇️ Descargar clientes (CSV)", df_cli.to_csv(index=False).encode("utf-8"), "clientes.csv")

# ==========================================================
# ABOGADOS (CRUD) – con Colegio Profesional, Distrito Judicial, Referencia y Notas
# ==========================================================
if menu == "Abogados":
 st.subheader("👨‍⚖️ Abogados")
 df_ab = load_df("abogados")
 is_readonly = st.session_state.get('rol') == 'asistente'
 accion = st.radio("Acción", ["Nuevo","Editar","Eliminar"], horizontal=True)
 if accion == "Nuevo":
  with st.form("nuevo_abogado"):
   nombre = st.text_input("Nombre")
   dni = st.text_input("DNI")
   celular = st.text_input("Celular")
   correo = st.text_input("Correo")
   coleg = st.text_input("Colegiatura")
   colegio_prof = st.text_input("Colegio Profesional")
   dom = st.text_input("Domicilio Procesal")
   referencia = st.text_input("Referencia domicilio procesal")
   cas_e = st.text_input("Casilla Electrónica")
   distrito = st.text_input("Distrito Judicial")
   cas_j = st.text_input("Casilla Judicial")
   notas = st.text_area("Notas", height=120)
   submit = st.form_submit_button("Guardar", disabled=is_readonly)
  if submit:
   new_id = next_id(df_ab)
   df_ab = add_row(df_ab, {
    "ID": new_id,
    "Nombre": nombre,
    "DNI": dni,
    "Celular": celular,
    "Correo": correo,
    "Colegiatura": coleg,
    "ColegioProfesional": colegio_prof,
    "Domicilio Procesal": dom,
    "ReferenciaDomicilio": referencia,
    "Casilla Electronica": cas_e,
    "DistritoJudicial": distrito,
    "Casilla Judicial": cas_j,
    "Notas": notas,
   }, "abogados")
   save_df("abogados", df_ab)
   try: _audit_log('ADD','abogados', new_id, nombre)
   except Exception: pass
   st.success("✅ Abogado registrado")
   st.rerun()
 elif accion == "Editar" and not df_ab.empty:
  sel = st.selectbox("Abogado ID", df_ab["ID"].tolist(), key='ab_edit_id')
  fila = df_ab[df_ab["ID"] == sel].iloc[0]
  with st.form("edit_abogado"):
   nombre = st.text_input("Nombre", value=str(fila.get("Nombre","")))
   dni = st.text_input("DNI", value=str(fila.get("DNI","")))
   celular = st.text_input("Celular", value=str(fila.get("Celular","")))
   correo = st.text_input("Correo", value=str(fila.get("Correo","")))
   coleg = st.text_input("Colegiatura", value=str(fila.get("Colegiatura","")))
   colegio_prof = st.text_input("Colegio Profesional", value=str(fila.get("ColegioProfesional","")))
   dom = st.text_input("Domicilio Procesal", value=str(fila.get("Domicilio Procesal","")))
   referencia = st.text_input("Referencia domicilio procesal", value=str(fila.get("ReferenciaDomicilio","")))
   cas_e = st.text_input("Casilla Electrónica", value=str(fila.get("Casilla Electronica","")))
   distrito = st.text_input("Distrito Judicial", value=str(fila.get("DistritoJudicial","")))
   cas_j = st.text_input("Casilla Judicial", value=str(fila.get("Casilla Judicial","")))
   notas = st.text_area("Notas", value=str(fila.get("Notas","")), height=120)
   submit = st.form_submit_button("Guardar cambios", disabled=is_readonly)
  if submit:
   idx = df_ab.index[df_ab["ID"] == sel][0]
   df_ab.loc[idx, [
    "Nombre","DNI","Celular","Correo","Colegiatura",
    "ColegioProfesional","Domicilio Procesal","ReferenciaDomicilio",
    "Casilla Electronica","DistritoJudicial","Casilla Judicial","Notas"
   ]] = [
    nombre,dni,celular,correo,coleg,
    colegio_prof,dom,referencia,
    cas_e,distrito,cas_j,notas
   ]
   save_df("abogados", df_ab)
   try: _audit_log('UPDATE','abogados', sel, nombre)
   except Exception: pass
   st.success("✅ Actualizado")
   st.rerun()
 elif accion == "Eliminar" and not df_ab.empty:
  sel = st.selectbox("Abogado ID a eliminar", df_ab["ID"].tolist(), key='ab_del_id')
  if st.button("Eliminar abogado", disabled=is_readonly, key='ab_del_btn'):
   df_ab = df_ab[df_ab["ID"] != sel].copy()
   save_df("abogados", df_ab)
   try: _audit_log('DELETE','abogados', sel, '')
   except Exception: pass
   st.success("✅ Eliminado")
   st.rerun()
 st.dataframe(df_ab, use_container_width=True)
 st.download_button("⬇️ Descargar abogados (CSV)", df_ab.to_csv(index=False).encode("utf-8"), "abogados.csv")

# ==========================================================
# CASOS (CRUD) — datos judiciales + DEFENSA CONJUNTA + DELEGACIÓN
# + Instancia incluye: "Por iniciar" e "Instancia Administrativa"
# ==========================================================
if menu == "Casos":
    st.subheader("📁 Casos")

    df_casos = load_df("casos")
    is_readonly = st.session_state.get('rol') == 'asistente'
    accion = st.radio("Acción", ["Nuevo","Editar","Eliminar"], horizontal=True, key="cas_accion")

    # ✅ Asegurar campos nuevos en el esquema (para que se guarden en CSV)
    try:
        if "casos" in SCHEMAS:
            for extra in ["DefensaConjunta", "NumAbogados", "AbogadosExtra",
                          "DelegacionActiva", "NumDelegados", "Delegados"]:
                if extra not in SCHEMAS["casos"]:
                    SCHEMAS["casos"].append(extra)
    except Exception:
        pass

    # -------------------------
    # Cargar clientes/abogados/usuarios UNA sola vez + tolerar columnas antiguas
    # -------------------------
    df_clientes = load_df("clientes")
    df_abogados = load_df("abogados")
    df_usuarios = load_df("usuarios")

    # Clientes: asegurar columna Nombre
    if df_clientes is None:
        df_clientes = pd.DataFrame()
    if not df_clientes.empty:
        if "Nombre" not in df_clientes.columns:
            for alt in ["Nombre completo", "NombreCompleto", "NOMBRE", "RazonSocial", "Razón Social", "Razon Social"]:
                if alt in df_clientes.columns:
                    df_clientes["Nombre"] = df_clientes[alt]
                    break
        if "Nombre" not in df_clientes.columns:
            df_clientes["Nombre"] = ""

    clientes_list = (
        df_clientes["Nombre"].dropna().astype(str).str.strip().tolist()
        if (df_clientes is not None and not df_clientes.empty and "Nombre" in df_clientes.columns)
        else []
    )
    clientes_list = [c for c in clientes_list if c != ""]

    # Abogados: asegurar columna Nombre
    if df_abogados is None:
        df_abogados = pd.DataFrame()
    if not df_abogados.empty:
        if "Nombre" not in df_abogados.columns:
            for alt in ["NOMBRE", "Nombre completo", "NombreCompleto"]:
                if alt in df_abogados.columns:
                    df_abogados["Nombre"] = df_abogados[alt]
                    break
        if "Nombre" not in df_abogados.columns:
            df_abogados["Nombre"] = ""

    abogados_list = (
        df_abogados["Nombre"].dropna().astype(str).str.strip().tolist()
        if (df_abogados is not None and not df_abogados.empty and "Nombre" in df_abogados.columns)
        else []
    )
    abogados_list = [a for a in abogados_list if a != ""]

    # Usuarios delegables (usuarios activos si existe columna Activo)
    delegables = []
    if df_usuarios is None:
        df_usuarios = pd.DataFrame()
    if not df_usuarios.empty and "Usuario" in df_usuarios.columns:
        tmpu = df_usuarios.copy()
        if "Activo" in tmpu.columns:
            tmpu = tmpu[tmpu["Activo"].astype(str) == "1"].copy()
        delegables = tmpu["Usuario"].astype(str).str.strip().tolist()
        delegables = [u for u in delegables if u != ""]

    # ✅ OPCIONES PARA EL CAMPO "Instancia" DEL CASO
    # (incluye lo que pediste sin depender de cambios globales)
    INSTANCIA_OPTS = ["Por iniciar", "Instancia Administrativa"] + list(ETAPAS_HONORARIOS)

    # ----------------------------------------------------------
    # NUEVO
    # ----------------------------------------------------------
    if accion == "Nuevo":
        if not clientes_list:
            st.warning("Primero registra clientes. (No se encontró columna Nombre utilizable en clientes.csv)")
        elif not abogados_list:
            st.warning("Primero registra abogados.")
        else:
            # ✅ Defensa conjunta (dinámico fuera del form)
            st.markdown("### ⚖️ Defensa del caso")
            abogado_principal_preview = st.selectbox("Abogado (principal)", abogados_list, key="cas_new_abogado_preview")
            defensa_conjunta = st.checkbox("✅ DEFENSA CONJUNTA (permitir más abogados)", value=False, key="cas_new_defconj")

            num_abogados = 1
            if defensa_conjunta:
                num_abogados = st.number_input(
                    "¿Cuántos abogados en total participarán?",
                    min_value=2, max_value=6, value=2, step=1,
                    key="cas_new_numab"
                )

            st.divider()

            # ✅ Delegación (dinámico fuera del form)
            st.markdown("### 🧩 Delegación del caso (asistentes / secretaria)")
            delegacion_activa = st.checkbox("✅ Delegar este caso a otros usuarios", value=False, key="cas_new_del_act")

            num_delegados = 0
            delegados_sel = []
            if delegacion_activa:
                if not delegables:
                    st.warning("No hay usuarios activos disponibles para delegar.")
                else:
                    num_delegados = st.number_input(
                        "¿Cuántos delegados tendrá el caso?",
                        min_value=1, max_value=6, value=1, step=1,
                        key="cas_new_num_del"
                    )
                    for i in range(int(num_delegados)):
                        delegados_sel.append(
                            st.selectbox(f"Delegado #{i+1} (usuario)", delegables, key=f"cas_new_del_{i}")
                        )

            st.divider()

            with st.form("nuevo_caso"):
                cliente = st.selectbox("Cliente", clientes_list, key="cas_new_cliente")

                abogado = st.selectbox(
                    "Abogado (principal)",
                    abogados_list,
                    index=abogados_list.index(abogado_principal_preview) if abogado_principal_preview in abogados_list else 0,
                    key="cas_new_abogado"
                )

                abogados_extra = []
                if defensa_conjunta:
                    st.markdown("### 👥 Abogados adicionales")
                    opciones_extra = [a for a in abogados_list if a != abogado] or abogados_list[:]
                    for i in range(int(num_abogados) - 1):
                        abogados_extra.append(
                            st.selectbox(f"Abogado adicional #{i+1}", opciones_extra, key=f"cas_new_ab_extra_{i}")
                        )

                expediente = st.text_input("Expediente", key="cas_new_exp")
                anio = st.text_input("Año", key="cas_new_anio")
                materia = st.text_input("Materia", key="cas_new_mat")

                # ✅ AQUÍ están las opciones nuevas de Instancia
                instancia = st.selectbox("Instancia", INSTANCIA_OPTS, key="cas_new_inst")

                pret = st.text_input("Pretensión", key="cas_new_pret")
                juzgado = st.text_input("Juzgado", key="cas_new_juz")
                distrito_jud = st.text_input("Distrito Judicial", key="cas_new_dist")
                contraparte = st.text_input("Contraparte", key="cas_new_contra")
                contraparte_doc = st.text_input("DNI/RUC Contraparte", key="cas_new_contradoc")
                obs = st.text_area("Observaciones", key="cas_new_obs")
                estado = st.selectbox("EstadoCaso", ["Activo","En pausa","Cerrado","Archivado"], key="cas_new_estado")
                fi = st.date_input("Fecha inicio", value=date.today(), key="cas_new_fi")

                submit = st.form_submit_button("Guardar", disabled=is_readonly)

            if submit:
                new_id = next_id(df_casos)
                df_casos = add_row(df_casos, {
                    "ID": new_id,
                    "Cliente": cliente,
                    "Abogado": abogado,
                    "Expediente": normalize_key(expediente),
                    "Año": anio,
                    "Materia": materia,
                    "Instancia": instancia,
                    "Pretension": pret,
                    "Juzgado": juzgado,
                    "DistritoJudicial": distrito_jud,
                    "Contraparte": contraparte,
                    "ContraparteDoc": contraparte_doc,
                    "Observaciones": obs,
                    "EstadoCaso": estado,
                    "FechaInicio": str(fi),

                    "DefensaConjunta": "1" if defensa_conjunta else "0",
                    "NumAbogados": int(num_abogados) if defensa_conjunta else 1,
                    "AbogadosExtra": " | ".join([a for a in abogados_extra if str(a).strip() != ""]),

                    "DelegacionActiva": "1" if delegacion_activa else "0",
                    "NumDelegados": int(num_delegados) if delegacion_activa else 0,
                    "Delegados": " | ".join([d for d in delegados_sel if str(d).strip() != ""]),
                }, "casos")

                save_df("casos", df_casos)
                try:
                    _audit_log('ADD','casos', new_id, expediente)
                except Exception:
                    pass
                st.success("✅ Caso registrado")
                st.rerun()

    # ----------------------------------------------------------
    # EDITAR
    # ----------------------------------------------------------
    elif accion == "Editar" and df_casos is not None and not df_casos.empty:
        exp_sel = st.selectbox("Expediente", df_casos["Expediente"].tolist(), key='cas_edit_exp')
        fila = df_casos[df_casos["Expediente"] == exp_sel].iloc[0]

        del_act_prev = str(fila.get("DelegacionActiva","0")) == "1"
        num_del_prev = int(pd.to_numeric(fila.get("NumDelegados", 0), errors="coerce") or 0)
        del_prev = str(fila.get("Delegados","") or "")
        del_list_prev = [x.strip() for x in del_prev.split("|") if x.strip()]

        with st.form("edit_caso"):
            cliente = st.selectbox(
                "Cliente",
                clientes_list,
                index=clientes_list.index(fila.get('Cliente','')) if fila.get('Cliente','') in clientes_list else 0,
                key="cas_edit_cliente"
            )
            abogado = st.selectbox(
                "Abogado (principal)",
                abogados_list,
                index=abogados_list.index(fila.get('Abogado','')) if fila.get('Abogado','') in abogados_list else 0,
                key="cas_edit_abogado"
            )

            # Defensa conjunta
            defconj_val = str(fila.get("DefensaConjunta","0")) == "1"
            defensa_conjunta = st.checkbox("✅ DEFENSA CONJUNTA (más abogados)", value=defconj_val, key="cas_edit_defconj")

            num_prev = int(pd.to_numeric(fila.get("NumAbogados", 1), errors="coerce") or 1)
            extras_prev = str(fila.get("AbogadosExtra","") or "")
            extras_list_prev = [x.strip() for x in extras_prev.split("|") if x.strip()]

            num_abogados = 1
            abogados_extra = []

            if defensa_conjunta:
                num_abogados = st.number_input(
                    "¿Cuántos abogados en total participarán?",
                    min_value=2, max_value=6,
                    value=max(2, num_prev),
                    step=1,
                    key="cas_edit_numab"
                )
                st.markdown("### 👥 Abogados adicionales")
                opciones_extra = [a for a in abogados_list if a != abogado] or abogados_list[:]
                for i in range(int(num_abogados) - 1):
                    default = extras_list_prev[i] if i < len(extras_list_prev) else (opciones_extra[0] if opciones_extra else "")
                    idx_def = opciones_extra.index(default) if default in opciones_extra else 0
                    abogados_extra.append(
                        st.selectbox(f"Abogado adicional #{i+1}", opciones_extra, index=idx_def, key=f"cas_edit_ab_extra_{i}")
                    )

            anio = st.text_input("Año", value=str(fila.get('Año','')), key="cas_edit_anio")
            materia = st.text_input("Materia", value=str(fila.get('Materia','')), key="cas_edit_mat")
            instancia = st.selectbox(
                "Instancia",
                INSTANCIA_OPTS,
                index=INSTANCIA_OPTS.index(fila.get('Instancia','')) if fila.get('Instancia','') in INSTANCIA_OPTS else 0,
                key="cas_edit_inst"
            )
            pret = st.text_input("Pretensión", value=str(fila.get('Pretension','')), key="cas_edit_pret")
            juzgado = st.text_input("Juzgado", value=str(fila.get('Juzgado','')), key="cas_edit_juz")
            distrito_jud = st.text_input("Distrito Judicial", value=str(fila.get('DistritoJudicial','')), key="cas_edit_dist")
            contraparte = st.text_input("Contraparte", value=str(fila.get('Contraparte','')), key="cas_edit_contra")
            contraparte_doc = st.text_input("DNI/RUC Contraparte", value=str(fila.get('ContraparteDoc','')), key="cas_edit_contradoc")
            obs = st.text_area("Observaciones", value=str(fila.get('Observaciones','')), key="cas_edit_obs")
            estado = st.selectbox(
                "EstadoCaso",
                ["Activo","En pausa","Cerrado","Archivado"],
                index=["Activo","En pausa","Cerrado","Archivado"].index(fila.get('EstadoCaso','Activo')) if fila.get('EstadoCaso','Activo') in ["Activo","En pausa","Cerrado","Archivado"] else 0,
                key="cas_edit_estado"
            )

            # Delegación
            st.markdown("### 🧩 Delegación del caso")
            delegacion_activa = st.checkbox("✅ Delegar este caso a otros usuarios", value=del_act_prev, key="cas_edit_del_act")
            num_delegados = 0
            delegados_sel = []

            if delegacion_activa:
                if not delegables:
                    st.warning("No hay usuarios activos disponibles para delegar.")
                else:
                    num_delegados = st.number_input(
                        "¿Cuántos delegados tendrá el caso?",
                        min_value=1, max_value=6,
                        value=max(1, num_del_prev) if num_del_prev else 1,
                        step=1,
                        key="cas_edit_num_del"
                    )
                    for i in range(int(num_delegados)):
                        default = del_list_prev[i] if i < len(del_list_prev) else delegables[0]
                        idx_def = delegables.index(default) if default in delegables else 0
                        delegados_sel.append(
                            st.selectbox(f"Delegado #{i+1} (usuario)", delegables, index=idx_def, key=f"cas_edit_del_{i}")
                        )

            submit = st.form_submit_button("Guardar cambios", disabled=is_readonly)

        if submit:
            idx = df_casos.index[df_casos["Expediente"] == exp_sel][0]
            df_casos.loc[idx, [
                "Cliente","Abogado","Año","Materia","Instancia","Pretension",
                "Juzgado","DistritoJudicial","Contraparte","ContraparteDoc","Observaciones","EstadoCaso",
                "DefensaConjunta","NumAbogados","AbogadosExtra",
                "DelegacionActiva","NumDelegados","Delegados"
            ]] = [
                cliente, abogado, anio, materia, instancia, pret,
                juzgado, distrito_jud, contraparte, contraparte_doc, obs, estado,
                "1" if defensa_conjunta else "0",
                int(num_abogados) if defensa_conjunta else 1,
                " | ".join([a for a in abogados_extra if str(a).strip() != ""]),
                "1" if delegacion_activa else "0",
                int(num_delegados) if delegacion_activa else 0,
                " | ".join([d for d in delegados_sel if str(d).strip() != ""]),
            ]

            save_df("casos", df_casos)
            try:
                _audit_log('UPDATE','casos', exp_sel, 'edit')
            except Exception:
                pass
            st.success("✅ Caso actualizado")
            st.rerun()

    # ----------------------------------------------------------
    # ELIMINAR
    # ----------------------------------------------------------
    elif accion == "Eliminar" and df_casos is not None and not df_casos.empty:
        exp_sel = st.selectbox("Expediente a eliminar", df_casos["Expediente"].tolist(), key='cas_del_exp')
        st.warning("⚠️ Esta acción no se puede deshacer")
        if st.button("Eliminar caso", disabled=is_readonly, key='cas_del_btn'):
            df_casos = df_casos[df_casos["Expediente"] != exp_sel].copy()
            save_df("casos", df_casos)
            try:
                _audit_log('DELETE','casos', exp_sel, '')
            except Exception:
                pass
            st.success("✅ Caso eliminado")
            st.rerun()

    st.dataframe(df_casos, use_container_width=True)
    st.download_button("⬇️ Descargar casos (CSV)", df_casos.to_csv(index=False).encode("utf-8"), "casos.csv")

# ==========================================================
# HONORARIOS (Total + Por etapa) ✅ DEFINITIVO
# - Se elimina definitivamente "honorarios por tipo (MARCA 006)"
# - Modo TOTAL vs ETAPAS
# - Edición / Borrado robusto
# ==========================================================
if menu == "Honorarios":
    st.subheader("🧾 Honorarios")

    is_readonly = st.session_state.get("rol") == "asistente"

    tab_total, tab_etapas = st.tabs(["Total (como siempre)", "Por etapa / instancia"])

    # =========================
    # Helpers locales (ROBUSTOS)
    # =========================
    def _num(x):
        return pd.to_numeric(pd.Series([x]), errors="coerce").fillna(0.0).iloc[0]

    def _norm(exp):
        return normalize_key(exp)

    def _sum_etapas(expediente: str) -> float:
        if honorarios_etapas is None or honorarios_etapas.empty:
            return 0.0
        exp = _norm(expediente)
        sub = honorarios_etapas[honorarios_etapas["Caso"].astype(str) == exp]
        if sub.empty:
            return 0.0
        return float(pd.to_numeric(sub["Monto Pactado"], errors="coerce").fillna(0.0).sum())

    def _sum_total(expediente: str) -> float:
        if honorarios is None or honorarios.empty:
            return 0.0
        exp = _norm(expediente)
        sub = honorarios[honorarios["Caso"].astype(str) == exp]
        if sub.empty:
            return 0.0
        return float(pd.to_numeric(sub["Monto Pactado"], errors="coerce").fillna(0.0).sum())

    def _sum_pagos(expediente: str) -> float:
        if pagos_honorarios is None or pagos_honorarios.empty:
            return 0.0
        exp = _norm(expediente)
        sub = pagos_honorarios[pagos_honorarios["Caso"].astype(str) == exp]
        if sub.empty:
            return 0.0
        return float(pd.to_numeric(sub["Monto"], errors="coerce").fillna(0.0).sum())

    exp_list = casos["Expediente"].tolist() if not casos.empty else []

    # ==========================================================
    # TAB TOTAL
    # ==========================================================
    with tab_total:
        st.markdown("### Honorario total")

        st.dataframe(honorarios.sort_values("ID", ascending=False), use_container_width=True)

        if exp_list:
            exp = st.selectbox("Expediente", exp_list, key="hon_total_exp")

            modo = st.radio(
                "Modo de registro",
                ["Monto total (registro directo)", "Por etapas / instancias (auto-suma)"],
                horizontal=True,
                key="hon_modo"
            )

            pactado_et = _sum_etapas(exp)
            pactado_total_reg = _sum_total(exp)
            pagado = _sum_pagos(exp)

            pactado_recomendado = pactado_et if pactado_et > 0 else pactado_total_reg
            pendiente = max(0.0, pactado_recomendado - pagado)

            m1, m2, m3 = st.columns(3)
            m1.metric("Pactado", f"S/ {pactado_recomendado:,.2f}")
            m2.metric("Pagado", f"S/ {pagado:,.2f}")
            m3.metric("Pendiente", f"S/ {pendiente:,.2f}")

            st.divider()

            if modo == "Monto total (registro directo)":
                monto = st.number_input(
                    "Monto pactado (total)",
                    min_value=0.0,
                    step=50.0,
                    key="hon_total_monto_new"
                )
                notas = st.text_input(
                    "Notas",
                    value="",
                    key="hon_total_notas_new"
                )

                if st.button(
                    "Guardar honorario total",
                    disabled=is_readonly,
                    key="hon_total_save_btn"
                ):
                    honorarios = add_row(honorarios, {
                        "ID": next_id(honorarios),
                        "Caso": _norm(exp),
                        "Monto Pactado": float(monto),
                        "Notas": notas,
                        "FechaRegistro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }, "honorarios")
                    save_df("honorarios", honorarios)
                    st.success("✅ Guardado")
                    st.rerun()

            else:
                st.info("Modo por etapas activo. El total se calcula automáticamente.")
                st.write(f"📌 Total por etapas: **S/ {pactado_et:,.2f}**")

        # ✅ Editar / borrar honorario total
        st.divider()
        st.markdown("### Editar / borrar honorario total")

        if not honorarios.empty:
            sel = st.selectbox(
                "Honorario ID",
                honorarios["ID"].tolist(),
                key="hon_total_sel_id"
            )
            fila = honorarios[honorarios["ID"] == sel].iloc[0]

            with st.form("hon_total_edit"):
                caso_e = st.text_input(
                    "Caso",
                    value=fila["Caso"],
                    key="hon_total_caso_edit"
                )
                monto_e = st.number_input(
                    "Monto Pactado",
                    min_value=0.0,
                    value=_num(fila["Monto Pactado"]),
                    step=50.0,
                    key="hon_total_monto_edit"
                )
                notas_e = st.text_input(
                    "Notas",
                    value=str(fila.get("Notas","")),
                    key="hon_total_notas_edit"
                )
                submit = st.form_submit_button(
                    "Guardar cambios",
                    disabled=is_readonly
                )

            if submit:
                idx = honorarios.index[honorarios["ID"] == sel][0]
                honorarios.loc[idx, ["Caso","Monto Pactado","Notas"]] = [_norm(caso_e), float(monto_e), notas_e]
                save_df("honorarios", honorarios)
                st.success("✅ Actualizado")
                st.rerun()

            st.warning("⚠️ Eliminar honorario total")
            confirm = st.text_input(
                "Escribe ELIMINAR para confirmar",
                key="hon_total_del_confirm"
            )
            if st.button(
                "🗑️ Borrar",
                disabled=is_readonly or confirm.strip().upper() != "ELIMINAR",
                key="hon_total_del_btn"
            ):
                honorarios = honorarios[honorarios["ID"] != sel].copy()
                save_df("honorarios", honorarios)
                st.success("✅ Eliminado")
                st.rerun()

        st.download_button(
            "⬇️ Descargar honorarios total (CSV)",
            honorarios.to_csv(index=False).encode("utf-8"),
            "honorarios.csv",
            key="hon_total_dl_csv"
        )

    # ==========================================================
    # TAB ETAPAS / INSTANCIAS
    # ==========================================================
    with tab_etapas:
        st.markdown("### Honorarios por etapa / instancia")

        st.dataframe(honorarios_etapas.sort_values("ID", ascending=False), use_container_width=True)

        if exp_list:
            exp = st.selectbox("Expediente", exp_list, key="hon_et_exp")
            etapa = st.selectbox(
                "Instancia / Etapa",
                ETAPAS_HONORARIOS,
                key="hon_et_etapa_new"
            )
            monto = st.number_input(
                "Monto pactado",
                min_value=0.0,
                step=50.0,
                key="hon_et_monto_new"
            )
            notas = st.text_input(
                "Notas",
                value="",
                key="hon_et_notas_new"
            )

            total_et = _sum_etapas(exp)
            st.caption(f"📌 Total por etapas: S/ {total_et:,.2f}")

            if st.button(
                "Guardar honorario por etapa",
                disabled=is_readonly,
                key="hon_et_save_btn"
            ):
                honorarios_etapas = add_row(honorarios_etapas, {
                    "ID": next_id(honorarios_etapas),
                    "Caso": _norm(exp),
                    "Etapa": etapa,
                    "Monto Pactado": float(monto),
                    "Notas": notas,
                    "FechaRegistro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }, "honorarios_etapas")
                save_df("honorarios_etapas", honorarios_etapas)
                st.success("✅ Guardado")
                st.rerun()

        st.divider()
        st.markdown("### Editar / borrar honorario por etapa")

        if not honorarios_etapas.empty:
            sel = st.selectbox(
                "Honorario Etapa ID",
                honorarios_etapas["ID"].tolist(),
                key="hon_et_sel_id"
            )
            fila = honorarios_etapas[honorarios_etapas["ID"] == sel].iloc[0]

            with st.form("hon_et_edit"):
                caso_e = st.text_input(
                    "Caso",
                    value=fila["Caso"],
                    key="hon_et_caso_edit"
                )
                etapa_e = st.selectbox(
                    "Instancia / Etapa",
                    ETAPAS_HONORARIOS,
                    index=ETAPAS_HONORARIOS.index(fila["Etapa"]) if fila["Etapa"] in ETAPAS_HONORARIOS else 0,
                    key="hon_et_etapa_edit"
                )
                monto_e = st.number_input(
                    "Monto Pactado",
                    min_value=0.0,
                    value=_num(fila["Monto Pactado"]),
                    step=50.0,
                    key="hon_et_monto_edit"
                )
                notas_e = st.text_input(
                    "Notas",
                    value=str(fila.get("Notas","")),
                    key="hon_et_notas_edit"
                )
                submit = st.form_submit_button(
                    "Guardar cambios",
                    disabled=is_readonly
                )

            if submit:
                idx = honorarios_etapas.index[honorarios_etapas["ID"] == sel][0]
                honorarios_etapas.loc[idx, ["Caso","Etapa","Monto Pactado","Notas"]] = [_norm(caso_e), etapa_e, float(monto_e), notas_e]
                save_df("honorarios_etapas", honorarios_etapas)
                st.success("✅ Actualizado")
                st.rerun()

            st.warning("⚠️ Eliminar honorario por etapa")
            confirm = st.text_input(
                "Escribe ELIMINAR para confirmar",
                key="hon_et_del_confirm"
            )
            if st.button(
                "🗑️ Borrar",
                disabled=is_readonly or confirm.strip().upper() != "ELIMINAR",
                key="hon_et_del_btn"
            ):
                honorarios_etapas = honorarios_etapas[honorarios_etapas["ID"] != sel].copy()
                save_df("honorarios_etapas", honorarios_etapas)
                st.success("✅ Eliminado")
                st.rerun()

        st.download_button(
            "⬇️ Descargar honorarios por etapa (CSV)",
            honorarios_etapas.to_csv(index=False).encode("utf-8"),
            "honorarios_etapas.csv",
            key="hon_et_dl_csv"
        )
# ==========================================================
# PAGOS HONORARIOS (por etapa + editar/borrar)
# ==========================================================
if menu == "Pagos Honorarios":
    st.subheader("💳 Pagos Honorarios (con etapa)")
    st.dataframe(pagos_honorarios.sort_values("FechaPago", ascending=False), use_container_width=True)

    exp_list = casos["Expediente"].tolist()
    if exp_list:
        exp = st.selectbox("Expediente", exp_list, key="ph_exp")
        etapa = st.selectbox("Etapa", ETAPAS_HONORARIOS, key="ph_etapa")
        fecha = st.date_input("Fecha pago", value=date.today(), key="ph_fecha")
        monto = st.number_input("Monto pagado", min_value=0.0, step=50.0, key="ph_monto")
        obs = st.text_input("Observación", value="", key="ph_obs")

        if st.button("Registrar pago honorarios"):
            new_id = next_id(pagos_honorarios)
            pagos_honorarios = add_row(pagos_honorarios, {
                "ID": new_id, "Caso": normalize_key(exp), "Etapa": etapa,
                "FechaPago": str(fecha), "Monto": float(monto), "Observacion": obs
            }, "pagos_honorarios")
            save_df("pagos_honorarios", pagos_honorarios)
            st.success("✅ Pago registrado")
            st.rerun()

    st.divider()
    st.markdown("### Editar / borrar pago (por ID)")
    if not pagos_honorarios.empty:
        sel = st.selectbox("Pago ID", pagos_honorarios["ID"].tolist(), key="ph_sel")
        fila = pagos_honorarios[pagos_honorarios["ID"] == sel].iloc[0]

        with st.form("ph_edit_form"):
            exp_e = st.selectbox("Expediente", exp_list, index=exp_list.index(fila["Caso"]) if fila["Caso"] in exp_list else 0)
            etapa_e = st.selectbox("Etapa", ETAPAS_HONORARIOS, index=ETAPAS_HONORARIOS.index(fila["Etapa"]) if fila["Etapa"] in ETAPAS_HONORARIOS else 0)
            fecha_e = st.text_input("FechaPago (YYYY-MM-DD)", value=str(fila["FechaPago"]))
            monto_e = st.number_input("Monto", min_value=0.0, value=money(fila["Monto"]), step=50.0)
            obs_e = st.text_input("Observación", value=str(fila["Observacion"]))
            submit = st.form_submit_button("Guardar cambios")
            if submit:
                idx = pagos_honorarios.index[pagos_honorarios["ID"] == sel][0]
                pagos_honorarios.loc[idx, :] = [sel, normalize_key(exp_e), etapa_e, fecha_e, float(monto_e), obs_e]
                save_df("pagos_honorarios", pagos_honorarios)
                st.success("✅ Actualizado")
                st.rerun()

        if st.button("🗑️ Borrar pago honorarios", key="ph_del"):
            pagos_honorarios = pagos_honorarios[pagos_honorarios["ID"] != sel].copy()
            save_df("pagos_honorarios", pagos_honorarios)
            st.success("✅ Eliminado")
            st.rerun()

    st.download_button("⬇️ Descargar pagos honorarios (CSV)", pagos_honorarios.to_csv(index=False).encode("utf-8"), "pagos_honorarios.csv")

# ==========================================================
# CUOTA LITIS
# ==========================================================
if menu == "Cuota Litis":
    st.subheader("⚖️ Cuota Litis")
    st.dataframe(cuota_litis.sort_values("ID", ascending=False), use_container_width=True)
    st.download_button("⬇️ Descargar cuota litis (CSV)", cuota_litis.to_csv(index=False).encode("utf-8"), "cuota_litis.csv")

# ==========================================================
# PAGOS CUOTA LITIS (editar/borrar)
# ==========================================================
if menu == "Pagos Cuota Litis":
    st.subheader("💳 Pagos Cuota Litis")
    st.dataframe(pagos_litis.sort_values("FechaPago", ascending=False), use_container_width=True)
    st.download_button("⬇️ Descargar pagos litis (CSV)", pagos_litis.to_csv(index=False).encode("utf-8"), "pagos_litis.csv")

# ==========================================================
# CRONOGRAMA DE CUOTAS (EDITAR / BORRAR / MARCAR PAGADA)
# - No cambia esquemas
# - "Marcar pagada" crea un pago real (honorarios o litis)
# ==========================================================
if menu == "Cronograma de Cuotas":
    st.subheader("📅 Cronograma de cuotas")

    is_readonly = st.session_state.get("rol") == "asistente"

    st.markdown("### Cuotas registradas")
    st.dataframe(cuotas.sort_values(["Caso","Tipo","NroCuota"], ascending=[True, True, True]), use_container_width=True)

    exp_list = casos["Expediente"].tolist()
    if not exp_list:
        st.warning("Primero registra casos para poder crear cuotas.")
    else:
        # ==========================================================
        # CREAR CUOTA
        # ==========================================================
        st.divider()
        st.markdown("### ➕ Crear cuota")

        caso_new = st.selectbox("Expediente", exp_list, key="cr_caso_new")
        tipo_new = st.selectbox("Tipo", TIPOS_CUOTA, key="cr_tipo_new")
        venc_new = st.date_input("Fecha vencimiento", value=date.today(), key="cr_venc_new")
        monto_new = st.number_input("Monto cuota", min_value=0.0, step=50.0, key="cr_monto_new")
        notas_new = st.text_input("Notas", value="", key="cr_notas_new")

        caso_norm = normalize_key(caso_new)
        tipo_norm = "Honorarios" if str(tipo_new) == "Honorarios" else "CuotaLitis"

        sub = cuotas[(cuotas["Caso"] == caso_norm) & (cuotas["Tipo"] == tipo_norm)].copy()
        sub["NroCuota"] = pd.to_numeric(sub["NroCuota"], errors="coerce").fillna(0).astype(int)
        nro = int(sub["NroCuota"].max()) + 1 if not sub.empty else 1

        if st.button("Guardar cuota", disabled=is_readonly, key="cr_save_new"):
            new_id = next_id(cuotas)
            cuotas = add_row(cuotas, {
                "ID": new_id,
                "Caso": caso_norm,
                "Tipo": tipo_norm,
                "NroCuota": nro,
                "FechaVenc": str(venc_new),
                "Monto": float(monto_new),
                "Notas": notas_new
            }, "cuotas")
            save_df("cuotas", cuotas)
            st.success("✅ Cuota creada")
            st.rerun()

        # ==========================================================
        # EDITAR / BORRAR CUOTA
        # ==========================================================
        st.divider()
        st.markdown("### ✏️ Editar / borrar cuota (por ID)")

        if cuotas.empty:
            st.info("No hay cuotas registradas.")
        else:
            sel_id = st.selectbox("Cuota ID", cuotas["ID"].tolist(), key="cr_sel_id")
            fila = cuotas[cuotas["ID"] == sel_id].iloc[0]

            with st.form("cr_edit_form"):
                caso_e = st.selectbox("Expediente", exp_list,
                                      index=exp_list.index(fila["Caso"]) if fila["Caso"] in exp_list else 0,
                                      key="cr_edit_caso")
                tipo_e = st.selectbox("Tipo", ["Honorarios","CuotaLitis"],
                                      index=0 if str(fila["Tipo"]) == "Honorarios" else 1,
                                      key="cr_edit_tipo")
                nro_e = st.number_input("NroCuota", min_value=1, step=1,
                                        value=int(pd.to_numeric(fila["NroCuota"], errors="coerce") or 1),
                                        key="cr_edit_nro")
                venc_e = st.text_input("FechaVenc (YYYY-MM-DD)", value=str(fila["FechaVenc"]), key="cr_edit_venc")
                monto_e = st.number_input("Monto", min_value=0.0, step=50.0,
                                          value=float(pd.to_numeric(fila["Monto"], errors="coerce") or 0),
                                          key="cr_edit_monto")
                notas_e = st.text_input("Notas", value=str(fila.get("Notas","")), key="cr_edit_notas")

                submit = st.form_submit_button("Guardar cambios", disabled=is_readonly)

            if submit:
                idx = cuotas.index[cuotas["ID"] == sel_id][0]
                cuotas.loc[idx, ["Caso","Tipo","NroCuota","FechaVenc","Monto","Notas"]] = [
                    normalize_key(caso_e),
                    "Honorarios" if str(tipo_e) == "Honorarios" else "CuotaLitis",
                    int(nro_e),
                    str(venc_e),
                    float(monto_e),
                    notas_e
                ]
                save_df("cuotas", cuotas)
                st.success("✅ Cuota actualizada")
                st.rerun()

            st.warning("⚠️ Eliminar cuota (irreversible)")
            confirm_del = st.text_input("Escribe ELIMINAR para confirmar", key="cr_del_confirm")
            if st.button("🗑️ Eliminar cuota", disabled=is_readonly or confirm_del.strip().upper() != "ELIMINAR", key="cr_del_btn"):
                cuotas = cuotas[cuotas["ID"] != sel_id].copy()
                save_df("cuotas", cuotas)
                st.success("✅ Cuota eliminada")
                st.rerun()

        # ==========================================================
        # MARCAR CUOTA COMO PAGADA (CREA PAGO REAL)
        # ==========================================================
        st.divider()
        st.markdown("### ✅ Marcar cuota como pagada (crea un pago real)")

        st.caption("Esto NO cambia un 'flag' en cuotas; crea un pago en Pagos Honorarios o Pagos Cuota Litis para que el sistema lo reconozca.")
        st.caption("Para revertir, elimina ese pago en el módulo correspondiente.")

        if cuotas.empty:
            st.info("No hay cuotas para marcar.")
        else:
            sel_pay = st.selectbox("Selecciona Cuota ID", cuotas["ID"].tolist(), key="cr_pay_sel")
            f = cuotas[cuotas["ID"] == sel_pay].iloc[0]
            tipo_pago = str(f["Tipo"])
            caso_pago = str(f["Caso"])
            monto_pago = float(pd.to_numeric(f["Monto"], errors="coerce") or 0.0)

            colA, colB = st.columns([1,2])
            with colA:
                fecha_pago = st.date_input("Fecha del pago", value=date.today(), key="cr_pay_fecha")
            with colB:
                obs_pago = st.text_input("Observación", value=f"Pago manual por cuota ID {sel_pay}", key="cr_pay_obs")

            if st.button("💳 Registrar pago equivalente", disabled=is_readonly or monto_pago <= 0, key="cr_pay_btn"):
                if tipo_pago == "Honorarios":
                    new_id = next_id(pagos_honorarios)
                    pagos_honorarios = add_row(pagos_honorarios, {
                        "ID": new_id,
                        "Caso": normalize_key(caso_pago),
                        "Etapa": "",  # opcional: si quieres asociarlo a una etapa, lo puedes editar luego
                        "FechaPago": str(fecha_pago),
                        "Monto": float(monto_pago),
                        "Observacion": obs_pago
                    }, "pagos_honorarios")
                    save_df("pagos_honorarios", pagos_honorarios)
                    st.success("✅ Pago de honorarios creado (afecta dashboard y estado de cuotas)")
                    st.rerun()

                else:  # CuotaLitis
                    new_id = next_id(pagos_litis)
                    pagos_litis = add_row(pagos_litis, {
                        "ID": new_id,
                        "Caso": normalize_key(caso_pago),
                        "FechaPago": str(fecha_pago),
                        "Monto": float(monto_pago),
                        "Observacion": obs_pago
                    }, "pagos_litis")
                    save_df("pagos_litis", pagos_litis)
                    st.success("✅ Pago de litis creado (afecta dashboard y estado de cuotas)")
                    st.rerun()

        # ==========================================================
        # ESTADO DE CUOTAS
        # ==========================================================
        st.divider()
        st.markdown("### Estado de cuotas (usa pagos existentes)")
        estado = cuotas_status_all()
        if estado is None or estado.empty:
            st.info("No hay cuotas o no hay estado aún.")
        else:
            st.dataframe(estado.sort_values(["Caso","Tipo","NroCuota"], ascending=[True,True,True]), use_container_width=True)

        st.download_button("⬇️ Descargar cronograma (CSV)", cuotas.to_csv(index=False).encode("utf-8"), "cuotas.csv", key="cr_dl_csv")

# ==========================================================
# ACTUACIONES (FICHA por caso/cliente + historial + reporte)
# + ✅ Aranceles / Otros gastos
# + ✅ Pagado / Pendiente por cliente
# ==========================================================
if menu == "Actuaciones":
    st.subheader("🧾 Actuaciones – Ficha por caso y cliente")

    if casos.empty:
        st.info("Primero registra casos.")
    else:
        tab_reg, tab_hist, tab_rep = st.tabs(
            ["Registrar actuación", "Historial (desplegable)", "Reporte"]
        )

        exp_list = casos["Expediente"].tolist()

        # ==================================================
        # REGISTRAR
        # ==================================================
        with tab_reg:
            exp = st.selectbox("Expediente", exp_list, key="act_exp")
            fila_caso = casos[casos["Expediente"] == exp].iloc[0]
            cliente = str(fila_caso.get("Cliente",""))

            st.write(f"**Cliente:** {cliente}")

            with st.form("act_new_form"):
                fecha = st.date_input("Fecha", value=date.today())
                tipo = st.text_input("Tipo de actuación (ej: Demanda, Audiencia, Sentencia, Apelación...)")
                resumen = st.text_area("Resumen (amplio)", height=160)
                prox = st.text_input("Próxima acción (opcional)")
                prox_fecha = st.text_input("Fecha próxima acción (YYYY-MM-DD opcional)")
                link = st.text_input("Link OneDrive (opcional)")

                col1, col2 = st.columns(2)
                with col1:
                    aranceles = st.number_input(
                        "Aranceles / Costas (S/)",
                        min_value=0.0, step=50.0, value=0.0
                    )
                with col2:
                    otros_gastos = st.number_input(
                        "Otros gastos (S/)",
                        min_value=0.0, step=50.0, value=0.0
                    )

                # ✅ CHECK PAGADO / PENDIENTE
                gastos_pagado = st.checkbox(
                    "✅ Gastos pagados por el cliente",
                    value=False,
                    help="Si está marcado, estos gastos NO se consideran deuda"
                )

                notas = st.text_area("Notas (opcional)", height=100)
                submit = st.form_submit_button("Guardar actuación")

                if submit:
                    new_id = next_id(actuaciones)
                    actuaciones = add_row(actuaciones, {
                        "ID": new_id,
                        "Caso": normalize_key(exp),
                        "Cliente": cliente,
                        "Fecha": str(fecha),
                        "TipoActuacion": tipo,
                        "Resumen": resumen,
                        "ProximaAccion": prox,
                        "FechaProximaAccion": prox_fecha,
                        "LinkOneDrive": link,
                        "CostasAranceles": float(aranceles),
                        "Gastos": float(otros_gastos),
                        "GastosPagado": "1" if gastos_pagado else "0",
                        "Notas": notas
                    }, "actuaciones")
                    save_df("actuaciones", actuaciones)
                    st.success("✅ Actuación registrada")
                    st.rerun()

        # ==================================================
        # HISTORIAL
        # ==================================================
        with tab_hist:
            exp_h = st.selectbox(
                "Selecciona expediente para historial",
                exp_list,
                key="act_hist_exp"
            )
            hist = actuaciones[
                actuaciones["Caso"] == normalize_key(exp_h)
            ].copy()

            if hist.empty:
                st.info("No hay actuaciones registradas para este caso.")
            else:
                hist["_Fecha_dt"] = hist["Fecha"].apply(
                    lambda x: pd.to_datetime(x, errors="coerce")
                )
                hist.sort_values(
                    ["_Fecha_dt","ID"],
                    ascending=[False, False],
                    inplace=True
                )

                for _, r in hist.iterrows():
                    titulo = f"{r['Fecha']} – {r['TipoActuacion']} (ID {r['ID']})"
                    with st.expander(titulo, expanded=False):
                        st.write(f"**Cliente:** {r.get('Cliente','')}")
                        st.write(f"**Resumen:** {r.get('Resumen','')}")
                        st.write(f"**Próxima acción:** {r.get('ProximaAccion','')}")
                        st.write(f"**Fecha próxima acción:** {r.get('FechaProximaAccion','')}")
                        if str(r.get("LinkOneDrive","")).strip():
                            st.markdown(f"**Link OneDrive:** {r.get('LinkOneDrive')}")

                        st.write(f"**Aranceles / Costas:** S/ {money(r.get('CostasAranceles',0)):,.2f}")
                        st.write(f"**Otros gastos:** S/ {money(r.get('Gastos',0)):,.2f}")

                        estado = "✅ Pagado" if str(r.get("GastosPagado","0")) == "1" else "⏳ Pendiente"
                        st.write(f"**Estado de gastos:** {estado}")

                        st.write(f"**Notas:** {r.get('Notas','')}")

                st.divider()
                st.markdown("### Editar / borrar actuación (por ID)")
                sel = st.selectbox(
                    "Actuación ID",
                    hist["ID"].tolist(),
                    key="act_edit_id"
                )
                fila = actuaciones[actuaciones["ID"] == sel].iloc[0]

                with st.form("act_edit_form"):
                    fecha_e = st.text_input("Fecha (YYYY-MM-DD)", value=str(fila["Fecha"]))
                    tipo_e = st.text_input("TipoActuacion", value=str(fila["TipoActuacion"]))
                    resumen_e = st.text_area("Resumen", value=str(fila["Resumen"]), height=160)
                    prox_e = st.text_input("ProximaAccion", value=str(fila["ProximaAccion"]))
                    prox_fecha_e = st.text_input(
                        "FechaProximaAccion",
                        value=str(fila["FechaProximaAccion"])
                    )
                    link_e = st.text_input(
                        "LinkOneDrive",
                        value=str(fila.get("LinkOneDrive",""))
                    )

                    col1, col2 = st.columns(2)
                    with col1:
                        aranceles_e = st.number_input(
                            "Aranceles / Costas (S/)",
                            min_value=0.0,
                            value=money(fila.get("CostasAranceles",0)),
                            step=50.0
                        )
                    with col2:
                        gastos_e = st.number_input(
                            "Otros gastos (S/)",
                            min_value=0.0,
                            value=money(fila.get("Gastos",0)),
                            step=50.0
                        )

                    gastos_pagado_e = st.checkbox(
                        "✅ Gastos pagados por el cliente",
                        value=str(fila.get("GastosPagado","0")) == "1"
                    )

                    notas_e = st.text_area(
                        "Notas",
                        value=str(fila["Notas"]),
                        height=100
                    )
                    submit = st.form_submit_button("Guardar cambios")

                    if submit:
                        idx = actuaciones.index[actuaciones["ID"] == sel][0]
                        actuaciones.loc[idx, [
                            "Fecha","TipoActuacion","Resumen",
                            "ProximaAccion","FechaProximaAccion",
                            "LinkOneDrive","CostasAranceles",
                            "Gastos","GastosPagado","Notas"
                        ]] = [
                            fecha_e, tipo_e, resumen_e,
                            prox_e, prox_fecha_e,
                            link_e, float(aranceles_e),
                            float(gastos_e),
                            "1" if gastos_pagado_e else "0",
                            notas_e
                        ]
                        save_df("actuaciones", actuaciones)
                        st.success("✅ Actualizado")
                        st.rerun()

                if st.button("🗑️ Borrar actuación", key="act_del_btn"):
                    actuaciones = actuaciones[actuaciones["ID"] != sel].copy()
                    save_df("actuaciones", actuaciones)
                    st.success("✅ Eliminado")
                    st.rerun()

        # ==================================================
        # REPORTE (GASTOS DEL CLIENTE – SIN SALDOS AMBIGUOS)
        # ==================================================
        with tab_rep:
            st.markdown("### Reporte de historial de actuaciones")

            exp_r = st.selectbox(
                "Expediente para reporte",
                exp_list,
                key="act_rep_exp"
            )

            rep = actuaciones[
                actuaciones["Caso"] == normalize_key(exp_r)
            ].copy()

            if rep.empty:
                st.info("No hay actuaciones registradas para este expediente.")
            else:
                # Normalizar valores
                rep["CostasAranceles"] = pd.to_numeric(rep.get("CostasAranceles", 0), errors="coerce").fillna(0.0)
                rep["Gastos"] = pd.to_numeric(rep.get("Gastos", 0), errors="coerce").fillna(0.0)
                rep["GastosPagado"] = rep.get("GastosPagado", "0").astype(str)

                # ✅ columnas explícitas (NO ambiguas)
                rep["Gasto Pendiente Cliente"] = rep.apply(
                    lambda r: (r["CostasAranceles"] + r["Gastos"]) if r["GastosPagado"] != "1" else 0.0,
                    axis=1
                )
                rep["Gasto Pagado Cliente"] = rep.apply(
                    lambda r: (r["CostasAranceles"] + r["Gastos"]) if r["GastosPagado"] == "1" else 0.0,
                    axis=1
                )

                total_pendiente = rep["Gasto Pendiente Cliente"].sum()
                total_pagado = rep["Gasto Pagado Cliente"].sum()

                c1, c2 = st.columns(2)
                c1.metric("⏳ Pendiente de pago por el cliente", f"S/ {total_pendiente:,.2f}")
                c2.metric("✅ Pagado por el cliente", f"S/ {total_pagado:,.2f}")

                # ✅ dataframe SIN columnas que parezcan saldo
                st.dataframe(
                    rep[[
                        "Fecha",
                        "TipoActuacion",
                        "Gasto Pendiente Cliente",
                        "Gasto Pagado Cliente",
                        "Notas"
                    ]].sort_values("Fecha", ascending=False),
                    use_container_width=True
                )

                st.download_button(
                    "⬇️ Descargar reporte de actuaciones (CSV)",
                    rep.to_csv(index=False).encode("utf-8"),
                    f"reporte_actuaciones_{exp_r.replace('/','_')}.csv"
                )

# ==========================================================
# CONSULTAS: autoguardado + abogado + costo + proforma + historial + reporte ingresos
# ==========================================================
if menu == "Consultas":
    st.subheader("🗂️ Consultas – Registro, proforma e historial (con autoguardado)")

    # -------- AUTOGUARDADO EN SESIÓN --------
    if "consulta_draft" not in st.session_state:
        st.session_state.consulta_draft = {
            "Fecha": str(date.today()),
            "Cliente": "",
            "Caso": "",
            "Abogado": "",
            "Consulta": "",
            "Estrategia": "",
            "CostoConsulta": 0.0,
            "HonorariosPropuestos": 0.0,
            "LinkOneDrive": "",
            "Notas": ""
        }

    draft = st.session_state.consulta_draft

    tab_new, tab_hist, tab_rep = st.tabs(["Nueva consulta", "Historial (desplegable)", "Reporte / Ingresos"])

    exp_list = casos["Expediente"].tolist() if not casos.empty else []
    clientes_list = clientes["Nombre"].tolist() if not clientes.empty else []
    abogados_list = abogados["Nombre"].tolist() if not abogados.empty else []

    # ==========================================================
    # TAB: NUEVA (SIN FORM -> autoguardado real al escribir)
    # ==========================================================
    with tab_new:
        st.markdown("### ✍️ Nueva consulta (autoguardado activado)")

        c1, c2, c3 = st.columns([1, 1, 1])
        with c1:
            fecha = st.date_input("Fecha", value=pd.to_datetime(draft["Fecha"]).date() if str(draft["Fecha"]).strip() else date.today(), key="cons_fecha")
            draft["Fecha"] = str(fecha)
        with c2:
            if clientes_list:
                cliente = st.selectbox("Cliente", clientes_list, index=clientes_list.index(draft["Cliente"]) if draft["Cliente"] in clientes_list else 0, key="cons_cliente")
            else:
                cliente = st.text_input("Cliente", value=draft["Cliente"], key="cons_cliente_txt")
            draft["Cliente"] = cliente
        with c3:
            if exp_list:
                opciones = [""] + exp_list
                caso = st.selectbox("Expediente (opcional)", opciones, index=opciones.index(draft["Caso"]) if draft["Caso"] in opciones else 0, key="cons_caso")
            else:
                caso = st.text_input("Expediente (opcional)", value=draft["Caso"], key="cons_caso_txt")
            draft["Caso"] = caso

        c4, c5 = st.columns([1, 1])
        with c4:
            if abogados_list:
                abogado = st.selectbox("Abogado a cargo", abogados_list, index=abogados_list.index(draft["Abogado"]) if draft["Abogado"] in abogados_list else 0, key="cons_abogado")
            else:
                abogado = st.text_input("Abogado a cargo", value=draft["Abogado"], key="cons_abogado_txt")
            draft["Abogado"] = abogado
        with c5:
            costo = st.number_input("Costo de consulta (S/)", min_value=0.0, step=50.0, value=float(draft.get("CostoConsulta", 0.0) or 0.0), key="cons_costo")
            draft["CostoConsulta"] = float(costo)

        consulta_txt = st.text_area("Consulta (amplio)", value=draft["Consulta"], height=180, key="cons_consulta")
        draft["Consulta"] = consulta_txt

        estrategia_txt = st.text_area("Propuesta de estrategia (amplio)", value=draft["Estrategia"], height=180, key="cons_estrategia")
        draft["Estrategia"] = estrategia_txt

        c6, c7 = st.columns([1, 1])
        with c6:
            honorarios_prop = st.number_input(
                "Honorarios propuestos (S/)",
                min_value=0.0,
                step=50.0,
                value=float(draft.get("HonorariosPropuestos", 0.0) or 0.0),
                key="cons_honor"
            )
            draft["HonorariosPropuestos"] = float(honorarios_prop)
        with c7:
            link = st.text_input("Link OneDrive (opcional)", value=draft["LinkOneDrive"], key="cons_link")
            draft["LinkOneDrive"] = link

        notas = st.text_area("Notas (opcional)", value=draft["Notas"], height=100, key="cons_notas")
        draft["Notas"] = notas

        # Proforma (texto)
        proforma = (
            f"PROFORMA – {APP_NAME}\n"
            f"Fecha: {draft['Fecha']}\n"
            f"Cliente: {draft['Cliente']}\n"
            f"Expediente: {draft['Caso']}\n"
            f"Abogado a cargo: {draft['Abogado']}\n"
            f"{'-'*60}\n"
            f"CONSULTA:\n{draft['Consulta']}\n\n"
            f"PROPUESTA DE ESTRATEGIA:\n{draft['Estrategia']}\n\n"
            f"COSTO DE CONSULTA: S/ {float(draft['CostoConsulta']):,.2f}\n"
            f"HONORARIOS PROPUESTOS: S/ {float(draft['HonorariosPropuestos']):,.2f}\n"
            f"{'-'*60}\n"
            f"Notas: {draft['Notas']}\n"
        )

        st.download_button(
            "⬇️ Descargar proforma (TXT)",
            proforma.encode("utf-8"),
            f"proforma_{date.today()}.txt",
            key="cons_pf_draft"
        )

        b1, b2, b3 = st.columns([1, 1, 2])
        with b1:
            if st.button("💾 Guardar consulta y proforma", key="cons_save"):
                new_id = next_id(consultas)
                consultas = add_row(consultas, {
                    "ID": new_id,
                    "Fecha": str(draft["Fecha"]),
                    "Cliente": draft["Cliente"],
                    "Caso": normalize_key(draft["Caso"]),
                    "Abogado": draft["Abogado"],
                    "Consulta": draft["Consulta"],
                    "Estrategia": draft["Estrategia"],
                    "CostoConsulta": float(draft["CostoConsulta"]),
                    "HonorariosPropuestos": float(draft["HonorariosPropuestos"]),
                    "Proforma": proforma,
                    "LinkOneDrive": draft["LinkOneDrive"],
                    "Notas": draft["Notas"]
                }, "consultas")
                save_df("consultas", consultas)

                # ✅ limpiar borrador (autoguardado) al guardar
                st.session_state.consulta_draft = {
                    "Fecha": str(date.today()),
                    "Cliente": "",
                    "Caso": "",
                    "Abogado": "",
                    "Consulta": "",
                    "Estrategia": "",
                    "CostoConsulta": 0.0,
                    "HonorariosPropuestos": 0.0,
                    "LinkOneDrive": "",
                    "Notas": ""
                }
                st.success("✅ Consulta guardada")
                st.rerun()

        with b2:
            if st.button("🧹 Limpiar borrador", key="cons_clear"):
                st.session_state.consulta_draft = {
                    "Fecha": str(date.today()),
                    "Cliente": "",
                    "Caso": "",
                    "Abogado": "",
                    "Consulta": "",
                    "Estrategia": "",
                    "CostoConsulta": 0.0,
                    "HonorariosPropuestos": 0.0,
                    "LinkOneDrive": "",
                    "Notas": ""
                }
                st.success("✅ Borrador limpiado")
                st.rerun()

        with b3:
            st.info("✅ Autoguardado activo: si la app se recarga, tus campos se mantienen en esta sesión.")

    # ==========================================================
    # TAB: HISTORIAL (DESPLEGABLE) + EDITAR/BORRAR
    # ==========================================================
    with tab_hist:
        if consultas.empty:
            st.info("Aún no hay consultas registradas.")
        else:
            def label_consulta(row):
                return f"ID {row['ID']} – {row.get('Fecha','')} – {row.get('Cliente','')} – {row.get('Abogado','')}"
            consultas_view = consultas.copy()
            consultas_view["_label"] = consultas_view.apply(label_consulta, axis=1)

            sel_label = st.selectbox("Selecciona una consulta", consultas_view["_label"].tolist(), key="cons_hist_sel")
            row = consultas_view[consultas_view["_label"] == sel_label].iloc[0]

            st.markdown("### Detalle de consulta")
            st.write(f"**Fecha:** {row.get('Fecha','')}")
            st.write(f"**Cliente:** {row.get('Cliente','')}")
            st.write(f"**Expediente:** {row.get('Caso','')}")
            st.write(f"**Abogado a cargo:** {row.get('Abogado','')}")
            st.write(f"**Costo consulta:** S/ {money(row.get('CostoConsulta',0)):,.2f}")

            if str(row.get("LinkOneDrive","")).strip():
                st.markdown(f"**Link OneDrive:** {row.get('LinkOneDrive')}")

            st.markdown("**Consulta:**")
            st.write(row.get("Consulta",""))
            st.markdown("**Estrategia:**")
            st.write(row.get("Estrategia",""))

            st.markdown(f"**Honorarios propuestos:** S/ {money(row.get('HonorariosPropuestos',0)):,.2f}")

            st.divider()
            st.markdown("### Proforma")
            st.text_area("Proforma (texto)", value=str(row.get("Proforma","")), height=260, key="cons_pf_view")
            st.download_button(
                "⬇️ Descargar proforma seleccionada (TXT)",
                str(row.get("Proforma","")).encode("utf-8"),
                f"proforma_consulta_{row['ID']}.txt",
                key="cons_pf_dl_hist"
            )

            st.divider()
            st.markdown("### Editar / borrar consulta (por ID)")
            sel_id = int(row["ID"])
            fila = consultas[consultas["ID"] == sel_id].iloc[0]

            with st.form("cons_edit_form"):
                fecha_e = st.text_input("Fecha", value=str(fila.get("Fecha","")))
                cliente_e = st.text_input("Cliente", value=str(fila.get("Cliente","")))
                caso_e = st.text_input("Expediente", value=str(fila.get("Caso","")))
                abogado_e = st.text_input("Abogado a cargo", value=str(fila.get("Abogado","")))
                costo_e = st.number_input("Costo consulta (S/)", min_value=0.0, value=money(fila.get("CostoConsulta",0)), step=50.0)

                consulta_e = st.text_area("Consulta", value=str(fila.get("Consulta","")), height=160)
                estrategia_e = st.text_area("Estrategia", value=str(fila.get("Estrategia","")), height=160)
                honor_e = st.number_input("Honorarios propuestos (S/)", min_value=0.0, value=money(fila.get("HonorariosPropuestos",0)), step=50.0)

                link_e = st.text_input("Link OneDrive", value=str(fila.get("LinkOneDrive","")))
                notas_e = st.text_area("Notas", value=str(fila.get("Notas","")), height=100)

                submit = st.form_submit_button("Guardar cambios")
                if submit:
                    proforma_new = (
                        f"PROFORMA – {APP_NAME}\n"
                        f"Fecha: {fecha_e}\n"
                        f"Cliente: {cliente_e}\n"
                        f"Expediente: {caso_e}\n"
                        f"Abogado a cargo: {abogado_e}\n"
                        f"{'-'*60}\n"
                        f"CONSULTA:\n{consulta_e}\n\n"
                        f"PROPUESTA DE ESTRATEGIA:\n{estrategia_e}\n\n"
                        f"COSTO DE CONSULTA: S/ {float(costo_e):,.2f}\n"
                        f"HONORARIOS PROPUESTOS: S/ {float(honor_e):,.2f}\n"
                        f"{'-'*60}\n"
                        f"Notas: {notas_e}\n"
                    )
                    idx = consultas.index[consultas["ID"] == sel_id][0]
                    consultas.loc[idx, ["Fecha","Cliente","Caso","Abogado","Consulta","Estrategia","CostoConsulta",
                                        "HonorariosPropuestos","Proforma","LinkOneDrive","Notas"]] = [
                        fecha_e, cliente_e, normalize_key(caso_e), abogado_e, consulta_e, estrategia_e,
                        float(costo_e), float(honor_e), proforma_new, link_e, notas_e
                    ]
                    save_df("consultas", consultas)
                    st.success("✅ Consulta actualizada")
                    st.rerun()

            if st.button("🗑️ Borrar consulta", key="cons_del"):
                consultas = consultas[consultas["ID"] != sel_id].copy()
                save_df("consultas", consultas)
                st.success("✅ Eliminada")
                st.rerun()

    # ==========================================================
    # TAB: REPORTE / INGRESOS
    # ==========================================================
    with tab_rep:
        st.markdown("### Reporte de consultas")
        st.dataframe(consultas.sort_values("Fecha", ascending=False), use_container_width=True)
        st.download_button("⬇️ Descargar consultas (CSV)", consultas.to_csv(index=False).encode("utf-8"), "consultas.csv", key="cons_csv_all")

        st.divider()
        st.markdown("### 📈 Ingresos por consultas (suma de costos)")
        if consultas.empty:
            st.info("No hay consultas registradas.")
        else:
            tmp = consultas.copy()
            tmp["CostoConsulta"] = pd.to_numeric(tmp.get("CostoConsulta", 0), errors="coerce").fillna(0.0)
            tmp["Abogado"] = tmp.get("Abogado", "").astype(str)

            rep = tmp.groupby("Abogado", as_index=False)["CostoConsulta"].sum().rename(
                columns={"CostoConsulta": "IngresosConsultas"}
            ).sort_values("IngresosConsultas", ascending=False)

            st.dataframe(rep, use_container_width=True)
            st.download_button("⬇️ Descargar ingresos (CSV)", rep.to_csv(index=False).encode("utf-8"), "ingresos_consultas.csv", key="cons_ing_csv")
# ==========================================================
# DOCUMENTOS
# ==========================================================
if menu == "Documentos":
    st.subheader("📎 Documentos")
    st.dataframe(documentos.sort_values("Fecha", ascending=False), use_container_width=True)
    st.download_button("⬇️ Descargar documentos (CSV)", documentos.to_csv(index=False).encode("utf-8"), "documentos.csv")

# ==========================================================
# PLANTILLAS DE CONTRATO (MEJORADO: CÓDIGOS + INSERTAR + VALIDAR + PREVIEW)
# ==========================================================
if menu == "Plantillas de Contrato":
    st.subheader("📝 Plantillas de Contrato (Modelos)")
    accion = st.radio("Acción", ["Nueva","Editar","Eliminar"], horizontal=True, key="tpl_accion")

    # -------------------------
    # Construcción dinámica de códigos disponibles (TODAS las columnas)
    # -------------------------
    def _tok(prefix: str, col: str) -> str:
        # Mantener exactamente el formato que usa build_context: {{PREFIX_<COLNAME_UPPER>}}
        return "{{" + f"{prefix}_{str(col).upper()}" + "}}"

    # Códigos básicos y económicos (manuales)
    base_tokens = {
        "{{EXPEDIENTE}}": "Número de expediente del caso seleccionado",
        "{{FECHA_HOY}}": "Fecha actual (fecha de generación del contrato)"
    }

    econ_tokens = {
        "{{MONTO_PACTADO}}": "Total de honorarios pactados del caso",
        "{{HONORARIO_PRIMERA_INSTANCIA}}": "Honorarios pactados por Primera Instancia",
        "{{HONORARIO_SEGUNDA_INSTANCIA}}": "Honorarios pactados por Segunda Instancia",
        "{{HONORARIO_CASACION}}": "Honorarios pactados por Casación",
        "{{HONORARIO_OTROS}}": "Otros honorarios pactados",
        "{{CUOTA_LITIS_BASE}}": "Monto base de cuota litis",
        "{{CUOTA_LITIS_PORCENTAJE}}": "Porcentaje pactado de cuota litis",
        "{{CRONOGRAMA_PAGOS}}": "Bloque/tabla con el cronograma completo de cuotas"
    }

    # Tokens dinámicos por columnas (caso/cliente/abogado)
    caso_cols = list(casos.columns) if 'casos' in globals() and casos is not None and not casos.empty else (list(casos.columns) if 'casos' in globals() else [])
    cliente_cols = list(clientes.columns) if 'clientes' in globals() and clientes is not None and not clientes.empty else (list(clientes.columns) if 'clientes' in globals() else [])
    abogado_cols = list(abogados.columns) if 'abogados' in globals() and abogados is not None and not abogados.empty else (list(abogados.columns) if 'abogados' in globals() else [])

    # Descripciones “bonitas” para algunos campos clave (opcional)
    pretty = {
        # Caso / Contraparte / Judicial
        "CONTRAPARTE": "Contraparte (nombre)",
        "CONTRAPARTEDOC": "Documento de la contraparte (DNI/RUC)",
        "JUZGADO": "Juzgado",
        "DISTRITOJUDICIAL": "Distrito Judicial",
        "MATERIA": "Materia",
        "INSTANCIA": "Instancia",
        "PRETENSION": "Pretensión",
        "ESTADOCASO": "Estado del caso",
        "FECHAINICIO": "Fecha de inicio del caso",
        # Cliente
        "TIPOCLIENTE": "Tipo de cliente (Natural/Jurídica)",
        "RAZONSOCIAL": "Razón social (persona jurídica)",
        "RUC": "RUC (persona jurídica)",
        "REPRESENTANTELEGAL": "Representante legal",
        "DNI": "DNI",
        "DIRECCION": "Dirección",
        "CORREO": "Correo",
        "CELULAR": "Celular",
        # Abogado
        "COLEGIATURA": "Colegiatura",
        "DOMICILIO PROCESAL": "Domicilio procesal",
        "CASILLA ELECTRONICA": "Casilla electrónica",
        "CASILLA JUDICIAL": "Casilla judicial",
    }

    def _desc_dynamic(prefix: str, col: str) -> str:
        key = str(col).upper()
        if key in pretty:
            return f"{prefix}: {pretty[key]}"
        return f"{prefix}: Campo '{col}'"

    # Armar catálogo completo
    catalog = []

    # BÁSICOS
    for k, v in base_tokens.items():
        catalog.append(("BÁSICOS", k, v))

    # CASO (todas las columnas)
    for col in caso_cols:
        catalog.append(("CASO (todas las columnas)", _tok("CASO", col), _desc_dynamic("CASO", col)))

    # CLIENTE (todas las columnas)
    for col in cliente_cols:
        catalog.append(("CLIENTE (todas las columnas)", _tok("CLIENTE", col), _desc_dynamic("CLIENTE", col)))

    # ABOGADO (todas las columnas)
    for col in abogado_cols:
        catalog.append(("ABOGADO (todas las columnas)", _tok("ABOGADO", col), _desc_dynamic("ABOGADO", col)))

    # ECONÓMICOS
    for k, v in econ_tokens.items():
        catalog.append(("ECONÓMICOS", k, v))

    # Lista de tokens válidos (para validador)
    allowed_tokens = set([t for _, t, _ in catalog])

    # -------------------------
    # Utilidades: insertar token, validar, preview
    # -------------------------
    import re
    token_pat = re.compile(r"{{[^{}]+}}")

    def _unknown_tokens(text: str):
        found = set(token_pat.findall(text or ""))
        unknown = sorted([t for t in found if t not in allowed_tokens])
        return unknown, sorted(found)

    def _render_preview_if_possible(texto: str):
        # usa build_context/render_template si existen
        if 'build_context' in globals() and 'render_template' in globals():
            if casos.empty:
                st.info("No hay casos para previsualizar.")
                return
            exp_prev = st.selectbox("Expediente para previsualizar", casos["Expediente"].tolist(), key="tpl_prev_exp")
            try:
                ctx = build_context(exp_prev)
                prev = render_template(texto or "", ctx)
                st.text_area("Vista previa (con datos reales)", value=prev, height=260, key="tpl_prev_out")
            except Exception as e:
                st.warning(f"No se pudo previsualizar: {e}")
        else:
            st.info("Vista previa deshabilitada: no se encontró build_context/render_template en la app.")

    # -------------------------
    # Panel fijo de códigos (visible mientras creas/editar)
    # -------------------------
    with st.expander("📌 Códigos disponibles (clic para insertar)", expanded=True):
        st.caption("Tip: Usa el buscador y presiona **Insertar** para añadir el código en tu contenido.")
        qcode = st.text_input("Buscar código o descripción", value="", key="tpl_code_search").strip().lower()

        # filtrar catálogo por búsqueda
        filtered = []
        for grp, tok, desc in catalog:
            if not qcode or qcode in tok.lower() or qcode in desc.lower() or qcode in grp.lower():
                filtered.append((grp, tok, desc))

        # selector de token
        options = [f"[{grp}] {tok} — {desc}" for grp, tok, desc in filtered]
        if not options:
            st.info("Sin resultados para ese filtro.")
        else:
            sel = st.selectbox("Selecciona un código", options, key="tpl_code_pick")

            # extraer el token real
            tok_selected = sel.split("] ", 1)[1].split(" — ", 1)[0].strip()

            st.code(tok_selected, language="text")

            st.caption("Puedes copiarlo manualmente o usar Insertar.")
            # Insertar se hará sobre una variable de sesión que usamos en el editor
            # (se define más abajo dependiendo de Nueva/Editar)
            st.session_state["tpl_tok_to_insert"] = tok_selected

    # -------------------------
    # NUEVA
    # -------------------------
    if accion == "Nueva":
        # Estado editable en session_state para poder “insertar”
        if "tpl_new_contenido" not in st.session_state:
            st.session_state["tpl_new_contenido"] = ""

        # Insertar token si el usuario lo pidió
        if st.button("➕ Insertar código en contenido (Nueva)", key="tpl_new_insert_btn"):
            tok = st.session_state.get("tpl_tok_to_insert", "")
            if tok:
                st.session_state["tpl_new_contenido"] = (st.session_state["tpl_new_contenido"] or "") + "\n" + tok

        with st.form("tpl_new"):
            nombre = st.text_input("Nombre", key="tpl_new_nombre")
            contenido = st.text_area("Contenido", height=300, key="tpl_new_contenido")
            notas = st.text_input("Notas", value="", key="tpl_new_notas")

            # Validador en vivo (dentro del form)
            unknown, found = _unknown_tokens(contenido)
            if unknown:
                st.warning("⚠️ Códigos desconocidos detectados (no serán reemplazados):\n" + "\n".join(unknown))
            else:
                st.success("✅ Códigos válidos (o no se detectaron códigos).")

            st.caption(f"Códigos detectados: {len(found)}")

            # Vista previa (si existe soporte)
            st.markdown("### 👁️ Vista previa")
            _render_preview_if_possible(contenido)

            submit = st.form_submit_button("Guardar plantilla")

            if submit:
                new_id = next_id(plantillas)
                plantillas = add_row(plantillas, {
                    "ID": new_id,
                    "Nombre": nombre,
                    "Contenido": contenido,
                    "Notas": notas,
                    "Creado": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                }, "plantillas")
                save_df("plantillas", plantillas)
                st.success("✅ Plantilla creada")
                # limpiar borrador
                st.session_state["tpl_new_contenido"] = ""
                st.rerun()

    # -------------------------
    # EDITAR
    # -------------------------
    elif accion == "Editar":
        if plantillas.empty:
            st.info("No hay plantillas.")
        else:
            sel = st.selectbox(
                "Selecciona plantilla",
                plantillas["ID"].tolist(),
                format_func=lambda x: f"ID {x} – {plantillas[plantillas['ID']==x].iloc[0]['Nombre']}",
                key="tpl_edit_sel"
            )

            fila = plantillas[plantillas["ID"] == sel].iloc[0]

            # Inicializar borrador editable al cambiar selección
            key_state = f"tpl_edit_contenido_{sel}"
            if key_state not in st.session_state:
                st.session_state[key_state] = str(fila.get("Contenido",""))

            # Insertar token en editor de Editar
            if st.button("➕ Insertar código en contenido (Editar)", key="tpl_edit_insert_btn"):
                tok = st.session_state.get("tpl_tok_to_insert", "")
                if tok:
                    st.session_state[key_state] = (st.session_state[key_state] or "") + "\n" + tok

            with st.form("tpl_edit"):
                nombre = st.text_input("Nombre", value=str(fila["Nombre"]), key=f"tpl_edit_nombre_{sel}")
                contenido = st.text_area("Contenido", value=st.session_state[key_state], height=300, key=key_state)
                notas = st.text_input("Notas", value=str(fila["Notas"]), key=f"tpl_edit_notas_{sel}")

                unknown, found = _unknown_tokens(contenido)
                if unknown:
                    st.warning("⚠️ Códigos desconocidos detectados (no serán reemplazados):\n" + "\n".join(unknown))
                else:
                    st.success("✅ Códigos válidos (o no se detectaron códigos).")

                st.caption(f"Códigos detectados: {len(found)}")

                st.markdown("### 👁️ Vista previa")
                _render_preview_if_possible(contenido)

                submit = st.form_submit_button("Guardar cambios")

                if submit:
                    idx = plantillas.index[plantillas["ID"] == sel][0]
                    plantillas.loc[idx, ["Nombre","Contenido","Notas"]] = [nombre, contenido, notas]
                    save_df("plantillas", plantillas)
                    st.success("✅ Plantilla actualizada")
                    st.rerun()

    # -------------------------
    # ELIMINAR
    # -------------------------
    elif accion == "Eliminar":
        if plantillas.empty:
            st.info("No hay plantillas.")
        else:
            sel = st.selectbox(
                "Selecciona plantilla a eliminar",
                plantillas["ID"].tolist(),
                format_func=lambda x: f"ID {x} – {plantillas[plantillas['ID']==x].iloc[0]['Nombre']}",
                key="tpl_del_sel"
            )

            st.warning("⚠️ Esta acción no se puede deshacer")
            confirm = st.text_input("Escribe ELIMINAR para confirmar", key="tpl_del_confirm")

            if st.button("🗑️ Eliminar plantilla", key="tpl_del_btn", disabled=confirm.strip().upper() != "ELIMINAR"):
                plantillas = plantillas[plantillas["ID"] != sel].copy()
                save_df("plantillas", plantillas)
                st.success("✅ Plantilla eliminada")
                st.rerun()

    st.divider()
    st.dataframe(plantillas, use_container_width=True)
    st.download_button(
        "⬇️ Descargar plantillas (CSV)",
        plantillas.to_csv(index=False).encode("utf-8"),
        "plantillas_contratos.csv",
        key="tpl_dl_csv"
    )

# ==========================================================
# GENERAR CONTRATO
# ==========================================================
def build_context(expediente: str):
    expediente = normalize_key(expediente)
    ctx = {}

    # =========================
    # CASO (todas las columnas)
    # =========================
    caso_row = casos[casos["Expediente"] == expediente]
    if not caso_row.empty:
        caso = caso_row.iloc[0].to_dict()
        for k, v in caso.items():
            ctx[f"{{{{CASO_{str(k).upper()}}}}}"] = str(v)

    # =========================
    # CLIENTE (todas las columnas)
    # =========================
    if not caso_row.empty:
        cliente_nombre = str(caso_row.iloc[0].get("Cliente",""))
        cli_row = clientes[clientes["Nombre"].astype(str) == cliente_nombre]
        if not cli_row.empty:
            cli = cli_row.iloc[0].to_dict()
            for k, v in cli.items():
                ctx[f"{{{{CLIENTE_{str(k).upper()}}}}}"] = str(v)

    # =========================
    # ABOGADO (todas las columnas)
    # =========================
    if not caso_row.empty:
        abogado_nombre = str(caso_row.iloc[0].get("Abogado",""))
        ab_row = abogados[abogados["Nombre"].astype(str) == abogado_nombre]
        if not ab_row.empty:
            ab = ab_row.iloc[0].to_dict()
            for k, v in ab.items():
                ctx[f"{{{{ABOGADO_{str(k).upper()}}}}}"] = str(v)

    # =========================
    # CAMPOS BÁSICOS
    # =========================
    ctx["{{EXPEDIENTE}}"] = expediente
    ctx["{{FECHA_HOY}}"] = date.today().strftime("%Y-%m-%d")

    return ctx


def render_template(text: str, ctx: dict) -> str:
    out = text
    for k, v in ctx.items():
        out = out.replace(k, v)
    return out


# ==========================================================
# WORD (.DOCX) – soporte seguro
# ==========================================================
from io import BytesIO

try:
    from docx import Document
except Exception:
    Document = None


def generar_docx(texto: str, titulo="Contrato"):
    """
    Genera un DOCX en memoria para descargar.
    Si python-docx no está instalado, devuelve None.
    """
    if Document is None:
        return None

    doc = Document()
    doc.add_heading(str(titulo), level=1)

    for linea in str(texto).split("\n"):
        doc.add_paragraph(linea)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================================
# UI: Generar Contrato + Editar manualmente + Guardar TXT/DOCX
# ==========================================================
if menu == "Generar Contrato":
    st.subheader("📄 Generar contrato automáticamente")
    if casos.empty:
        st.info("Primero registra casos.")
    elif plantillas.empty:
        st.info("Primero crea una plantilla.")
    else:
        exp = st.selectbox("Expediente", casos["Expediente"].tolist(), key="gc_exp")
        tpl_id = st.selectbox("Plantilla ID", plantillas["ID"].tolist(), key="gc_tpl")
        tpl = plantillas[plantillas["ID"] == tpl_id].iloc[0]

        ctx = build_context(exp)
        generado = render_template(str(tpl["Contenido"]), ctx)

        # ✅ NUEVO: edición manual opcional
        st.markdown("### ✏️ Edición manual (opcional)")
        editar_manual = st.checkbox(
            "Modificar manualmente este contrato",
            value=False,
            key="gc_edit_mode"
        )

        if editar_manual:
            texto_final = st.text_area(
                "Contrato (editable)",
                value=generado,
                height=420,
                key="gc_text_edit"
            )
        else:
            st.text_area("Vista previa", value=generado, height=350)
            texto_final = generado

        # Nombre archivo base
        nombre_archivo = f"Contrato_{exp.replace('/','_')}_{str(tpl['Nombre']).replace(' ','_')}.txt"

        # ✅ Descargar TXT (si editó, descarga lo editado)
        st.download_button(
            "⬇️ Descargar contrato (TXT)",
            data=texto_final.encode("utf-8"),
            file_name=nombre_archivo,
            key="gc_dl_txt"
        )

        # ✅ Descargar WORD (si editó, descarga lo editado)
        docx_file = generar_docx(texto_final, titulo=str(tpl.get("Nombre","Contrato")))
        if docx_file is None:
            st.warning("⚠️ No se pudo generar Word. Falta instalar python-docx en requirements.txt")
        else:
            st.download_button(
                "⬇️ Descargar contrato (WORD)",
                data=docx_file,
                file_name=nombre_archivo.replace(".txt", ".docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_word_{exp}_{tpl_id}"
            )

        # ✅ Guardar en generados/ (TXT + DOCX si existe)
        if st.button("💾 Guardar en carpeta generados/", key="gc_save"):
            # Guardar TXT (si editó, guarda editado)
            out_txt = os.path.join(GENERADOS_DIR, nombre_archivo)
            with open(out_txt, "w", encoding="utf-8") as f:
                f.write(texto_final)

            # Guardar DOCX (si python-docx está disponible)
            out_docx = None
            if docx_file is not None:
                out_docx = os.path.join(GENERADOS_DIR, nombre_archivo.replace(".txt", ".docx"))
                with open(out_docx, "wb") as f:
                    f.write(docx_file.getbuffer())

            st.success("✅ Guardado en carpeta generados/")
            st.info(out_txt)
            if out_docx:
                st.info(out_docx)

# ==========================================================
# USUARIOS + ROLES + PERMISOS (CON CONTRASEÑA + VINCULO ABOGADO + NOMBRE/DNI)
# ==========================================================
if menu == "Usuarios":
    st.subheader("👥 Usuarios del Sistema")

    ROLES_DISPONIBLES = [
        "Admin",
        "Abogado",
        "Personal Administrativo",
        "Secretaria/o",
        "Solo Lectura"
    ]

    DEFAULT_PERMISOS = {
        "Admin": {"ver": True, "agregar": True, "modificar": True, "borrar": True},
        "Abogado": {"ver": True, "agregar": True, "modificar": True, "borrar": True},
        "Personal Administrativo": {"ver": True, "agregar": True, "modificar": True, "borrar": False},
        "Secretaria/o": {"ver": True, "agregar": True, "modificar": False, "borrar": False},
        "Solo Lectura": {"ver": True, "agregar": False, "modificar": False, "borrar": False},
    }

    def _norm_role(x: str) -> str:
        return str(x or "").strip().lower()

    def _is_admin() -> bool:
        return _norm_role(st.session_state.get("rol", "")) in ["admin", "administrador"]

    # Asegurar columnas extra en usuarios
    try:
        if "usuarios" in SCHEMAS:
            for extra in ["NombreCompleto", "DNI"]:
                if extra not in SCHEMAS["usuarios"]:
                    SCHEMAS["usuarios"].append(extra)
    except Exception:
        pass

    # =========================
    # PERMISOS
    # =========================
    permisos = load_df("permisos")
    if permisos is None or permisos.empty:
        rows = []
        for rol, perms in DEFAULT_PERMISOS.items():
            rows.append({
                "Rol": rol,
                "Ver": int(perms["ver"]),
                "Agregar": int(perms["agregar"]),
                "Modificar": int(perms["modificar"]),
                "Borrar": int(perms["borrar"]),
            })
        permisos = pd.DataFrame(rows)
        save_df("permisos", permisos)

    def has_perm(accion: str) -> bool:
        rol_norm = _norm_role(st.session_state.get("rol", ""))
        if not rol_norm:
            return False
        tmp = permisos.copy()
        tmp["Rol_norm"] = tmp["Rol"].astype(str).str.strip().str.lower()
        fila = tmp[tmp["Rol_norm"] == rol_norm]
        if fila.empty:
            return False
        col = {"ver": "Ver", "agregar": "Agregar", "modificar": "Modificar", "borrar": "Borrar"}.get(accion)
        if not col:
            return False
        return bool(int(fila.iloc[0].get(col, 0)))

    globals()["has_perm"] = has_perm

    if _is_admin():
        with st.expander("⚙️ Configuración de permisos por rol", expanded=False):
            st.caption("Define qué puede hacer cada rol en el sistema.")
            edit = permisos.copy()
            for i in edit.index:
                st.markdown(f"**{edit.at[i,'Rol']}**")
                c1, c2, c3, c4 = st.columns(4)
                edit.at[i,"Ver"] = int(c1.checkbox("Ver", value=bool(int(edit.at[i,"Ver"])), key=f"p_ver_{i}"))
                edit.at[i,"Agregar"] = int(c2.checkbox("Agregar", value=bool(int(edit.at[i,"Agregar"])), key=f"p_add_{i}"))
                edit.at[i,"Modificar"] = int(c3.checkbox("Modificar", value=bool(int(edit.at[i,"Modificar"])), key=f"p_mod_{i}"))
                edit.at[i,"Borrar"] = int(c4.checkbox("Borrar", value=bool(int(edit.at[i,"Borrar"])), key=f"p_del_{i}"))
                st.divider()
            if st.button("💾 Guardar permisos", key="p_save"):
                permisos = edit.copy()
                save_df("permisos", permisos)
                st.success("✅ Permisos actualizados")
                st.rerun()
    else:
        st.info("Solo Admin puede editar la matriz de permisos. Puedes ver usuarios, pero no cambiar permisos.")

    # =========================
    # ABOGADOS DISPONIBLES (solo lectura)
    # =========================
    df_ab = load_df("abogados")
    if df_ab is None:
        df_ab = pd.DataFrame()

    if not df_ab.empty and "Nombre" not in df_ab.columns:
        for alt in ["NOMBRE", "Nombre completo", "NombreCompleto"]:
            if alt in df_ab.columns:
                df_ab["Nombre"] = df_ab[alt]
                break
    if "Nombre" not in df_ab.columns:
        df_ab["Nombre"] = ""

    # Preferimos ID si existe; si no, usamos Nombre como clave
    if (not df_ab.empty) and ("ID" in df_ab.columns) and df_ab["ID"].astype(str).str.strip().ne("").any():
        key_col = "ID"
    else:
        key_col = "Nombre"

    if key_col in df_ab.columns:
        df_ab[key_col] = df_ab[key_col].astype(str).str.strip()
    df_ab["Nombre"] = df_ab["Nombre"].astype(str).str.strip()

    df_ab2 = df_ab.copy()
    df_ab2 = df_ab2[(df_ab2["Nombre"] != "") & (df_ab2.get(key_col, "").astype(str).str.strip() != "")]
    abogado_opts = df_ab2[key_col].astype(str).tolist() if (not df_ab2.empty and key_col in df_ab2.columns) else []
    abogado_label = {str(r[key_col]).strip(): str(r["Nombre"]).strip() for _, r in df_ab2.iterrows()}

    # =========================
    # GESTIÓN DE USUARIOS
    # =========================
    st.markdown("## 👤 Gestión de usuarios")
    accion_u = st.radio("Acción", ["Nuevo","Editar","Eliminar"], horizontal=True, key="usr_acc")

    # ---------- NUEVO ----------
    if accion_u == "Nuevo":
        with st.form("usr_new"):
            # 1) Rol primero
            rol = st.selectbox("Rol", ROLES_DISPONIBLES, key="usr_new_role")

            # 2) Credenciales siempre
            usuario = st.text_input("Usuario", key="usr_new_user")
            pwd = st.text_input("Contraseña", type="password", key="usr_new_pwd")
            pwd2 = st.text_input("Repetir contraseña", type="password", key="usr_new_pwd2")

            # 3) Datos según rol
            abogado_id = ""
            nombre_completo = ""
            dni_personal = ""

            if rol == "Abogado":
                st.markdown("### 👨‍⚖️ Abogado asociado (obligatorio)")
                if not abogado_opts:
                    st.error("❌ No hay abogados registrados para asociar. Registra abogados primero.")
                else:
                    abogado_id = st.selectbox(
                        "Selecciona abogado",
                        options=abogado_opts,
                        format_func=lambda x: f"{abogado_label.get(str(x),'')} ({key_col} {x})",
                        key="usr_new_abogado"
                    )
            else:
                st.markdown("### 👤 Datos personales")
                nombre_completo = st.text_input("Nombre completo", key="usr_new_nombre")
                dni_personal = st.text_input("DNI", key="usr_new_dni")

            # si rol=abogado y no hay abogados, deshabilitar submit
            disabled_submit = (rol == "Abogado" and not abogado_opts)
            submit = st.form_submit_button("Crear usuario", disabled=disabled_submit)

            if submit:
                if not str(usuario).strip():
                    st.error("Usuario es obligatorio.")
                    st.stop()
                if (usuarios["Usuario"].astype(str) == str(usuario).strip()).any():
                    st.error("Ese usuario ya existe.")
                    st.stop()
                if not pwd:
                    st.error("Contraseña es obligatoria.")
                    st.stop()
                if pwd != pwd2:
                    st.error("Las contraseñas no coinciden.")
                    st.stop()
                if rol == "Abogado" and not str(abogado_id).strip():
                    st.error("Debes seleccionar un abogado asociado (obligatorio).")
                    st.stop()

                usuarios = add_row(usuarios, {
                    "Usuario": str(usuario).strip(),
                    "PasswordHash": sha256(pwd),
                    "Rol": rol,
                    "AbogadoID": str(abogado_id).strip() if rol == "Abogado" else "",
                    "Activo": "1",
                    "Creado": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "NombreCompleto": str(nombre_completo).strip() if rol != "Abogado" else "",
                    "DNI": str(dni_personal).strip() if rol != "Abogado" else "",
                }, "usuarios")
                save_df("usuarios", usuarios)

                # ✅ Mostrar con quién se vinculó y pasar a EDITAR automáticamente
                if rol == "Abogado":
                    st.success(f"✅ Usuario creado y asociado a: {abogado_label.get(str(abogado_id),'')} ({key_col} {abogado_id})")
                else:
                    st.success("✅ Usuario creado")

                # limpiar campos para evitar repetir "crear"
                for k in ["usr_new_user","usr_new_pwd","usr_new_pwd2","usr_new_nombre","usr_new_dni"]:
                    if k in st.session_state:
                        st.session_state[k] = ""

                # ir directo a Editar para "Actualizar" si se requiere
                st.session_state["usr_acc"] = "Editar"
                st.session_state["usr_edit_sel"] = str(usuario).strip()
                st.rerun()

    # ---------- EDITAR ----------
    elif accion_u == "Editar":
        if usuarios.empty:
            st.info("No hay usuarios.")
        else:
            sel_user = st.selectbox(
                "Selecciona usuario",
                usuarios["Usuario"].astype(str).tolist(),
                key="usr_edit_sel"
            )
            fila = usuarios[usuarios["Usuario"].astype(str) == str(sel_user)].iloc[0]

            with st.form("usr_edit"):
                # Rol primero
                rol_new = st.selectbox(
                    "Rol",
                    ROLES_DISPONIBLES,
                    index=ROLES_DISPONIBLES.index(str(fila.get("Rol","Solo Lectura"))) if str(fila.get("Rol","Solo Lectura")) in ROLES_DISPONIBLES else 0,
                    key="usr_edit_role"
                )

                usuario_new = st.text_input("Usuario", value=str(fila.get("Usuario","")), key="usr_edit_user")
                activo_new = st.selectbox("Activo", ["1","0"], index=0 if str(fila.get("Activo","1")) == "1" else 1, key="usr_edit_activo")

                st.markdown("### 🔑 Cambiar contraseña (opcional)")
                new_pwd = st.text_input("Nueva contraseña", type="password", key="usr_edit_pwd")
                new_pwd2 = st.text_input("Repetir nueva contraseña", type="password", key="usr_edit_pwd2")

                abogado_id_new = str(fila.get("AbogadoID","")).strip()
                nombre_new = str(fila.get("NombreCompleto","")).strip()
                dni_new = str(fila.get("DNI","")).strip()

                if rol_new == "Abogado":
                    st.markdown("### 👨‍⚖️ Abogado asociado (obligatorio)")
                    if not abogado_opts:
                        st.error("❌ No hay abogados registrados para asociar. Registra abogados primero.")
                        abogado_id_new = ""
                    else:
                        if abogado_id_new:
                            st.caption(f"Actualmente asociado a: {abogado_label.get(abogado_id_new,'')} ({key_col} {abogado_id_new})")
                        idx_def = abogado_opts.index(abogado_id_new) if abogado_id_new in abogado_opts else 0
                        abogado_id_new = st.selectbox(
                            "Selecciona abogado",
                            options=abogado_opts,
                            index=idx_def,
                            format_func=lambda x: f"{abogado_label.get(str(x),'')} ({key_col} {x})",
                            key="usr_edit_abogado"
                        )
                    nombre_new = ""
                    dni_new = ""
                else:
                    nombre_new = st.text_input("Nombre completo", value=nombre_new, key="usr_edit_nombre")
                    dni_new = st.text_input("DNI", value=dni_new, key="usr_edit_dni")
                    abogado_id_new = ""

                disabled_submit = (rol_new == "Abogado" and not abogado_opts)
                submit = st.form_submit_button("Actualizar usuario", disabled=disabled_submit)

                if submit:
                    if not str(usuario_new).strip():
                        st.error("Usuario es obligatorio.")
                        st.stop()
                    if (str(usuario_new).strip() != str(sel_user).strip()) and (usuarios["Usuario"].astype(str) == str(usuario_new).strip()).any():
                        st.error("Ese usuario ya existe.")
                        st.stop()
                    if new_pwd and new_pwd != new_pwd2:
                        st.error("Las contraseñas no coinciden.")
                        st.stop()
                    if rol_new == "Abogado" and not str(abogado_id_new).strip():
                        st.error("Debes seleccionar un abogado asociado (obligatorio).")
                        st.stop()

                    idx = usuarios.index[usuarios["Usuario"].astype(str) == str(sel_user)][0]
                    usuarios.at[idx, "Usuario"] = str(usuario_new).strip()
                    usuarios.at[idx, "Rol"] = rol_new
                    usuarios.at[idx, "Activo"] = str(activo_new)
                    usuarios.at[idx, "AbogadoID"] = str(abogado_id_new).strip()
                    usuarios.at[idx, "NombreCompleto"] = str(nombre_new).strip()
                    usuarios.at[idx, "DNI"] = str(dni_new).strip()

                    if new_pwd:
                        usuarios.at[idx, "PasswordHash"] = sha256(new_pwd)

                    save_df("usuarios", usuarios)
                    st.success("✅ Usuario actualizado")
                    st.rerun()

    # ---------- ELIMINAR ----------
    elif accion_u == "Eliminar":
        if usuarios.empty:
            st.info("No hay usuarios.")
        else:
            sel_user = st.selectbox("Selecciona usuario a eliminar", usuarios["Usuario"].astype(str).tolist(), key="usr_del_sel")
            st.warning("⚠️ Esta acción no se puede deshacer")
            confirm = st.text_input("Escribe ELIMINAR para confirmar", key="usr_del_confirm")
            if st.button("🗑️ Eliminar usuario", key="usr_del_btn", disabled=confirm.strip().upper() != "ELIMINAR"):
                usuarios = usuarios[usuarios["Usuario"].astype(str) != str(sel_user)].copy()
                save_df("usuarios", usuarios)
                st.success("✅ Usuario eliminado")
                st.rerun()

    st.divider()

    # Mostrar relación usuario → abogado en tabla
    df_show = usuarios.drop(columns=["PasswordHash"], errors="ignore").copy()
    if "AbogadoID" in df_show.columns:
        df_show["AbogadoAsociado"] = df_show["AbogadoID"].astype(str).map(lambda x: abogado_label.get(str(x), "") if str(x).strip() else "")
    st.dataframe(df_show, use_container_width=True)
# ==========================================================
# REPORTES (FILTRADOS POR ROL)
# ==========================================================
if menu == "Reportes":
    st.subheader("📊 Reportes")

    rol = str(st.session_state.get("rol","")).strip().lower()

    # Cargar casos y aplicar visibilidad por rol (BLOQUE A)
    df_casos_all = load_df("casos")
    df_casos_vis = filtrar_casos_por_rol(df_casos_all)

    # ------------------------------------------------------
    # REPORTE: CASOS VISIBLES
    # ------------------------------------------------------
    st.markdown("### 📁 Casos visibles")

    if df_casos_vis is None or df_casos_vis.empty:
        st.info("No tienes casos asignados o delegados.")
    else:
        columnas = [
            c for c in [
                "Expediente","Cliente","Abogado","Materia",
                "Instancia","EstadoCaso","Contraparte","DistritoJudicial"
            ]
            if c in df_casos_vis.columns
        ]

        st.dataframe(df_casos_vis[columnas], use_container_width=True)

        st.download_button(
            "⬇️ Descargar mis casos (CSV)",
            df_casos_vis.to_csv(index=False).encode("utf-8"),
            "mis_casos.csv"
        )

    st.divider()

    # ------------------------------------------------------
    # REPORTE: ACTUACIONES (solo de casos visibles)
    # ------------------------------------------------------
    st.markdown("### 🧾 Actuaciones")

    df_act = load_df("actuaciones")

    if df_act is None or df_act.empty or df_casos_vis is None or df_casos_vis.empty:
        st.info("No hay actuaciones para mostrar.")
    else:
        if "Caso" in df_act.columns and "Expediente" in df_casos_vis.columns:
            visibles = (
                df_casos_vis["Expediente"]
                .astype(str)
                .apply(normalize_key)
                .tolist()
            )

            actv = df_act[
                df_act["Caso"]
                .astype(str)
                .apply(normalize_key)
                .isin(visibles)
            ].copy()

            if actv.empty:
                st.info("No hay actuaciones para tus casos.")
            else:
                st.dataframe(
                    actv.sort_values("Fecha", ascending=False),
                    use_container_width=True
                )

                st.download_button(
                    "⬇️ Descargar actuaciones (CSV)",
                    actv.to_csv(index=False).encode("utf-8"),
                    "actuaciones_mis_casos.csv"
                )

    st.divider()

    # ------------------------------------------------------
    # REPORTE FINANCIERO (SOLO ADMIN / ADMINISTRATIVO)
    # ------------------------------------------------------
    if rol in ["admin", "personal administrativo"]:
        st.markdown("### 💰 Reporte financiero")

        # Usa TU función existente (no la inventamos)
        df_fin = resumen_financiero_df()

        if df_fin is None or df_fin.empty:
            st.info("No hay información financiera.")
        else:
            st.dataframe(df_fin, use_container_width=True)

            st.download_button(
                "⬇️ Descargar reporte financiero (CSV)",
                df_fin.to_csv(index=False).encode("utf-8"),
                "reporte_financiero.csv"
            )
    else:
        st.info("🔒 Tu rol no tiene acceso a reportes financieros.")

# ==========================================================
# AUDITORÍA (huérfanos)
# ==========================================================
if menu == "Auditoría":
    st.subheader("🧹 Auditoría")
    st.info("Muestra registros huérfanos (Caso no existe en casos). Útil para detectar montos/pagos fantasmas.")

    casos_set = set(casos["Expediente"].tolist())

    def orphans(df):
        if df.empty or "Caso" not in df.columns:
            return df
        tmp = df.copy()
        tmp["Caso"] = tmp["Caso"].apply(normalize_key)
        return tmp[~tmp["Caso"].isin(casos_set)].copy()

    st.markdown("### Honorarios total huérfanos")
    st.dataframe(orphans(honorarios), use_container_width=True)

    st.markdown("### Honorarios por etapa huérfanos")
    st.dataframe(orphans(honorarios_etapas), use_container_width=True)

    st.markdown("### Pagos honorarios huérfanos")
    st.dataframe(orphans(pagos_honorarios), use_container_width=True)

    st.markdown("### Cuota litis huérfana")
    st.dataframe(orphans(cuota_litis), use_container_width=True)

    st.markdown("### Pagos litis huérfanos")
    st.dataframe(orphans(pagos_litis), use_container_width=True)

    st.markdown("### Cuotas cronograma huérfanas")
    st.dataframe(orphans(cuotas), use_container_width=True)

    st.markdown("### Actuaciones huérfanas")
    st.dataframe(orphans(actuaciones.rename(columns={"Caso":"Caso"})), use_container_width=True)

    st.markdown("### Consultas huérfanas (si expediente no existe)")
    if not consultas.empty:
        tmp = consultas.copy()
        tmp["Caso"] = tmp["Caso"].apply(normalize_key)
        hu = tmp[(tmp["Caso"] != "") & (~tmp["Caso"].isin(casos_set))].copy()
        st.dataframe(hu, use_container_width=True)
    else:
        st.info("No hay consultas.")
# ==========================================================
# ======= PARCHE MARCA 004 (INSERTAR DESPUÉS DE LÍNEA 1575) ==
# Objetivo: NO recortar nada, solo mejorar lo indicado:
#  - Dashboard: cronograma visible aunque df_estado esté vacío
#  - Resumen económico: tabla consolidada visible
#  - Cuota Litis y Pagos Litis: CRUD completo visible
#  - Reporte por abogado y sus casos
# ==========================================================

# ----------------------------------------------------------
# 1) DASHBOARD: Cronograma visible aunque no haya pagos
#    (No toca tu dashboard, solo añade salida si corresponde)
# ----------------------------------------------------------
try:
    if 'menu' in globals() and menu == "Dashboard":
        # Mostrar cronograma base si hay cuotas aunque df_estado sea vacío
        st.divider()
        st.markdown("### 📅 Cronograma (visible aunque no haya pagos)")

        # Si existe cuotas.csv se debe mostrar siempre
        if 'cuotas' in globals() and not cuotas.empty:
            # Si df_estado está vacío o no tiene columnas esperadas, mostrar cronograma base
            if ('df_estado' not in globals()) or (df_estado is None) or (getattr(df_estado, "empty", True)) or ("SaldoCuota" not in getattr(df_estado, "columns", [])):
                st.info("Hay cuotas registradas, pero aún no hay pagos asignados. Mostrando cronograma base.")
                st.dataframe(
                    cuotas.sort_values(["Caso", "Tipo", "NroCuota"], ascending=[True, True, True]),
                    use_container_width=True
                )
            else:
                st.dataframe(
                    df_estado.sort_values(["Caso", "Tipo", "NroCuota"], ascending=[True, True, True]),
                    use_container_width=True
                )
        else:
            st.info("No hay cuotas registradas aún.")

        # Resumen económico consolidado (tabla)
        st.divider()
        st.markdown("### 📊 Resumen económico consolidado")
        if 'resumen_financiero_df' in globals():
            df_res_local = resumen_financiero_df()
            st.dataframe(df_res_local, use_container_width=True)
except Exception:
    pass

# ----------------------------------------------------------
# 2) CUOTA LITIS: asegurar CRUD visible (no borra tu menú)
# ----------------------------------------------------------
def _patch_render_cuota_litis_menu():
    st.subheader("⚖️ Cuota Litis (restaurado)")
    st.dataframe(cuota_litis.sort_values("ID", ascending=False), use_container_width=True)

    exp_list = casos["Expediente"].tolist() if not casos.empty else []
    if exp_list:
        exp = st.selectbox("Expediente", exp_list, key="patch_cl_exp")
        base = st.number_input("Monto base", min_value=0.0, step=100.0, key="patch_cl_base")
        porc = st.number_input("Porcentaje (%)", min_value=0.0, step=1.0, key="patch_cl_porc")
        notas = st.text_input("Notas", value="", key="patch_cl_notas")

        if st.button("Guardar cuota litis", key="patch_cl_save"):
            new_id = next_id(cuota_litis)
            nueva = {
                "ID": new_id,
                "Caso": normalize_key(exp),
                "Monto Base": float(base),
                "Porcentaje": float(porc),
                "Notas": notas,
                "FechaRegistro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            df2 = add_row(cuota_litis, nueva, "cuota_litis")
            save_df("cuota_litis", df2)
            st.success("✅ Cuota litis guardada")
            st.rerun()

    st.divider()
    st.markdown("### Editar / borrar (por ID)")
    if not cuota_litis.empty:
        sel = st.selectbox("Cuota Litis ID", cuota_litis["ID"].tolist(), key="patch_cl_sel")
        fila = cuota_litis[cuota_litis["ID"] == sel].iloc[0]

        with st.form("patch_cl_edit_form"):
            caso_e = st.text_input("Expediente", value=str(fila["Caso"]), key="patch_cl_caso_e")
            base_e = st.number_input("Monto base", min_value=0.0, value=money(fila["Monto Base"]), step=100.0, key="patch_cl_base_e")
            porc_e = st.number_input("Porcentaje", min_value=0.0, value=money(fila["Porcentaje"]), step=1.0, key="patch_cl_porc_e")
            notas_e = st.text_input("Notas", value=str(fila["Notas"]), key="patch_cl_notas_e")
            submit = st.form_submit_button("Guardar cambios")
            if submit:
                idx = cuota_litis.index[cuota_litis["ID"] == sel][0]
                cuota_litis.loc[idx, ["Caso","Monto Base","Porcentaje","Notas"]] = [normalize_key(caso_e), float(base_e), float(porc_e), notas_e]
                save_df("cuota_litis", cuota_litis)
                st.success("✅ Actualizado")
                st.rerun()

        if st.button("🗑️ Borrar cuota litis", key="patch_cl_del"):
            df2 = cuota_litis[cuota_litis["ID"] != sel].copy()
            save_df("cuota_litis", df2)
            st.success("✅ Eliminado")
            st.rerun()

    st.download_button("⬇️ Descargar cuota litis (CSV)", cuota_litis.to_csv(index=False).encode("utf-8"), "cuota_litis.csv")

# ----------------------------------------------------------
# 3) PAGOS LITIS: asegurar CRUD visible (no borra tu menú)
# ----------------------------------------------------------
def _patch_render_pagos_litis_menu():
    st.subheader("💳 Pagos Cuota Litis (restaurado)")
    st.dataframe(pagos_litis.sort_values("FechaPago", ascending=False), use_container_width=True)

    exp_list = casos["Expediente"].tolist() if not casos.empty else []
    if exp_list:
        exp = st.selectbox("Expediente", exp_list, key="patch_pl_exp")
        fecha = st.date_input("Fecha pago", value=date.today(), key="patch_pl_fecha")
        monto = st.number_input("Monto pagado", min_value=0.0, step=50.0, key="patch_pl_monto")
        obs = st.text_input("Observación", value="", key="patch_pl_obs")

        if st.button("Registrar pago litis", key="patch_pl_save"):
            new_id = next_id(pagos_litis)
            nueva = {
                "ID": new_id,
                "Caso": normalize_key(exp),
                "FechaPago": str(fecha),
                "Monto": float(monto),
                "Observacion": obs
            }
            df2 = add_row(pagos_litis, nueva, "pagos_litis")
            save_df("pagos_litis", df2)
            st.success("✅ Pago litis guardado")
            st.rerun()

    st.divider()
    st.markdown("### Editar / borrar (por ID)")
    if not pagos_litis.empty:
        sel = st.selectbox("Pago ID", pagos_litis["ID"].tolist(), key="patch_pl_sel")
        fila = pagos_litis[pagos_litis["ID"] == sel].iloc[0]

        with st.form("patch_pl_edit_form"):
            caso_e = st.text_input("Expediente", value=str(fila["Caso"]), key="patch_pl_caso_e")
            fecha_e = st.text_input("FechaPago (YYYY-MM-DD)", value=str(fila["FechaPago"]), key="patch_pl_fecha_e")
            monto_e = st.number_input("Monto", min_value=0.0, value=money(fila["Monto"]), step=50.0, key="patch_pl_monto_e")
            obs_e = st.text_input("Observación", value=str(fila["Observacion"]), key="patch_pl_obs_e")
            submit = st.form_submit_button("Guardar cambios")
            if submit:
                idx = pagos_litis.index[pagos_litis["ID"] == sel][0]
                pagos_litis.loc[idx, :] = [sel, normalize_key(caso_e), fecha_e, float(monto_e), obs_e]
                save_df("pagos_litis", pagos_litis)
                st.success("✅ Actualizado")
                st.rerun()

        if st.button("🗑️ Borrar pago litis", key="patch_pl_del"):
            df2 = pagos_litis[pagos_litis["ID"] != sel].copy()
            save_df("pagos_litis", df2)
            st.success("✅ Eliminado")
            st.rerun()

    st.download_button("⬇️ Descargar pagos litis (CSV)", pagos_litis.to_csv(index=False).encode("utf-8"), "pagos_litis.csv")

# ----------------------------------------------------------
# 4) Reporte por Abogado: bloque seguro (no destruye tabs existentes)
#    Se activa dentro de menu == "Reportes" si existe df_res y casos
# ----------------------------------------------------------
def _patch_reporte_por_abogado():
    st.markdown("### 👨‍⚖️ Reporte por abogado y sus casos (añadido)")
    if casos.empty:
        st.info("No hay casos registrados.")
        return
    df_res = resumen_financiero_df() if 'resumen_financiero_df' in globals() else pd.DataFrame()
    if df_res.empty:
        st.info("No hay datos financieros aún.")
        return

    # unir abogado
    dfm = df_res.merge(casos[["Expediente","Abogado"]], on="Expediente", how="left")

    rep_ab = dfm.groupby("Abogado", as_index=False).agg({
        "Honorario Pactado":"sum",
        "Honorario Pagado":"sum",
        "Honorario Pendiente":"sum",
        "Saldo Litis":"sum"
    })

    st.dataframe(rep_ab, use_container_width=True)
    st.download_button("⬇️ Descargar reporte por abogado (CSV)", rep_ab.to_csv(index=False).encode("utf-8"), "reporte_por_abogado.csv")

    st.divider()
    st.markdown("#### Detalle de casos por abogado")
    abogados_disp = [a for a in rep_ab["Abogado"].dropna().tolist() if str(a).strip() != ""]
    if not abogados_disp:
        st.info("No hay abogado asociado en los casos.")
        return

    ab_sel = st.selectbox("Selecciona abogado", abogados_disp, key="patch_ab_sel")
    det = dfm[dfm["Abogado"] == ab_sel].copy()
    st.dataframe(det, use_container_width=True)
    st.download_button("⬇️ Descargar detalle abogado (CSV)", det.to_csv(index=False).encode("utf-8"), f"reporte_detalle_{ab_sel}.csv")

# ----------------------------------------------------------
# Activación “sin tocar tu código”: si el menú coincide, muestra el bloque
# ----------------------------------------------------------
try:
    if 'menu' in globals() and menu == "Cuota Litis":
        _patch_render_cuota_litis_menu()
except Exception:
    pass

try:
    if 'menu' in globals() and menu == "Pagos Cuota Litis":
        _patch_render_pagos_litis_menu()
except Exception:
    pass

try:
    if 'menu' in globals() and menu == "Reportes":
        st.divider()
        _patch_reporte_por_abogado()
except Exception:
    pass

# =================== FIN PARCHE MARCA 004 ==================

# ==========================================================
# ======= MARCA 006 – EXTENSIONES (Roles, Auditoría, Búsqueda, Dashboard Proactivo)
# ==========================================================

def _audit_log(accion, entidad='', entidad_id='', detalle=''):
    try:
        df = load_df('auditoria_mod')
        new_id = next_id(df)
        df = add_row(df, {
            'ID': new_id,
            'Fecha': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'Usuario': st.session_state.get('usuario',''),
            'Rol': st.session_state.get('rol',''),
            'Accion': accion,
            'Entidad': entidad,
            'EntidadID': str(entidad_id),
            'Detalle': detalle
        }, 'auditoria_mod')
        save_df('auditoria_mod', df)
    except Exception:
        pass


def _can_edit():
    return st.session_state.get('rol') in ['admin','abogado']

def _is_readonly():
    return st.session_state.get('rol') == 'asistente'


def _filter_cases_by_role(df_casos):
    # Abogado: solo sus casos (match por nombre de abogado)
    if st.session_state.get('rol') == 'abogado':
        user = st.session_state.get('usuario','')
        # si existe abogados.csv y abogado_id, intentamos mapear a Nombre
        ab_name = None
        try:
            if 'abogados' in globals() and not abogados.empty:
                # Si el usuario tiene abogado_id, usarlo
                aid = st.session_state.get('abogado_id','')
                if str(aid).strip() and (abogados['ID'].astype(str) == str(aid)).any():
                    ab_name = abogados[abogados['ID'].astype(str)==str(aid)].iloc[0].get('Nombre','')
        except Exception:
            pass
        name = ab_name if ab_name else user
        return df_casos[df_casos['Abogado'].astype(str) == str(name)].copy()
    return df_casos

# Aplicar filtro visual a casos y tablas dependientes
try:
    if 'casos' in globals() and not casos.empty:
        _casos_f = _filter_cases_by_role(casos)
        if len(_casos_f) != len(casos):
            casos = _casos_f
except Exception:
    pass

# ==========================================================
# BÚSQUEDA GLOBAL (sidebar) – expediente, cliente, DNI, abogado, materia
# ==========================================================
try:
    with st.sidebar.expander('🔎 Búsqueda', expanded=False):
        q = st.text_input('Buscar (expediente/cliente/DNI/abogado/materia)')
        if q:
            ql = q.lower().strip()
            results = []
            if 'casos' in globals() and not casos.empty:
                for _, r in casos.iterrows():
                    if any(ql in str(r.get(k,'')).lower() for k in ['Expediente','Cliente','Abogado','Materia']):
                        results.append({'Tipo':'Caso','Clave':r.get('Expediente',''),'Detalle':f"{r.get('Cliente','')} | {r.get('Abogado','')} | {r.get('Materia','')}"})
            if 'clientes' in globals() and not clientes.empty:
                for _, r in clientes.iterrows():
                    if any(ql in str(r.get(k,'')).lower() for k in ['Nombre','DNI','RUC','RazonSocial']):
                        results.append({'Tipo':'Cliente','Clave':r.get('Nombre',''),'Detalle':str(r.get('DNI','')) or str(r.get('RUC',''))})
            if 'abogados' in globals() and not abogados.empty:
                for _, r in abogados.iterrows():
                    if any(ql in str(r.get(k,'')).lower() for k in ['Nombre','DNI','Colegiatura','ColegioProfesional']):
                        results.append({'Tipo':'Abogado','Clave':r.get('Nombre',''),'Detalle':str(r.get('Colegiatura',''))})
            if results:
                st.dataframe(pd.DataFrame(results), use_container_width=True)
            else:
                st.info('Sin resultados.')
except Exception:
    pass

# ==========================================================
# DASHBOARD – Actuaciones pendientes con semáforo
# ==========================================================
try:
    if 'menu' in globals() and menu == 'Dashboard':
        st.divider()
        st.markdown('### ⏱️ Actuaciones pendientes (semáforo)')
        if 'actuaciones' in globals() and not actuaciones.empty:
            def _dias(x):
                d = to_date_safe(x)
                return None if d is None else (d - date.today()).days
            tmp = actuaciones.copy()
            tmp['Dias'] = tmp['FechaProximaAccion'].apply(_dias)
            tmp = tmp[tmp['Dias'].notna()].copy()
            tmp = tmp[tmp['Dias'] >= 0].copy()
            def _sem(d):
                if d <= 2: return '🔴'
                if 3 <= d <= 5: return '🟡'
                return '🟢'
            tmp['Sem'] = tmp['Dias'].apply(_sem)
            tmp.sort_values('Dias', inplace=True)
            st.dataframe(tmp[['Sem','Caso','TipoActuacion','ProximaAccion','FechaProximaAccion','Dias']], use_container_width=True)
        else:
            st.info('No hay actuaciones con próxima acción registrada.')
except Exception:
    pass


# ==========================================================
# CONSULTAS – abogado a cargo + costo + reporte ingresos
# ==========================================================
try:
    if 'menu' in globals() and menu == 'Consultas':
        st.divider()
        st.markdown('## 💼 Consultas – Ingresos')
        if 'consultas' in globals() and not consultas.empty:
            dfc = consultas.copy()
            dfc['CostoConsulta'] = pd.to_numeric(dfc.get('CostoConsulta', 0), errors='coerce').fillna(0.0)
            rep = dfc.groupby('Abogado', as_index=False)['CostoConsulta'].sum().rename(columns={'CostoConsulta':'IngresosConsultas'})
            st.dataframe(rep, use_container_width=True)
            st.download_button('⬇️ Descargar ingresos por consulta (CSV)', rep.to_csv(index=False).encode('utf-8'), 'ingresos_consultas.csv')
        else:
            st.info('No hay consultas registradas aún.')
except Exception:
    pass



# ==========================================================
# INSTANCIAS (REGISTRAR / EDITAR / BORRAR)
# + Instancia Administrativa con Cosa Decidida
# ==========================================================
if menu == "Instancias":
    st.subheader("⚖️ Instancias del Caso")

    if casos.empty:
        st.info("Primero registra casos.")
    else:
        exp = st.selectbox("Expediente", casos["Expediente"].tolist(), key="inst_exp")
        exp_n = normalize_key(exp)

        # ✅ Cargar desde CSV (como en tu arquitectura real)
        df_i = load_df("instancias")
        if not df_i.empty and "Caso" in df_i.columns:
            df_i["Caso"] = df_i["Caso"].apply(normalize_key)

        sub = df_i[df_i["Caso"] == exp_n].copy() if not df_i.empty else pd.DataFrame()

        # =========================
        # LISTADO
        # =========================
        st.markdown("### Instancias registradas")
        if sub.empty:
            st.info("No hay instancias registradas para este expediente.")
        else:
            st.dataframe(
                sub.sort_values("ID", ascending=False),
                use_container_width=True
            )

        st.divider()

        # =========================
        # REGISTRAR NUEVA
        # =========================
        st.markdown("### ➕ Registrar instancia")

        with st.form("inst_new_form"):
            tipo = st.selectbox(
                "Tipo de instancia",
                [
                    "Instancia Administrativa",
                    "Primera Instancia",
                    "Segunda Instancia",
                    "Casación",
                    "Otros"
                ]
            )

            estado = st.text_input("Estado actual")
            resultado = st.text_input("Resultado")
            accion = st.text_input("Acción / Próximo paso")
            honorarios = st.number_input("Honorarios (S/)", min_value=0.0, step=100.0)

            # ---- SOLO ADMINISTRATIVA ----
            sede_admin = ""
            cosa_decidida = "0"
            fecha_cosa = ""

            if tipo == "Instancia Administrativa":
                st.markdown("**Datos de instancia administrativa**")
                sede_admin = st.text_input("Sede Administrativa")
                cosa = st.checkbox("¿Adquirió calidad de Cosa Decidida?")
                cosa_decidida = "1" if cosa else "0"
                if cosa:
                    fecha_cosa = st.text_input("Fecha de cosa decidida (YYYY-MM-DD)")

            submit = st.form_submit_button("Guardar instancia")

            if submit:
                new_id = next_id(df_i)
                df_i = add_row(df_i, {
                    "ID": new_id,
                    "Caso": exp_n,
                    "TipoInstancia": tipo,
                    "EstadoActual": estado,
                    "Resultado": resultado,
                    "Accion": accion,
                    "Honorarios": float(honorarios),
                    "SedeAdministrativa": sede_admin,
                    "CosaDecidida": cosa_decidida,
                    "FechaCosaDecidida": fecha_cosa
                }, "instancias")
                save_df("instancias", df_i)
                st.success("✅ Instancia registrada")
                st.rerun()

        st.divider()

        # =========================
        # EDITAR / BORRAR
        # =========================
        st.markdown("### ✏️ Editar / borrar instancia")

        if not sub.empty:
            sel_id = st.selectbox("Instancia ID", sub["ID"].tolist(), key="inst_edit_id")
            fila = df_i[df_i["ID"] == sel_id].iloc[0]

            with st.form("inst_edit_form"):
                tipo_e = st.selectbox(
                    "Tipo de instancia",
                    [
                        "Instancia Administrativa",
                        "Primera Instancia",
                        "Segunda Instancia",
                        "Casación",
                        "Otros"
                    ],
                    index=[
                        "Instancia Administrativa",
                        "Primera Instancia",
                        "Segunda Instancia",
                        "Casación",
                        "Otros"
                    ].index(fila.get("TipoInstancia", "Otros"))
                )

                estado_e = st.text_input("Estado actual", value=str(fila.get("EstadoActual", "")))
                resultado_e = st.text_input("Resultado", value=str(fila.get("Resultado", "")))
                accion_e = st.text_input("Acción", value=str(fila.get("Accion", "")))
                honor_e = st.number_input(
                    "Honorarios (S/)",
                    min_value=0.0,
                    value=float(pd.to_numeric(fila.get("Honorarios", 0), errors="coerce") or 0.0),
                    step=100.0
                )

                sede_admin_e = fila.get("SedeAdministrativa", "")
                cosa_decidida_e = fila.get("CosaDecidida", "0") == "1"
                fecha_cosa_e = fila.get("FechaCosaDecidida", "")

                if tipo_e == "Instancia Administrativa":
                    st.markdown("**Datos de instancia administrativa**")
                    sede_admin_e = st.text_input("Sede Administrativa", value=str(sede_admin_e))
                    cosa_decidida_e = st.checkbox("¿Cosa Decidida?", value=cosa_decidida_e)
                    if cosa_decidida_e:
                        fecha_cosa_e = st.text_input(
                            "Fecha de cosa decidida (YYYY-MM-DD)",
                            value=str(fecha_cosa_e)
                        )
                    else:
                        fecha_cosa_e = ""

                submit_edit = st.form_submit_button("Guardar cambios")

                if submit_edit:
                    idx = df_i.index[df_i["ID"] == sel_id][0]
                    df_i.loc[idx, [
                        "TipoInstancia", "EstadoActual", "Resultado", "Accion",
                        "Honorarios", "SedeAdministrativa",
                        "CosaDecidida", "FechaCosaDecidida"
                    ]] = [
                        tipo_e, estado_e, resultado_e, accion_e,
                        float(honor_e), sede_admin_e,
                        "1" if cosa_decidida_e else "0",
                        fecha_cosa_e
                    ]
                    save_df("instancias", df_i)
                    st.success("✅ Instancia actualizada")
                    st.rerun()

            st.warning("⚠️ Eliminar instancia (irreversible)")
            confirm = st.text_input("Escribe ELIMINAR para confirmar", key="inst_del_confirm")
            if st.button("🗑️ Eliminar instancia", disabled=confirm.strip().upper() != "ELIMINAR"):
                df_i = df_i[df_i["ID"] != sel_id].copy()
                save_df("instancias", df_i)
                st.success("✅ Instancia eliminada")
                st.rerun()
# ==========================================================
# MARCA 006 – CLIENTES (Extendido): Natural/Jurídico + Emergencia
# ==========================================================
try:
    if 'menu' in globals() and menu == 'Clientes (Extendido)':
        st.subheader('👥 Clientes (Extendido)')
        if clientes.empty:
            st.info('No hay clientes registrados.')
        else:
            sel = st.selectbox('Cliente ID', clientes['ID'].tolist())
            fila = clientes[clientes['ID'] == sel].iloc[0]
            tipo = st.selectbox('Tipo de cliente', ['Natural','Jurídica'], index=0 if str(fila.get('TipoCliente','Natural'))!='Jurídica' else 1)
            contacto = st.text_input('Contacto de emergencia', value=str(fila.get('ContactoEmergencia','')))
            cel_cont = st.text_input('Celular de contacto (emergencia)', value=str(fila.get('CelularEmergencia','')))
            # Jurídica
            rs = st.text_input('Razón Social (si jurídica)', value=str(fila.get('RazonSocial','')))
            ruc = st.text_input('RUC (si jurídica)', value=str(fila.get('RUC','')))
            rep = st.text_input('Representante legal', value=str(fila.get('RepresentanteLegal','')))
            partida = st.text_input('Partida electrónica', value=str(fila.get('PartidaElectronica','')))
            sede = st.text_input('Sede registral', value=str(fila.get('SedeRegistral','')))
            if st.button('💾 Guardar extendido', disabled=_is_readonly()):
                idx = clientes.index[clientes['ID'] == sel][0]
                clientes.loc[idx, ['TipoCliente','ContactoEmergencia','CelularEmergencia','RazonSocial','RUC','RepresentanteLegal','PartidaElectronica','SedeRegistral']] = [
                    tipo, contacto, cel_cont, rs, ruc, rep, partida, sede
                ]
                save_df('clientes', clientes)
                _audit_log('UPDATE','clientes',sel,'extendido')
                st.success('✅ Guardado')
                st.rerun()
except Exception:
    pass
# ==========================================================
# PARCHE SQLITE – SINCRONIZACIÓN CSV ↔ SQLITE (NO OBLIGATORIO)
# ==========================================================
import sqlite3

SQLITE_DB = "data.db"

def csv_to_sqlite():
    conn = sqlite3.connect(SQLITE_DB)
    for key, path in FILES.items():
        try:
            df = load_df(key)
            df.to_sql(key, conn, if_exists="replace", index=False)
        except Exception as e:
            st.warning(f"No se pudo exportar {key}: {e}")
    conn.close()

def sqlite_to_csv():
    conn = sqlite3.connect(SQLITE_DB)
    cur = conn.cursor()
    tables = cur.execute(
        "SELECT name FROM sqlite_master WHERE type='table'"
    ).fetchall()

    for (table,) in tables:
        try:
            df = pd.read_sql(f"SELECT * FROM {table}", conn)
            if table in FILES:
                df.to_csv(FILES[table], index=False)
        except Exception as e:
            st.warning(f"No se pudo importar {table}: {e}")
    conn.close()

# UI mínima (oculta, segura)
with st.sidebar.expander("🗄️ Base de Datos (SQLite)", expanded=False):
    st.caption("CSV sigue siendo la fuente principal")
    if st.button("⬆️ Exportar CSV → SQLite"):
        csv_to_sqlite()
        st.success("✅ CSV exportado a SQLite (data.db)")

    if st.button("⬇️ Importar SQLite → CSV"):
        sqlite_to_csv()
        st.success("✅ SQLite importado a CSV")
# ==========================================================
# PARCHE WORD – DESCARGAR CONTRATOS EN .DOCX
# ==========================================================
from docx import Document
from io import BytesIO

def generar_docx(texto: str, titulo="Contrato"):
    doc = Document()
    doc.add_heading(titulo, level=1)

    for linea in texto.split("\n"):
        doc.add_paragraph(linea)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# ==========================================================
# REPOSITORIO CONTRATOS – Borrador/Firmado + Numeración oficial + Vigente/Histórico + Versionado/Historial
# ==========================================================

REPO_CONTRATOS_FILE = os.path.join(DATA_DIR, 'repo_contratos.csv')
REPO_CONTRATOS_SCHEMA = [
    'ID','Archivo','Ruta','Extension','Expediente','Cliente','Abogado','NombreContrato',
    'Estado','Numero','Año','Sigla','FechaCreado','FechaFirmado','Notas','Hash','Existe','Visible',
    'Version','Historial','CategoriaContrato'
]

# -------------------------
# Helpers robustos
# -------------------------
def _repo_fix_ids(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    if 'ID' not in df.columns:
        df['ID'] = ''
    ids = pd.to_numeric(df['ID'], errors='coerce')
    max_id = int(ids.dropna().max()) if ids.notna().any() else 0
    bad = ids.isna() | (ids <= 0)
    for i in df.index[bad].tolist():
        max_id += 1
        df.at[i, 'ID'] = max_id
    df['ID'] = pd.to_numeric(df['ID'], errors='coerce').fillna(0).astype(int)
    return df

def _repo_get_idx(repo: pd.DataFrame, sel_id):
    if repo is None or repo.empty or 'ID' not in repo.columns:
        return None
    idxs = repo.index[repo['ID'].astype(str) == str(sel_id)].tolist()
    return idxs[0] if idxs else None

def _repo_append_hist(old: str, msg: str) -> str:
    base = (old or "").strip()
    return (base + "\n" + msg) if base else msg

def _repo_norm_cat(x: str) -> str:
    x = str(x or "").strip()
    return x if x in ["Vigente","Histórico"] else "Histórico"

# -------------------------
# IO
# -------------------------
def _repo_contratos_ensure():
    if not os.path.exists(REPO_CONTRATOS_FILE):
        pd.DataFrame(columns=REPO_CONTRATOS_SCHEMA).to_csv(REPO_CONTRATOS_FILE, index=False)
        return

    try:
        df = pd.read_csv(REPO_CONTRATOS_FILE)
    except Exception:
        df = pd.DataFrame(columns=REPO_CONTRATOS_SCHEMA)

    df = drop_unnamed(df)

    for c in REPO_CONTRATOS_SCHEMA:
        if c not in df.columns:
            df[c] = ''

    if 'Visible' not in df.columns:
        df['Visible'] = '1'
    if 'Version' not in df.columns:
        df['Version'] = 1
    if 'Historial' not in df.columns:
        df['Historial'] = ''
    if 'CategoriaContrato' not in df.columns:
        df['CategoriaContrato'] = 'Histórico'

    df = df.reindex(columns=REPO_CONTRATOS_SCHEMA)
    df = _repo_fix_ids(df)
    df['Version'] = pd.to_numeric(df.get('Version', 1), errors='coerce').fillna(1).astype(int)
    df['CategoriaContrato'] = df.get('CategoriaContrato','Histórico').apply(_repo_norm_cat)

    df.to_csv(REPO_CONTRATOS_FILE, index=False)

def _repo_contratos_load():
    _repo_contratos_ensure()
    try:
        df = pd.read_csv(REPO_CONTRATOS_FILE)
    except Exception:
        df = pd.DataFrame(columns=REPO_CONTRATOS_SCHEMA)

    df = drop_unnamed(df)

    for c in REPO_CONTRATOS_SCHEMA:
        if c not in df.columns:
            df[c] = ''

    if 'Visible' not in df.columns:
        df['Visible'] = '1'
    if 'Version' not in df.columns:
        df['Version'] = 1
    if 'Historial' not in df.columns:
        df['Historial'] = ''
    if 'CategoriaContrato' not in df.columns:
        df['CategoriaContrato'] = 'Histórico'

    df = df.reindex(columns=REPO_CONTRATOS_SCHEMA)
    df = _repo_fix_ids(df)
    df['Version'] = pd.to_numeric(df.get('Version', 1), errors='coerce').fillna(1).astype(int)
    df['CategoriaContrato'] = df.get('CategoriaContrato','Histórico').apply(_repo_norm_cat)

    return df

def _repo_contratos_save(df):
    try:
        backup_file(REPO_CONTRATOS_FILE)
    except Exception:
        pass

    df = drop_unnamed(df)
    for c in REPO_CONTRATOS_SCHEMA:
        if c not in df.columns:
            df[c] = ''

    if 'Visible' not in df.columns:
        df['Visible'] = '1'
    if 'Version' not in df.columns:
        df['Version'] = 1
    if 'Historial' not in df.columns:
        df['Historial'] = ''
    if 'CategoriaContrato' not in df.columns:
        df['CategoriaContrato'] = 'Histórico'

    df = df.reindex(columns=REPO_CONTRATOS_SCHEMA)
    df = _repo_fix_ids(df)
    df['Version'] = pd.to_numeric(df.get('Version', 1), errors='coerce').fillna(1).astype(int)
    df['CategoriaContrato'] = df.get('CategoriaContrato','Histórico').apply(_repo_norm_cat)

    df.to_csv(REPO_CONTRATOS_FILE, index=False)

# -------------------------
# Scan / Sync
# -------------------------
def _repo_file_sha(path: str) -> str:
    h = hashlib.sha256()
    try:
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(1024*1024), b''):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return ''

def _repo_scan_generados():
    os.makedirs(GENERADOS_DIR, exist_ok=True)
    out = []
    for fn in os.listdir(GENERADOS_DIR):
        if not fn.lower().endswith(('.txt','.docx')):
            continue
        path = os.path.join(GENERADOS_DIR, fn)
        if not os.path.isfile(path):
            continue

        ext = os.path.splitext(fn)[1].lower().replace('.','').upper()
        nombre, expediente = fn, ''

        if '_BORRADOR_' in fn:
            nombre = fn.split('_BORRADOR_')[0].replace('_',' ')
            expediente = fn.split('_BORRADOR_')[-1].split('.')[0].replace('__','_')

        cliente, abogado = '', ''
        try:
            if 'casos' in globals() and not casos.empty and expediente:
                r = casos[casos['Expediente'].astype(str)==str(expediente)].iloc[0]
                cliente = str(r.get('Cliente',''))
                abogado = str(r.get('Abogado',''))
        except Exception:
            pass

        try:
            fc = datetime.fromtimestamp(os.path.getmtime(path)).strftime('%Y-%m-%d %H:%M:%S')
        except Exception:
            fc = ''

        out.append({
            'Archivo': fn,'Ruta': path,'Extension': ext,
            'Expediente': expediente,'Cliente': cliente,'Abogado': abogado,
            'NombreContrato': nombre,'FechaCreado': fc,
            'Hash': _repo_file_sha(path),'Existe': '1'
        })
    return out

def _repo_sync_contratos():
    repo = _repo_contratos_load()
    scanned = _repo_scan_generados()
    scanned_by = {r['Archivo']: r for r in scanned}

    # Actualiza existentes y versiona si cambió el hash
    if not repo.empty:
        repo['Existe'] = '0'
        for i in repo.index:
            fn = str(repo.at[i,'Archivo'])
            if fn in scanned_by:
                repo.at[i,'Existe'] = '1'

                old_hash = str(repo.at[i,'Hash'])
                new_hash = str(scanned_by[fn].get('Hash',''))

                if old_hash and new_hash and old_hash != new_hash:
                    old_ver = pd.to_numeric(repo.at[i,'Version'], errors='coerce')
                    old_ver = 1 if pd.isna(old_ver) else int(old_ver)
                    new_ver = old_ver + 1
                    repo.at[i,'Version'] = new_ver

                    stamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    msg = f"[{stamp}] Versión {new_ver} (cambio de archivo)"
                    repo.at[i,'Historial'] = _repo_append_hist(str(repo.at[i,'Historial']), msg)

                for k in ['Ruta','Extension','Hash','FechaCreado','Cliente','Abogado','Expediente','NombreContrato']:
                    repo.at[i,k] = scanned_by[fn].get(k, repo.at[i,k])

    existing = set(repo['Archivo'].astype(str).tolist()) if not repo.empty else set()

    # ✅ FIX ROBUSTO: evita ValueError si max es NaN
    ids = pd.to_numeric(repo.get('ID', pd.Series(dtype='float')), errors='coerce')
    max_id = ids.dropna().max()
    max_id = 0 if pd.isna(max_id) else int(max_id)
    nid = max_id + 1

    # Nuevos siempre entran como HISTÓRICO (Opción C)
    new_rows = []
    for r in scanned:
        if r['Archivo'] not in existing:
            stamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            new_rows.append({
                'ID': nid,
                'Archivo': r['Archivo'],
                'Ruta': r['Ruta'],
                'Extension': r['Extension'],
                'Expediente': r.get('Expediente',''),
                'Cliente': r.get('Cliente',''),
                'Abogado': r.get('Abogado',''),
                'NombreContrato': r.get('NombreContrato','Contrato'),
                'Estado': 'Borrador',
                'Numero': '',
                'Año': '',
                'Sigla': 'CLS',
                'FechaCreado': r.get('FechaCreado',''),
                'FechaFirmado': '',
                'Notas': '',
                'Hash': r.get('Hash',''),
                'Existe': '1',
                'Visible': '1',
                'Version': 1,
                'Historial': f"[{stamp}] Creado (sincronizado desde generados/)",
                'CategoriaContrato': 'Histórico'
            })
            nid += 1

    if new_rows:
        repo = pd.concat([repo, pd.DataFrame(new_rows)], ignore_index=True)

    for c in REPO_CONTRATOS_SCHEMA:
        if c not in repo.columns:
            repo[c] = ''

    repo = repo.reindex(columns=REPO_CONTRATOS_SCHEMA)
    _repo_contratos_save(repo)
    return repo

# =========================
# UI del repositorio (CON BOTONES + RESET TOTAL + DETALLE POR CASO + BORRAR GENERADOS)
# =========================
if 'menu' in globals() and menu == 'Repositorio Contratos':
    st.subheader('📦 Repositorio de Contratos')

    is_admin = st.session_state.get('rol') == 'admin'

    # ===== Barra superior =====
    col1, col2, col3 = st.columns([1, 1, 1])

    with col1:
        if st.button('🔄 Sincronizar desde generados/', key='rc_sync'):
            _repo_sync_contratos()
            st.success('✅ Sincronizado')
            st.rerun()

    with col2:
        st.download_button(
            '⬇️ Descargar repositorio (CSV)',
            _repo_contratos_load().to_csv(index=False).encode('utf-8'),
            'repo_contratos.csv',
            key='rc_export'
        )

    with col3:
        if is_admin:
            with st.expander("⚠️ Reset total del repositorio (solo ADMIN)", expanded=False):
                st.warning("Esto borra: 1) repo_contratos.csv y 2) TODOS los archivos dentro de generados/.")
                confirm = st.checkbox("Entiendo el riesgo y deseo borrar todo", key="rc_reset_confirm")
                confirm2 = st.text_input("Escribe BORRAR para confirmar", value="", key="rc_reset_confirm2")
                if st.button("🧨 BORRAR TODO DEFINITIVAMENTE", key="rc_reset_btn",
                             disabled=(not confirm) or (confirm2.strip().upper() != "BORRAR")):
                    try:
                        if os.path.exists(REPO_CONTRATOS_FILE):
                            os.remove(REPO_CONTRATOS_FILE)
                        if os.path.exists(GENERADOS_DIR):
                            for fn in os.listdir(GENERADOS_DIR):
                                try:
                                    os.remove(os.path.join(GENERADOS_DIR, fn))
                                except Exception:
                                    pass
                        st.success("✅ Repositorio y generados/ eliminados por completo")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al resetear: {e}")
        else:
            st.write("")

    repo = _repo_contratos_load()
    if repo.empty:
        st.info('No hay contratos aún. Guarda un borrador y sincroniza.')
    else:
        # ===== Filtros + categoría =====
        f1, f2, f3, f4 = st.columns([1, 1, 1, 2])
        with f1:
            est = st.selectbox('Estado', ['Todos', 'Borrador', 'Firmado'], key='rc_estado')
        with f2:
            vis = st.selectbox('Visibilidad', ['Incluye ocultos', 'Solo visibles', 'Solo ocultos'], index=0, key='rc_vis')
        with f3:
            cat = st.selectbox('Categoría', ['Solo vigentes','Solo históricos','Todos'], index=0, key='rc_cat')
        with f4:
            q = st.text_input('Buscar', key='rc_buscar').strip().lower()

        view = repo.copy()

        if vis == 'Solo visibles':
            view = view[view['Visible'].astype(str) == '1']
        elif vis == 'Solo ocultos':
            view = view[view['Visible'].astype(str) != '1']

        if est != 'Todos':
            view = view[view['Estado'].astype(str) == est]

        if cat == 'Solo vigentes':
            view = view[view['CategoriaContrato'].astype(str) == 'Vigente']
        elif cat == 'Solo históricos':
            view = view[view['CategoriaContrato'].astype(str) != 'Vigente']

        if q:
            mask = (
                view['Archivo'].astype(str).str.lower().str.contains(q, na=False) |
                view['Expediente'].astype(str).str.lower().str.contains(q, na=False) |
                view['Cliente'].astype(str).str.lower().str.contains(q, na=False) |
                view['Abogado'].astype(str).str.lower().str.contains(q, na=False) |
                view['NombreContrato'].astype(str).str.lower().str.contains(q, na=False)
            )
            view = view[mask]

        st.dataframe(
            view[['ID','Archivo','Expediente','Cliente','Abogado','NombreContrato','Estado','Numero','Año','Sigla','CategoriaContrato','Version','Existe','Visible']],
            use_container_width=True
        )

        # ===== Detalle por expediente + BOTONES BORRAR GENERADOS =====
        st.divider()
        st.markdown("## 🔎 Detalle / Visualización por caso (expediente)")
        exp_opts = sorted([e for e in repo['Expediente'].astype(str).unique().tolist() if str(e).strip() != ""])
        exp_sel = None

        if exp_opts:
            exp_sel = st.selectbox("Expediente", exp_opts, key="rc_exp_det")
            det = repo[repo['Expediente'].astype(str) == str(exp_sel)].copy()

            # ✅ FIX: sort_values no admite errors='ignore'
            if 'CategoriaContrato' in det.columns:
                det['CategoriaContrato'] = det['CategoriaContrato'].astype(str)
            if 'Estado' in det.columns:
                det['Estado'] = det['Estado'].astype(str)
            if 'Version' in det.columns:
                det['Version'] = pd.to_numeric(det['Version'], errors='coerce').fillna(1).astype(int)

            sort_cols = [c for c in ['CategoriaContrato','Estado','Version'] if c in det.columns]
            if sort_cols:
                asc = [True, True, False][:len(sort_cols)]
                det.sort_values(sort_cols, inplace=True, ascending=asc)

            st.dataframe(det[['ID','Archivo','Estado','CategoriaContrato','Version','FechaCreado','FechaFirmado','Visible']], use_container_width=True)

            # ✅ BOTONES ADMIN: BORRAR GENERADOS (POR CASO / TODO)
            if is_admin:
                with st.expander("🧹 ADMIN: Borrar contratos generados (archivos en generados/)", expanded=False):
                    st.warning("Estas acciones borran ARCHIVOS físicos en generados/. No modifican honorarios ni datos del caso.")

                    cA, cB = st.columns(2)

                    with cA:
                        st.markdown("### 🧾 Borrar SOLO por expediente")
                        only_hist = st.checkbox("Borrar SOLO históricos (recomendado)", value=True, key="rc_delgen_only_hist")
                        conf1 = st.text_input("Escribe BORRAR-CASO para confirmar", key="rc_delgen_conf_case")

                        if st.button("🧨 Borrar generados del expediente", key="rc_delgen_case_btn",
                                     disabled=(conf1.strip().upper() != "BORRAR-CASO")):
                            repo2 = _repo_contratos_load()
                            mask_exp = repo2['Expediente'].astype(str) == str(exp_sel)
                            if only_hist:
                                mask_exp = mask_exp & (repo2.get('CategoriaContrato','Histórico').astype(str) != 'Vigente')

                            to_delete = repo2[mask_exp].copy()
                            borrados = 0

                            for ridx in to_delete.index:
                                ruta_del = str(repo2.at[ridx, 'Ruta'])
                                if ruta_del and os.path.exists(ruta_del):
                                    try:
                                        os.remove(ruta_del)
                                        borrados += 1
                                    except Exception:
                                        pass

                                    if ruta_del.lower().endswith(".txt"):
                                        docx_path = ruta_del[:-4] + ".docx"
                                        if os.path.exists(docx_path):
                                            try:
                                                os.remove(docx_path)
                                            except Exception:
                                                pass

                                stamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                                repo2.at[ridx, 'Existe'] = '0'
                                repo2.at[ridx, 'Historial'] = _repo_append_hist(
                                    str(repo2.at[ridx,'Historial']),
                                    f"[{stamp}] Archivos físicos borrados en generados/ (por expediente)"
                                )

                            _repo_contratos_save(repo2)
                            st.success(f"✅ Borrados {borrados} archivos del expediente {exp_sel}")
                            st.rerun()

                    with cB:
                        st.markdown("### 🧨 Borrar TODO generados/")
                        conf2 = st.text_input("Escribe BORRAR-TODO para confirmar", key="rc_delgen_conf_all")

                        if st.button("🧨 Borrar TODO generados/", key="rc_delgen_all_btn",
                                     disabled=(conf2.strip().upper() != "BORRAR-TODO")):
                            borrados = 0
                            try:
                                if os.path.exists(GENERADOS_DIR):
                                    for fn in os.listdir(GENERADOS_DIR):
                                        try:
                                            os.remove(os.path.join(GENERADOS_DIR, fn))
                                            borrados += 1
                                        except Exception:
                                            pass

                                repo2 = _repo_contratos_load()
                                if not repo2.empty:
                                    repo2['Existe'] = '0'
                                    stamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                                    for i in repo2.index:
                                        repo2.at[i,'Historial'] = _repo_append_hist(
                                            str(repo2.at[i,'Historial']),
                                            f"[{stamp}] Limpieza total generados/ (archivos físicos borrados)"
                                        )
                                    _repo_contratos_save(repo2)

                                st.success(f"✅ Borrados {borrados} archivos de generados/")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error al borrar generados/: {e}")
        else:
            st.info("Aún no hay expedientes detectados en el repositorio.")

        st.divider()

        if view.empty:
            st.info('No hay resultados con esos filtros.')
        else:
            view2 = view.copy()
            view2['ID'] = pd.to_numeric(view2['ID'], errors='coerce').fillna(0).astype(int)
            view2['_label'] = view2.apply(
                lambda r: f"ID {r['ID']} – {r.get('NombreContrato','')} – {r.get('Expediente','')} – {r.get('Estado','')} – {r.get('CategoriaContrato','Histórico')} – v{r.get('Version',1)}",
                axis=1
            )

            sel_label = st.selectbox('Selecciona contrato', view2['_label'].tolist(), key='rc_sel_id')
            sel_id = str(view2[view2['_label'] == sel_label].iloc[0]['ID'])

            row = repo[repo['ID'].astype(str) == str(sel_id)].iloc[0]

            st.markdown(f"**Versión:** {int(pd.to_numeric(row.get('Version',1), errors='coerce') or 1)}")
            with st.expander("📜 Historial", expanded=False):
                st.text_area("Historial", value=str(row.get('Historial','')), height=160, key='rc_hist')

            notas = st.text_area('Notas / Observaciones', value=str(row.get('Notas','')), height=90, key='rc_notas')
            sigla = st.text_input('Sigla', value=str(row.get('Sigla','CLS')), key='rc_sigla')

            # ----- Editar metadatos (solo ADMIN) -----
            with st.expander("✏️ Editar metadatos (solo ADMIN)", expanded=False):
                if not is_admin:
                    st.info("Solo ADMIN puede editar metadatos.")
                else:
                    nombre_e = st.text_input("NombreContrato", value=str(row.get('NombreContrato','')), key='rc_edit_nombre')
                    expediente_e = st.text_input("Expediente", value=str(row.get('Expediente','')), key='rc_edit_exp')
                    cliente_e = st.text_input("Cliente", value=str(row.get('Cliente','')), key='rc_edit_cli')
                    abogado_e = st.text_input("Abogado", value=str(row.get('Abogado','')), key='rc_edit_ab')
                    cat_e = st.selectbox("Categoría", ["Vigente","Histórico"], index=0 if str(row.get('CategoriaContrato'))=="Vigente" else 1, key='rc_edit_cat')
                    estado_e = st.selectbox("Estado", ["Borrador","Firmado"], index=0 if str(row.get('Estado'))!="Firmado" else 1, key='rc_edit_estado')

                    if st.button("💾 Guardar metadatos", key="rc_save_meta"):
                        repo2 = _repo_contratos_load()
                        i2 = _repo_get_idx(repo2, sel_id)
                        if i2 is not None:
                            repo2.at[i2,'NombreContrato'] = nombre_e
                            repo2.at[i2,'Expediente'] = expediente_e
                            repo2.at[i2,'Cliente'] = cliente_e
                            repo2.at[i2,'Abogado'] = abogado_e
                            repo2.at[i2,'CategoriaContrato'] = cat_e
                            repo2.at[i2,'Estado'] = estado_e
                            repo2.at[i2,'Sigla'] = sigla
                            repo2.at[i2,'Notas'] = notas
                            _repo_contratos_save(repo2)
                            st.success("✅ Metadatos guardados")
                            st.rerun()

            # Guardar notas rápido (solo admin)
            idx = _repo_get_idx(repo, sel_id)
            if idx is not None and is_admin:
                repo.at[idx, 'Notas'] = notas
                repo.at[idx, 'Sigla'] = sigla
                _repo_contratos_save(repo)

            # ===== Botones (solo ADMIN) =====
            b1, b2, b3, b4 = st.columns(4)

            with b1:
                if st.button('✅ Firmar (asigna N° y vuelve Vigente)', key='rc_firmar_btn', disabled=not is_admin):
                    repo2 = _repo_contratos_load()
                    idx2 = _repo_get_idx(repo2, sel_id)
                    if idx2 is not None:
                        repo2.at[idx2, 'Notas'] = notas
                        repo2.at[idx2, 'Sigla'] = sigla
                    repo2 = _repo_firmar(repo2, sel_id, sigla)

                    # Opción C: al firmar -> este contrato Vigente, otros del mismo expediente Histórico
                    try:
                        idx2 = _repo_get_idx(repo2, sel_id)
                        exp_of = str(repo2.at[idx2,'Expediente'])
                        same = repo2['Expediente'].astype(str) == exp_of
                        repo2.loc[same, 'CategoriaContrato'] = 'Histórico'
                        repo2.at[idx2, 'CategoriaContrato'] = 'Vigente'
                        stamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        repo2.at[idx2, 'Historial'] = _repo_append_hist(str(repo2.at[idx2,'Historial']), f"[{stamp}] Marcado como VIGENTE (firmado)")
                        _repo_contratos_save(repo2)
                    except Exception:
                        pass

                    st.success('✅ Firmado, numerado y marcado como Vigente')
                    st.rerun()

            with b2:
                if st.button('📝 Marcar BORRADOR', key='rc_borr_btn', disabled=not is_admin):
                    repo2 = _repo_contratos_load()
                    idx2 = _repo_get_idx(repo2, sel_id)
                    if idx2 is not None:
                        repo2.at[idx2, 'Notas'] = notas
                        repo2.at[idx2, 'Sigla'] = sigla
                    repo2 = _repo_borrador(repo2, sel_id)
                    st.success('✅ Marcado como borrador')
                    st.rerun()

            with b3:
                if str(row.get('Visible','1')) == '1':
                    if st.button('🗑️ Quitar', key='rc_quitar_btn', disabled=not is_admin):
                        repo2 = _repo_contratos_load()
                        repo2 = _repo_quitar(repo2, sel_id)
                        st.success('✅ Quitado (no revive)')
                        st.rerun()
                else:
                    if st.button('♻️ Restaurar', key='rc_rest_btn', disabled=not is_admin):
                        repo2 = _repo_contratos_load()
                        repo2 = _repo_restaurar(repo2, sel_id)
                        st.success('✅ Restaurado')
                        st.rerun()

            with b4:
                if st.button('🗑️ Eliminar (CSV y opcional archivo)', key='rc_del_btn', disabled=not is_admin):
                    st.session_state['rc_del_mode'] = True

            if is_admin and st.session_state.get('rc_del_mode', False):
                st.warning("⚠️ Eliminar contrato del repositorio (irreversible).")
                del_confirm = st.text_input("Escribe ELIMINAR para confirmar", key="rc_del_confirm")
                del_file = st.checkbox("Borrar también el archivo físico en generados/", value=False, key="rc_del_file")

                if st.button("✅ Confirmar eliminación", key="rc_del_do", disabled=del_confirm.strip().upper() != "ELIMINAR"):
                    repo2 = _repo_contratos_load()
                    idx2 = _repo_get_idx(repo2, sel_id)
                    ruta_del = ""
                    if idx2 is not None:
                        ruta_del = str(repo2.at[idx2,'Ruta'])
                    repo2 = repo2[repo2['ID'].astype(str) != str(sel_id)].copy()
                    _repo_contratos_save(repo2)

                    if del_file and ruta_del and os.path.exists(ruta_del):
                        try:
                            os.remove(ruta_del)
                        except Exception:
                            pass
                        if ruta_del.lower().endswith(".txt"):
                            docx_path = ruta_del[:-4] + ".docx"
                            if os.path.exists(docx_path):
                                try:
                                    os.remove(docx_path)
                                except Exception:
                                    pass

                    st.session_state['rc_del_mode'] = False
                    st.success("✅ Eliminado")
                    st.rerun()

            # ===== Descarga + Vista previa =====
            st.divider()
            ruta = str(row.get('Ruta',''))
            if str(row.get('Existe','')) == '1' and ruta and os.path.exists(ruta):
                ext = os.path.splitext(ruta)[1].lower()

                if ext == '.txt':
                    try:
                        with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
                            txt_prev = f.read()
                        with st.expander("👁️ Vista previa del contrato (TXT)", expanded=False):
                            st.text_area("Contenido", value=txt_prev, height=260, key="rc_preview_txt")
                    except Exception:
                        pass

                mime = 'text/plain' if ext == '.txt' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                with open(ruta, 'rb') as f:
                    data = f.read()

                st.download_button(
                    '⬇️ Descargar archivo',
                    data=data,
                    file_name=os.path.basename(ruta),
                    mime=mime,
                    key=f'rc_dl_{sel_id}'
                )

                if ext == '.txt':
                    docx_path = ruta[:-4] + '.docx'
                    if os.path.exists(docx_path):
                        with open(docx_path, 'rb') as f:
                            d2 = f.read()
                        st.download_button(
                            '⬇️ Descargar Word (.docx)',
                            data=d2,
                            file_name=os.path.basename(docx_path),
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                            key=f'rc_dl_docx_{sel_id}'
                        )
            else:
                st.warning('Archivo no encontrado en generados/.')
